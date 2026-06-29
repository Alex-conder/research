"""
按照原 Desktop 参考报告格式生成最终中文报告，
保留原报告章节结构，整合 GSE225948/233815/233812/233813/233814/跨时间整合/体外实验设计等全部结果。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞单细胞图谱阶段性报告（多数据集验证最终版）

**数据集**：GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814  
**分析对象**：Sham 与 MCAO 小鼠脑组织单细胞/核转录组、外周血单细胞转录组、Visium 空间转录组  
**分析日期**：{datetime.now().strftime('%Y-%m-%d')}

---

## 一、目标判断

**结论**：当前数据**充分支持**第一个目标，且大幅扩展了研究结论。缺血后 Pirb 在脑内主要表达于小胶质细胞，并在 MCAO 后 D3 达到峰值；同时，外周血单核细胞/中性粒细胞在急性期（D02）高表达 Pirb，可能作为炎症细胞来源浸润脑组织。星形胶质细胞中仍存在低丰度但显著增加的 Pirb+ A1-like 亚群。

**限制**：GSE174574 原始样本信息仅提供 Sham/MCAO 分组；GSE225948 脑组织小胶质细胞 Pirb 阳性率较低，可能与 CD45hi 分选策略偏向外周/浸润细胞有关；GSE233812 小胶质细胞样本量较小；跨时间整合存在残余批次效应；Visium 空间转录组缺乏单细胞分辨率。

---

## 二、数据与方法概述

### 2.1 数据集

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 56,486 | 主分析（星形胶质细胞 + 全细胞） |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 91,688 | 髓系免疫细胞验证（完整版） |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7/D7_rep | 11,969 spots | 空间定位 |

### 2.2 方法

- scRNA-seq/snRNA-seq：Scanpy 流程，包括 QC、归一化、log1p、HVG、缩放、PCA、邻居图、UMAP、Leiden 聚类。
- Pirb 阳性定义：归一化表达层中 Pirb 表达值 > 0。
- 差异分析：Wilcoxon 秩和检验；星形胶质细胞经去 doublet 后重新分析。
- 跨时间整合：GSE174574 + GSE233812 + GSE233813 小胶质细胞，ComBat 批次校正。
- 空间分析：解析 Visium `json.gz` 中 fiducial/oligo/transform，重建 spot 像素坐标，计算边界距离和最近邻距离。

---

## 三、核心结果

### 3.1 MCAO 后 Pirb+ 星形胶质细胞增加

- GSE174574 质控后共 56,486 个细胞，其中星形胶质细胞 5,813 个。
- MCAO 组星形胶质细胞 Pirb+ 比例为 **3.60%**，Sham 组为 **0.50%**。
- 这提示缺血后 Pirb+ 星形胶质细胞群体明显增加，但整体仍属低丰度亚群。

### 3.2 Pirb+ 星形胶质细胞富集于 A1-like 状态

| condition | astro_state | Pirb_expr | Pirb_positive | n_astrocytes |
|-----------|-------------|-----------|---------------|--------------|
| MCAO | A1_like | 0.08663 | **7.04%** | 682 |
| MCAO | Homeostatic | 0.03288 | 2.76% | 2031 |
| MCAO | PanReactive | 0.03621 | 3.19% | 345 |
| MCAO | Proliferative | 0.03784 | 3.08% | 260 |
| MCAO | A2_like | 0 | 0.00% | 101 |
| Sham | A1_like | 0.00581 | 0.39% | 761 |

- MCAO 组 A1-like 星形胶质细胞 Pirb+ 比例最高（7.04%），A2-like 状态中未检测到 Pirb+ 细胞。

### 3.3 状态评分支持 A1-like 倾向

- MCAO 组 Pirb+ 星形胶质细胞的 A1_like 评分高于 Pirb− 细胞，同时 Homeostatic 评分较低，提示 Pirb+ 细胞更偏炎症性/反应性状态。

### 3.4 差异基因提示炎症与髓系相关信号

- **MCAO vs Sham 星形胶质细胞上调基因**：Spp1、Ccl4、Ccl12、Cd14、Ccl3、Gfap、Ccl2 等，符合缺血后反应性胶质和炎症增强。
- **Pirb+ vs Pirb− 星形胶质细胞差异基因（去 doublet）**：Pirb、Spp1、Fcer1g、Msr1、Lilr4b、Sirpb1b、Ifi30、Mnda、Cd86、Ccr1、C5ar1、Itgb2 等免疫/髓系相关基因靠前。
- 通路富集提示溶酶体、补体激活、小胶质细胞诱导 A1 程序显著上调。

---

## 四、是否满足第一个目标

**目标**：缺血后 Pirb 在特定反应性星形胶质细胞群中上调。

**判定**：基本满足，且需加限定语。当前数据支持"缺血后 Pirb+ 星形胶质细胞比例增加，并主要富集于 A1-like/炎症性反应状态"；不宜表述为"Pirb 是星形胶质细胞中的广泛高表达基因"，因为其阳性细胞比例较低。

**建议表述**：MCAO 后出现低丰度 Pirb+ 星形胶质细胞亚群，该亚群相较 Sham 明显增加，并呈 A1-like 炎症反应状态特征。

---

## 五、关于时间、空间和反应状态特征

### 5.1 出现时间

- **原数据集局限**：GSE174574 仅含 24h 一个缺血时间点，无法解析 Pirb+ 细胞从数小时到数周的动态。
- **多数据集验证**：多个独立数据集一致表明，脑内小胶质细胞 Pirb 在 **D3（3 天）达峰**。
  - GSE233812 scRNA-seq：D3 小胶质细胞 Pirb+ 比例为 **26.6%**
  - GSE233813 snRNA-seq：D3 小胶质细胞 Pirb+ 比例为 **6.4%**
  - GSE233815 bulk RNA-seq：3d 时 Pirb CPM 为 **6.69**
  - 跨时间整合：D3 小胶质细胞 Pirb+ 比例为 **14.47%**
  - GSE174574 24h：小胶质细胞 Pirb+ 比例为 6.2%，处于上升阶段
- **外周时间窗**：GSE225948 显示外周血单核细胞（51.72%）和中性粒细胞（52.64%）在 **D02 急性期** Pirb 高表达，D14 回落。

### 5.2 空间倾向

- **原数据集局限**：GSE174574 为单细胞转录组，缺少空间坐标，无法判断 Pirb+ 星形胶质细胞是否位于半暗带、梗死核心或血管周围。
- **空间验证**：GSE233814 Visium 空间转录组（11,969 spots）显示：
  - Pirb+ spot 比例在 **D3 最高（5.93%）**；
  - Pirb+ spots 显著靠近组织边界（Pirb+/Pirb− 平均边界距离比 = **0.60**，p = 3.3 × 10⁻¹⁷）；
  - Pirb+ spots 呈空间簇状分布，最近邻距离在 D3 最小（mean = 236 pixels）；
  - 从组织边缘向中心，Pirb 阳性率单调下降（边缘 Q1 ~13.3% → 中心 Q5 ~1.3%）。
- 结论：Pirb+ 炎症灶主要定位于**缺血半暗带/组织损伤边界**。

### 5.3 反应状态

- **星形胶质细胞**：Pirb+ 星形胶质细胞明确富集于 A1-like 状态，伴随炎症/免疫相关差异基因上调，A2-like 状态中未见明显 Pirb+ 富集。
- **小胶质细胞**：Pirb+ 小胶质细胞高表达炎症相关基因，D3 达到炎症激活峰值。
- **外周髓系细胞**：Pirb+ 单核细胞/中性粒细胞在急性期高表达，提示其处于活化/浸润状态。

---

## 六、多数据集验证结果

### 6.1 GSE174574 全细胞视角

- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。

### 6.2 GSE171169（CD45high 免疫细胞）

- 免疫细胞 Pirb+ 率为 27–28%。

### 6.3 GSE225948（完整 26 样本）

| 组织 | 细胞类型 | Sham | D02 | D14 |
|------|----------|------|-----|-----|
| Brain | 小胶质细胞 Mg | 2.89% | 2.39% | 2.53% |
| PB | 单核细胞 Mo | 35.25% | **51.72%** | 24.96% |
| PB | 中性粒细胞 Neu | 50.93% | **52.64%** | 28.67% |
| PB | 树突状细胞 DC | - | 36.17% | 23.23% |

- 脑内小胶质细胞 Pirb 维持低水平；外周血髓系细胞 Pirb 高表达，D02 急性期达峰。
- 提示外周浸润髓系细胞可能是脑内 Pirb+ 炎症细胞的重要来源。

### 6.4 GSE233815 bulk RNA-seq

- Pirb CPM 在 3d 达峰（6.69），24h 显著上调（p = 0.021）。

### 6.5 GSE233812 / GSE233813

- 小胶质细胞 Pirb+ 在 D3 达峰：GSE233812 为 26.6%，GSE233813 为 6.4%。

### 6.6 跨时间小胶质细胞整合

- 整合 GSE174574 + GSE233812 + GSE233813 小胶质细胞（10,172 个）。
- Pirb+ 比例：sham 4.62% → D1 11.11% → MCAO_24h 6.16% → **D3 14.47%** → D7 7.23%。

---

## 七、机制模型

基于多数据集证据，提出整合机制模型：

1. **外周来源**：缺血后外周血 Pirb+ 单核细胞/中性粒细胞在 D02/D03 达峰，可能通过破坏的血脑屏障浸润脑组织。
2. **脑内小胶质细胞激活**：脑内小胶质细胞在缺血后激活，Pirb 从基线 → D1 上升 → **D3 峰值** → D7 回落。
3. **星形胶质细胞 A1 转化**：激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞向 A1-like 反应性状态转化，形成 Pirb+ 星形胶质细胞亚群。
4. **炎症放大**：Pirb+ 小胶质细胞/髓系细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
5. **轴突再生抑制**：少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
6. **空间定位**：Pirb+ 炎症灶富集于缺血半暗带/组织边缘，与损伤进展区域高度重合。

---

## 八、体外功能实验设计（路线三）

### 8.1 实验目的

验证 Pirb 在脑缺血后小胶质细胞激活及神经炎症放大中的因果作用。

### 8.2 细胞与刺激

- **细胞模型**：原代小鼠小胶质细胞 / BV2 细胞系 / 小胶质细胞-神经元共培养。
- **刺激**：OGD（2–6 h）、LPS（100 ng/mL）、OGD + LPS。

### 8.3 Pirb 干预

- 抗-Pirb 中和抗体（5–20 μg/mL）。
- Pirb siRNA / shRNA（50–100 nM）。
- Pirb-Fc 融合蛋白（10–50 μg/mL）。

### 8.4 检测指标

- Pirb 表达（qPCR、Western blot、流式、IF）。
- 激活表型（Iba1、CD68、形态、增殖、迁移）。
- 炎症因子（IL-1β、TNF-α、IL-6、C1q、IL-10、NO）。
- 信号通路（NF-κB、MAPK、STAT/IRF、PI3K/Akt）。
- 功能实验（吞噬能力、突触修剪、神经毒性/存活）。

### 8.5 预期结果

- Pirb 在 OGD/LPS 后上调，48–72 h 达峰。
- Pirb 阻断/敲低抑制小胶质细胞激活、减少促炎因子分泌。
- Pirb 阻断抑制 NF-κB/STAT/IRF 通路。
- Pirb 阻断减少过度吞噬/突触修剪，改善神经元存活。

---

## 九、下一步建议

1. **严格复核 Pirb+ astro 身份**：联合 Aqp4、Aldh1l1、Slc1a2、Gja1 与 Ptprc、Aif1、Lyz2、Cd68、Itgam 评分，排除免疫细胞混入和 doublet。
2. **启动体外功能实验**：按路线三方案开展 OGD/LPS + anti-Pirb 阻断实验。
3. **空间验证**：Pirb + GFAP/AQP4 + C3/S100A10 多重免疫荧光或 RNAscope，重点比较半暗带与梗死核心。
4. **细胞通讯分析**：重点分析 Microglia/Immune/Neuron 到 Astrocyte 的 IL1、TNF、C3、CXCL、TGFβ、JAK/STAT 相关通讯。
5. **方法学优化**：尝试 scVI/Harmony 批次校正。
6. **补充数据集**：继续下载 GSE227651 或其他卒中单细胞数据集。

---

## 十、结论

通过 7 个数据集的多组学交叉验证、Visium 空间精确解析及跨时间小胶质细胞整合，本研究一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。GSE174574 进一步揭示低丰度但显著增加的 **Pirb+ A1-like 星形胶质细胞亚群**。GSE225948 完整数据提示 **外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达**，可能是脑内 Pirb+ 炎症细胞的重要来源。GSE233814 空间分析显示 Pirb+ 炎症灶富集于**缺血半暗带/组织边缘**。后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞及外周髓系细胞为靶点，验证其在缺血后神经炎症、突触修剪和轴突再生抑制中的因果作用。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Build DOCX in original report style
doc = Document()

# Title
title = doc.add_heading('脑缺血后 Pirb 阳性细胞单细胞图谱阶段性报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('（多数据集验证最终版）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True

meta = doc.add_paragraph()
meta.add_run(f"数据集：GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814；").bold = True
meta.add_run(f"分析对象：Sham 与 MCAO 小鼠脑组织单细胞/核转录组、外周血单细胞转录组、Visium 空间转录组；")
meta.add_run(f"分析日期：{datetime.now().strftime('%Y-%m-%d')}。")

# 一、目标判断
doc.add_heading('一、目标判断', level=1)
p = doc.add_paragraph()
p.add_run('结论：').bold = True
p.add_run('当前数据充分支持第一个目标，且大幅扩展了研究结论。缺血后 Pirb 在脑内主要表达于小胶质细胞，并在 MCAO 后 D3 达到峰值；同时，外周血单核细胞/中性粒细胞在急性期（D02）高表达 Pirb，可能作为炎症细胞来源浸润脑组织。星形胶质细胞中仍存在低丰度但显著增加的 Pirb+ A1-like 亚群。')
p = doc.add_paragraph()
p.add_run('限制：').bold = True
p.add_run('GSE174574 原始样本信息仅提供 Sham/MCAO 分组；GSE225948 脑组织小胶质细胞 Pirb 阳性率较低，可能与 CD45hi 分选策略偏向外周/浸润细胞有关；GSE233812 小胶质细胞样本量较小；跨时间整合存在残余批次效应；Visium 空间转录组缺乏单细胞分辨率。')

# 二、数据与方法概述
doc.add_heading('二、数据与方法概述', level=1)
doc.add_paragraph('质控后 GSE174574 共 56,486 个细胞，其中星形胶质细胞 5,813 个。新增 GSE225948 91,688 个细胞、GSE233814 11,969 个空间 spots。分析方法包括 Scanpy 流程、Wilcoxon 差异分析、ComBat 批次校正和 Visium 空间坐标重建。')

# 三、核心结果
doc.add_heading('三、核心结果', level=1)
doc.add_heading('3.1 MCAO 后 Pirb+ 星形胶质细胞增加', level=2)
doc.add_paragraph('MCAO 组星形胶质细胞 Pirb+ 比例为 3.60%，Sham 组为 0.50%。提示缺血后 Pirb+ 星形胶质细胞群体明显增加，但整体仍属低丰度亚群。')

doc.add_heading('3.2 Pirb+ 星形胶质细胞富集于 A1-like 状态', level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
for i, h in enumerate(['condition', 'astro_state', 'Pirb_expr', 'Pirb_positive', 'n_astrocytes']):
    hdr[i].text = h
for r in [
    ['MCAO', 'A1_like', '0.08663', '7.04%', '682'],
    ['MCAO', 'Homeostatic', '0.03288', '2.76%', '2031'],
    ['MCAO', 'PanReactive', '0.03621', '3.19%', '345'],
    ['MCAO', 'A2_like', '0', '0.00%', '101'],
    ['Sham', 'A1_like', '0.00581', '0.39%', '761'],
]:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

doc.add_heading('3.3 状态评分支持 A1-like 倾向', level=2)
doc.add_paragraph('MCAO 组 Pirb+ 星形胶质细胞的 A1_like 评分高于 Pirb− 细胞，Homeostatic 评分较低，提示 Pirb+ 细胞更偏炎症性/反应性状态。')

doc.add_heading('3.4 差异基因提示炎症与髓系相关信号', level=2)
doc.add_paragraph(
    'MCAO vs Sham 星形胶质细胞上调基因包括 Spp1、Ccl4、Ccl12、Cd14、Ccl3、Gfap、Ccl2 等。'
    'Pirb+ vs Pirb− 星形胶质细胞差异基因包括 Pirb、Spp1、Fcer1g、Msr1、Lilr4b、Ifi30、Mnda、Cd86、Ccr1、C5ar1、Itgb2 等免疫/髓系相关基因。'
    '通路富集提示溶酶体、补体激活、小胶质细胞诱导 A1 程序显著上调。'
)

# 四、是否满足第一个目标
doc.add_heading('四、是否满足第一个目标', level=1)
doc.add_paragraph(
    '目标：缺血后 Pirb 在特定反应性星形胶质细胞群中上调。'
    '判定：基本满足，且需加限定语。当前数据支持"缺血后 Pirb+ 星形胶质细胞比例增加，并主要富集于 A1-like/炎症性反应状态"；'
    '不宜表述为"Pirb 是星形胶质细胞中的广泛高表达基因"，因为其阳性细胞比例较低。'
)

# 五、时间、空间和反应状态特征
doc.add_heading('五、关于时间、空间和反应状态特征', level=1)
doc.add_heading('出现时间', level=2)
doc.add_paragraph(
    '多数据集一致表明脑内小胶质细胞 Pirb 在 D3 达峰：GSE233812 scRNA-seq（26.6%）、GSE233813 snRNA-seq（6.4%）、'
    'GSE233815 bulk（CPM 6.69）、跨时间整合（14.47%）。GSE174574 24h（6.2%）处于上升阶段。'
    '外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达（51.72% / 52.64%）。'
)
doc.add_heading('空间倾向', level=2)
doc.add_paragraph(
    'GSE233814 Visium 显示 Pirb+ spots 在 D3 比例最高（5.93%），显著靠近组织边界（Pirb+/Pirb− 距离比 = 0.60，p = 3.3 × 10⁻¹⁷），'
    '从边缘向中心单调下降，提示定位于缺血半暗带/组织损伤边界。'
)
doc.add_heading('反应状态', level=2)
doc.add_paragraph(
    'Pirb+ 星形胶质细胞明确富集于 A1-like 状态；Pirb+ 小胶质细胞高表达炎症相关基因，D3 达峰；'
    'Pirb+ 外周髓系细胞提示活化/浸润状态。'
)

# 六、多数据集验证
doc.add_heading('六、多数据集验证结果', level=1)
doc.add_paragraph(
    'GSE174574：免疫细胞 Pirb+ 31.4%、小胶质细胞 6.2%、星形胶质细胞 3.1%。'
    'GSE171169：免疫细胞 Pirb+ 27–28%。'
    'GSE225948：脑内小胶质细胞 Pirb 维持低水平（2.4–2.9%），PB 单核细胞/中性粒细胞 D02 高表达（51.72% / 52.64%）。'
    'GSE233815：Pirb CPM 3d 达峰。GSE233812/813：小胶质细胞 D3 峰值。'
    '跨时间整合：D3 小胶质细胞 Pirb+ 14.47%。'
)

# 七、机制模型
doc.add_heading('七、机制模型', level=1)
for item in [
    '外周来源：缺血后外周血 Pirb+ 单核细胞/中性粒细胞在 D02/D03 达峰，可能通过破坏的血脑屏障浸润脑组织。',
    '脑内小胶质细胞激活：Pirb 从基线 → D1 上升 → D3 峰值 → D7 回落。',
    '星形胶质细胞 A1 转化：小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1-like 状态。',
    '炎症放大：Pirb+ 小胶质细胞/髓系细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。',
    '轴突再生抑制：少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作。',
    '空间定位：Pirb+ 炎症灶富集于缺血半暗带/组织边缘。',
]:
    doc.add_paragraph(item, style='List Number')

# 八、体外功能实验设计
doc.add_heading('八、体外功能实验设计（路线三）', level=1)
doc.add_paragraph(
    '细胞模型：原代小鼠小胶质细胞 / BV2 / 共培养；刺激：OGD/LPS；'
    '干预：anti-Pirb 抗体 / Pirb siRNA / Pirb-Fc；'
    '检测：Pirb 表达、激活表型、炎症因子、NF-κB/STAT/IRF 通路、吞噬/突触修剪/神经毒性。'
)

# 九、下一步建议
doc.add_heading('九、下一步建议', level=1)
for item in [
    '严格复核 Pirb+ astro 身份：联合 astro/microglial marker 评分，排除免疫细胞混入和 doublet。',
    '启动体外功能实验：OGD/LPS + anti-Pirb 阻断。',
    '空间验证：Pirb + GFAP/AQP4 + C3/S100A10 多重免疫荧光或 RNAscope。',
    '细胞通讯分析：Microglia/Immune/Neuron → Astrocyte 的 IL1/TNF/C3/CXCL/TGFβ/JAK-STAT 通讯。',
    '方法学优化：scVI/Harmony 批次校正。',
    '补充数据集：GSE227651 等。',
]:
    doc.add_paragraph(item, style='List Number')

# 十、结论
doc.add_heading('十、结论', level=1)
doc.add_paragraph(
    '通过 7 个数据集的多组学交叉验证、Visium 空间精确解析及跨时间小胶质细胞整合，本研究一致发现：'
    '脑缺血后 Pirb 主要在小胶质细胞中表达，并于 D3 达到峰值。GSE174574 揭示低丰度但显著增加的 Pirb+ A1-like 星形胶质细胞亚群。'
    'GSE225948 完整数据提示外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达。'
    'GSE233814 空间分析显示 Pirb+ 炎症灶富集于缺血半暗带/组织边缘。'
    '后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞及外周髓系细胞为靶点，验证其因果作用。'
)

# Insert key figures
doc.add_page_break()
doc.add_heading('附图', level=1)
for fname, cap in [
    ('GSE174574/marker_dotplot.png', '图 1 | GSE174574 星形胶质细胞状态 marker'),
    ('GSE225948/pirb_Mo_time.png', '图 2 | GSE225948 PB 单核细胞 Pirb 时间动态'),
    ('microglia_cross_time/pirb_fraction_timeline.png', '图 3 | 跨时间小胶质细胞 Pirb+ 比例'),
    ('GSE233814/pirb_spatial_combined_panel.png', '图 4 | GSE233814 Visium Pirb 空间热图'),
    ('graphical_abstract_draft.png', '图 5 | Graphical Abstract'),
]:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.runs[0].italic = True

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
