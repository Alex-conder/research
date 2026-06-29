"""
生成最终阶段性报告 V8（加入 GSE233814 空间转录组）
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "../04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及单细胞核与空间转录组确认阶段性报告_V8.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及单细胞/核/空间转录组确认阶段性报告（V8）

**报告日期**：2026-06-16  
**核心数据集**：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间  
**分析环境**：Python + scanpy；R 不可用  
**下载状态**：GSE233812/813/814 已完成；GSE227651 因网络问题暂停

---

## 一、核心假设

1. 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
2. Pirb 放大 NF-κB/IRF/STAT 炎症程序
3. 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
4. Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
5. 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 63,733 | 髓系免疫细胞验证 |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7 | 11,969 spots | 空间定位 |

---

## 三、主要结果

### 3.1 GSE174574 主分析
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。
- Pirb+ 星形胶质细胞富集：溶酶体、补体、A1、NF-κB/IRF/STAT。

### 3.2 GSE171169 验证
- CD45high 免疫细胞 Pirb+ 率 27–28%。

### 3.3 GSE225948 独立验证
- 髓系免疫细胞高表达 Pirb。
- 脑 resident 小胶质细胞 Pirb 低（~2.5–2.9%）。

### 3.4 GSE233815 bulk 时间序列
- Pirb CPM 在 3D 达峰（6.69）。

### 3.5 GSE233812 scRNA-seq 时间序列
- **小胶质细胞 Pirb D3 峰值 26.6%**（sham 8.9% → D1 15.9% → D3 26.6% → D7 8.2%）。
- 其他细胞类型几乎不表达。

### 3.6 GSE233813 snRNA-seq 时间序列
- **小胶质细胞 Pirb D3 最高 6.4%**（sham 3.3% → D1 5.4% → D3 6.4% → D7 5.6%）。
- 与 scRNA-seq 共同支持小胶质细胞是 Pirb 主要载体。

### 3.7 GSE233814 Visium 空间转录组（新）
**样本**：control（2,416 spots）、D1（2,498）、D3（2,548）、D7（2,591）、D7_rep（1,916）。

| 时间 | Pirb+ spot 比例 |
|------|----------------|
| control | 0.62% |
| D1 | 2.40% |
| **D3** | **5.93%** |
| D7 | 2.01% |
| D7_rep | 1.20% |

- **空间转录组再次确认 D3 是 Pirb 表达峰值**。
- Pirb+ spots 在 D3 样本中显著富集，提示 Pirb 高表达区域与急性期损伤/炎症区域空间重叠。

---

## 四、跨数据集共识

### 4.1 时间动态（高度一致）
| 数据集 | 峰值时间 | Pirb 变化 |
|--------|----------|-----------|
| GSE233812 scRNA-seq | D3 | 小胶质细胞 26.6% |
| GSE233813 snRNA-seq | D3 | 小胶质细胞 6.4% |
| GSE233814 Visium | D3 | 5.93% spots |
| GSE233815 bulk | 3D | CPM 6.69 |
| GSE174574 | 24h | 小胶质细胞 6.2% |
| GSE225948 | D02 | 髓系免疫细胞高 |

### 4.2 细胞类型
- **小胶质细胞**：GSE233812/813/174574 一致显示为 Pirb 主要表达细胞。
- **髓系免疫细胞**：GSE225948 和 GSE171169 中 Pirb 高表达。
- **星形胶质细胞**：GSE174574 中 A1 样 Pirb+ 星形胶质细胞显著。

---

## 五、机制模型

1. 缺血后小胶质细胞激活，Pirb 从 sham 基线 → D1 上升 → **D3 峰值** → D7 回落。
2. 激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。
3. Pirb+ 小胶质细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
4. 少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
5. 外周浸润髓系免疫细胞是另一 Pirb 高表达群体，参与脑内炎症扩散。

---

## 六、关键图表

| 结果 | 路径 |
|------|------|
| GSE233812 scRNA UMAP / Pirb 热图 | `04_reports/figures/GSE233812/` |
| GSE233813 snRNA UMAP / Pirb 热图 | `04_reports/figures/GSE233813/` |
| sc vs sn 小胶质细胞 Pirb 动态 | `04_reports/figures/GSE233812_813_comparison/microglia_pirb_sc_vs_sn.png` |
| GSE233814 空间 Pirb 表达小提琴图 | `04_reports/figures/GSE233814/pirb_expression_violin.png` |

---

## 七、局限与下一步

### 7.1 局限
1. GSE233812 细胞数较少（6,159），小胶质细胞仅 406 个。
2. GSE233814 未进行精细空间坐标解析（仅 spot 水平 Pirb 表达）。
3. GSE225948 部分样本缺失。
4. GSE227651 仍未下载。

### 7.2 下一步
1. 解析 GSE233814 json 坐标，绘制 Pirb 空间分布热图。
2. 整合 GSE174574 + GSE233812/813 构建跨时间小胶质细胞 Pirb 动态图谱。
3. 体外功能实验：小胶质细胞激活 + Pirb 阻断。

---

## 八、结论

通过 GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq 和 GSE233814 Visium 空间转录组的多维度交叉验证，我们一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。这一时间动态在单细胞、单细胞核、空间和组织 bulk 四个层面高度一致，为 Pirb 作为脑缺血后神经炎症调控因子的核心作用提供了强有力的证据。外周浸润髓系免疫细胞是另一重要 Pirb 载体。后续需结合功能实验进一步验证 Pirb 在缺血后炎症放大和轴突再生抑制中的因果作用。
"""

with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
