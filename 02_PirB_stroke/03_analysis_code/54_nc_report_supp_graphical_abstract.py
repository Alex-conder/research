"""
1. 生成 Supplementary Tables (Excel)
2. 生成 Nature Communications 格式的完整报告 (MD + DOCX)
3. 生成 Graphical Abstract 草图 (PNG)
"""
import os, sys
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Circle, FancyArrowPatch

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
SUP_DIR = "D:/Pirb_stroke_project/04_reports/supplementary"
os.makedirs(SUP_DIR, exist_ok=True)

SUP_XLSX = os.path.join(SUP_DIR, "Supplementary_Tables_Pirb_Stroke.xlsx")
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.docx")
GA_PATH = os.path.join(REPORT_DIR, "figures", "graphical_abstract_draft.png")

# ------------------- Supplementary Tables -------------------
def load_de(path, top_n=200, pcol='pvals_adj', lfc='logfoldchange'):
    df = pd.read_csv(path)
    # standardize columns
    for c in ['pvals_adj', 'pval_adj']:
        if c in df.columns:
            df['pvals_adj'] = df[c]
    for c in ['logfoldchange']:
        if c in df.columns:
            df['log2FC'] = df[c]
    # sort
    df = df.sort_values(['pvals_adj', 'log2FC'], ascending=[True, False])
    df = df.head(top_n)
    cols = ['gene', 'log2FC', 'pvals_adj']
    avail = [c for c in cols if c in df.columns]
    return df[avail].reset_index(drop=True)

# Read tables
de_astro = load_de('04_reports/figures/GSE174574/doublet_qc/DE_PirbPos_vs_Neg_Astrocyte_no_doublet.csv', top_n=300)
de_micro = load_de('04_reports/figures/GSE174574/de_pirb/DE_PirbPos_vs_Neg_Microglia.csv', top_n=300)
de_imm = load_de('04_reports/figures/GSE174574/de_pirb/DE_PirbPos_vs_Neg_Immune.csv', top_n=300)
de_cross = pd.read_csv('04_reports/figures/microglia_cross_time/pirb_pos_vs_neg_microglia_de.csv')
pathway = pd.read_csv('04_reports/figures/GSE174574/de_pirb/manual_pathway_enrichment.csv')
spatial_summary = pd.read_csv('04_reports/figures/GSE233814/pirb_spatial_pattern_summary.csv')

# Dataset summary
dataset_summary = pd.DataFrame({
    'Dataset': ['GSE174574', 'GSE171169', 'GSE225948', 'GSE233815', 'GSE233812', 'GSE233813', 'GSE233814'],
    'Data_type': ['scRNA-seq', 'scRNA-seq', 'scRNA-seq', 'bulk RNA-seq', 'scRNA-seq', 'snRNA-seq', 'Visium spatial'],
    'Species_Model': ['Mouse MCAO', 'Mouse MCAO CD45high', 'Mouse tMCAO PB+brain', 'Mouse tMCAO', 'Mouse tMCAO', 'Mouse tMCAO', 'Mouse tMCAO'],
    'Time_points': ['Sham, 24h', '5d, 14d', 'Sham, D02, D14', '3h,12h,24h,3d,7d', 'sham,D1,D3,D7', 'sham,D1,D3,D7', 'control,D1,D3,D7,D7_rep'],
    'Cells_or_spots': ['56,486', '10,295', '91,688', '48 samples', '6,159', '8,374', '11,969 spots'],
    'Key_finding': [
        'A1-like astrocytes Pirb+ 7.04%; microglia Pirb+ 6.2%',
        'Immune cells Pirb+ ~27%',
        'PB Mo/Neu Pirb+ peaks D02; brain Mg remains low',
        'Pirb peaks at 3d (CPM 6.69)',
        'Microglia Pirb+ peaks D3 26.6%',
        'Microglia Pirb+ peaks D3 6.4%',
        'Pirb+ spots peak D3 5.93%; enriched at tissue boundary'
    ]
})

with pd.ExcelWriter(SUP_XLSX, engine='openpyxl') as writer:
    dataset_summary.to_excel(writer, sheet_name='ST1 Dataset summary', index=False)
    de_astro.to_excel(writer, sheet_name='ST2 Astro Pirb+ vs Pirb-', index=False)
    de_micro.to_excel(writer, sheet_name='ST3 Microglia Pirb+ vs Pirb-', index=False)
    de_imm.to_excel(writer, sheet_name='ST4 Immune Pirb+ vs Pirb-', index=False)
    de_cross.to_excel(writer, sheet_name='ST5 Cross-time microglia DE', index=False)
    pathway.to_excel(writer, sheet_name='ST6 Pathway enrichment', index=False)
    spatial_summary.to_excel(writer, sheet_name='ST7 Spatial summary', index=False)

print(f"[SAVE] {SUP_XLSX}")

# ------------------- Graphical Abstract -------------------
fig, ax = plt.subplots(figsize=(14, 8))
ax.set_xlim(0, 14)
ax.set_ylim(0, 8)
ax.axis('off')

# title
ax.text(7, 7.6, 'Pirb drives post-ischemic neuroinflammation', fontsize=18, ha='center', fontweight='bold')

# --- Panel 1: Healthy brain ---
ax.add_patch(FancyBboxPatch((0.5, 4.5), 3.2, 2.5, boxstyle="round,pad=0.05", facecolor='#e8f4f8', edgecolor='#333'))
ax.text(2.1, 6.7, 'Healthy brain', fontsize=12, ha='center', fontweight='bold')
ax.text(2.1, 6.3, 'Pirb low/undetectable', fontsize=10, ha='center', color='#555')
# neurons
c1 = Circle((1.2, 5.3), 0.3, color='#a6cee3')
c2 = Circle((2.5, 5.6), 0.3, color='#a6cee3')
c3 = Circle((1.8, 4.9), 0.3, color='#a6cee3')
ax.add_patch(c1); ax.add_patch(c2); ax.add_patch(c3)
ax.text(2.1, 4.6, 'Homeostatic microglia\nResting astrocytes', fontsize=9, ha='center')

# --- Panel 2: Ischemic core / penumbra ---
ax.add_patch(FancyBboxPatch((5.2, 4.5), 4.0, 2.5, boxstyle="round,pad=0.05", facecolor='#fff0f0', edgecolor='#d62728', linewidth=2))
ax.text(7.2, 6.7, 'Ischemic penumbra (D3 peak)', fontsize=12, ha='center', fontweight='bold', color='#d62728')
# Pirb+ microglia
ax.add_patch(Circle((6.0, 5.6), 0.35, color='#ff7f0e'))
ax.text(6.0, 5.6, 'Pirb+', ha='center', va='center', fontsize=8, color='white', fontweight='bold')
ax.text(6.0, 5.0, 'Microglia', ha='center', fontsize=9)
# Pirb+ astrocyte
ax.add_patch(Circle((7.5, 5.6), 0.35, color='#2ca02c'))
ax.text(7.5, 5.6, 'Pirb+', ha='center', va='center', fontsize=8, color='white', fontweight='bold')
ax.text(7.5, 5.0, 'A1 astrocyte', ha='center', fontsize=9)
# infiltrating monocyte/neutrophil
ax.add_patch(Circle((8.8, 5.6), 0.35, color='#9467bd'))
ax.text(8.8, 5.6, 'Pirb+', ha='center', va='center', fontsize=8, color='white', fontweight='bold')
ax.text(8.8, 5.0, 'Mo/Neu', ha='center', fontsize=9)
# arrows: infiltration
ax.annotate('', xy=(5.9, 5.6), xytext=(5.3, 6.0),
            arrowprops=dict(arrowstyle='->', color='#9467bd', lw=1.5))

# --- Panel 3: Mechanism ---
ax.add_patch(FancyBboxPatch((10.0, 4.5), 3.5, 2.5, boxstyle="round,pad=0.05", facecolor='#f5f5dc', edgecolor='#333'))
ax.text(11.75, 6.7, 'Molecular mechanism', fontsize=12, ha='center', fontweight='bold')
mech_items = [
    'IL-1β / TNF / C1q',
    'NF-κB / STAT / IRF',
    'Synapse pruning↑',
    'Axon regeneration↓'
]
for i, t in enumerate(mech_items):
    ax.text(11.75, 6.2 - i*0.35, f'• {t}', fontsize=10, ha='center')

# --- Timeline ---
ax.plot([1, 13], [3.8, 3.8], 'k-', lw=2)
for x, t in [(2, 'Sham'), (4.5, '6h'), (6.5, 'D1'), (8.5, 'D3'), (10.5, 'D7'), (12.5, 'D14')]:
    ax.plot([x, x], [3.7, 3.9], 'k-', lw=2)
    ax.text(x, 3.45, t, fontsize=10, ha='center')
ax.text(7, 3.15, 'Pirb expression peaks at D3 in microglia; PB Mo/Neu peaks at D02', fontsize=10, ha='center', style='italic')

# --- Legend ---
legend_elements = [
    mpatches.Patch(color='#ff7f0e', label='Pirb+ microglia'),
    mpatches.Patch(color='#2ca02c', label='Pirb+ A1 astrocyte'),
    mpatches.Patch(color='#9467bd', label='Pirb+ Mo/Neu infiltrate'),
]
ax.legend(handles=legend_elements, loc='lower left', bbox_to_anchor=(0.02, 0.02), fontsize=10)

plt.tight_layout()
plt.savefig(GA_PATH, dpi=300, bbox_inches='tight')
plt.close()
print(f"[SAVE] {GA_PATH}")

# ------------------- Nature Communications report -------------------
md_content = f"""# Pirb-positive cells orchestrate post-ischemic neuroinflammation: a multi-dataset single-cell and spatial transcriptomics study

**Authors**: [Author names]¹,²,*  
**Affiliations**: ¹ Department, Institution, City, Country; ² Department, Institution, City, Country  
**Corresponding author**: *email@institution.edu  
**Date**: {datetime.now().strftime('%Y-%m-%d')}

---

## Abstract

**Background**: Paired immunoglobulin-like receptor B (Pirb, mouse ortholog of human LILRB2) is an immune inhibitory receptor implicated in neuroinflammation and axon regeneration failure, yet its cell-type-specific dynamics after cerebral ischemia remain poorly defined.  
**Results**: Integrating seven public datasets (GSE174574, GSE171169, GSE225948, GSE233815, GSE233812, GSE233813, GSE233814), we show that Pirb is sharply upregulated in microglia and infiltrating peripheral myeloid cells after experimental stroke, peaking at post-ischemic day 3 (D3). In GSE174574, Pirb-positive astrocytes were enriched in the A1-like reactive state (7.04% vs 0.39% in sham). Cross-dataset microglial integration confirmed a D3 peak (14.47%). GSE225948 revealed high Pirb expression in peripheral blood monocytes (51.72%) and neutrophils (52.64%) at acute D02, while brain microglia remained low (~2.4–2.9%). Visium spatial transcriptomics localized D3 Pirb-expressing spots to the ischemic boundary/penumbra (Pirb+ spots 0.60-fold closer to tissue edge, p = 3.3 × 10⁻¹⁷).  
**Conclusions**: Pirb marks a transient, spatially restricted inflammatory cell population after stroke. Our findings position Pirb as a candidate therapeutic target for modulating post-ischemic neuroinflammation and provide a roadmap for in vitro functional validation.

**Keywords**: Pirb, LILRB2, stroke, microglia, astrocyte, neuroinflammation, single-cell RNA sequencing, spatial transcriptomics

---

## Highlights

- Pirb expression peaks at post-ischemic day 3 in brain microglia across three independent datasets.
- Pirb-positive reactive astrocytes adopt an A1-like inflammatory state after MCAO.
- Peripheral blood monocytes and neutrophils show acute Pirb upregulation and may infiltrate the injured brain.
- Visium spatial transcriptomics places Pirb-positive inflammatory foci at the ischemic penumbra boundary.
- A functional validation roadmap targeting Pirb in microglia is proposed.

---

## Introduction

Cerebral ischemia triggers a rapid neuroinflammatory response in which resident microglia and infiltrating peripheral myeloid cells release cytokines, phagocytose synapses, and shape tissue repair... [placeholder for extended introduction]

Pirb (paired immunoglobulin-like receptor B) and its human orthologs LILRB1/2 are immune inhibitory receptors expressed by myeloid cells... In the central nervous system, Pirb/LILRB signaling has been linked to myelin-mediated axon growth inhibition, microglial activation, and synapse elimination.

Here, we integrate seven independent mouse stroke datasets to (i) map the cell-type-specific expression of Pirb, (ii) define its temporal trajectory, (iii) localize Pirb-expressing cells in situ using Visium spatial transcriptomics, and (iv) propose an in vitro functional validation framework.

---

## Results

### Pirb is induced in A1-like reactive astrocytes after MCAO

Using GSE174574 (Sham vs MCAO 24 h), we identified 5,813 astrocytes among 56,486 high-quality cells. Pirb-positive astrocytes were rare but significantly enriched after MCAO (3.60% vs 0.50% in sham). State scoring revealed that Pirb+ astrocytes were concentrated in the A1-like reactive state (7.04% Pirb+ in MCAO A1-like vs 0.39% in sham A1-like; Supplementary Table 2). Differentially expressed genes in Pirb+ astrocytes included immune/microglial markers such as Cd14, Fcer1g, Msr1, Lilr4b, and C5ar1 (Supplementary Table 2), and pathway analysis implicated lysosome, complement activation, and microglia-induced A1 astrocyte programs (Supplementary Table 6).

### Temporal dynamics of Pirb across independent datasets

Multiple datasets converged on a D3 peak for Pirb expression in brain microglia: GSE233812 scRNA-seq (26.6% Pirb+ microglia at D3), GSE233813 snRNA-seq (6.4%), GSE233815 bulk RNA-seq (CPM 6.69 at 3 d), and cross-dataset integration of GSE174574 + GSE233812 + GSE233813 (14.47% Pirb+ microglia at D3). GSE174574 at 24 h showed 6.2% Pirb+ microglia, consistent with an ascending phase toward the D3 peak.

### Peripheral myeloid cells are a major Pirb-expressing compartment

Analysis of the complete 26-sample GSE225948 dataset (91,688 cells after QC) showed that brain microglia maintained low Pirb expression across sham, D02, and D14 (~2.4–2.9%). In contrast, peripheral blood monocytes and neutrophils displayed high acute Pirb expression that peaked at D02 (monocytes 51.72%; neutrophils 52.64%) and declined by D14 (Supplementary Table 7). These findings suggest that circulating Pirb+ myeloid cells may infiltrate the ischemic brain and contribute to the inflammatory Pirb+ pool.

### Spatial localization of Pirb expression to the ischemic boundary

GSE233814 Visium spatial transcriptomics (11,969 spots) showed that Pirb-positive spot fraction peaked at D3 (5.93%) relative to control (0.62%), D1 (2.40%), and D7 (~2%). Pirb+ spots were spatially clustered (nearest-neighbor distance mean 236 pixels at D3) and significantly closer to the tissue boundary than Pirb− spots (mean distance ratio Pirb+/Pirb− = 0.60, p = 3.3 × 10⁻¹⁷), consistent with localization to the ischemic penumbra. Pirb+ spots displayed elevated microglial and inflammatory module scores (Supplementary Table 7).

### An integrated model of Pirb-mediated neuroinflammation

We propose an integrated model in which (1) ischemia acutely mobilizes Pirb+ peripheral monocytes/neutrophils; (2) resident microglia upregulate Pirb, peaking at D3; (3) activated microglia secrete IL-1α/TNF/C1q, driving astrocytes toward an A1-like Pirb+ state; (4) Pirb+ myeloid/astroglial cells amplify neuroinflammation via NF-κB/STAT/IRF programs; and (5) oligodendrocyte-derived Pirb ligands (MAG/OMgp) further suppress axon regeneration.

---

## Discussion

Our multi-dataset analysis provides convergent evidence that Pirb is transiently upregulated in microglia and peripheral myeloid cells after cerebral ischemia and that Pirb+ inflammatory foci localize to the ischemic boundary. The D3 microglial peak aligns with the established delayed neuroinflammatory phase, while the D02 peripheral peak identifies a window for early immune intervention.

The low-frequency but highly inflammatory Pirb+ A1 astrocyte population observed in GSE174574 suggests that Pirb may mark a subset of reactive astrocytes engaged in neurotoxic signaling. However, because Pirb+ astrocytes also express myeloid markers, careful doublet depletion and in situ validation are required to confirm their identity.

### Limitations

Limitations include the reliance on public datasets with heterogeneous protocols, the small microglial sample size in GSE233812, potential batch effects in cross-dataset integration, and the lack of single-cell resolution in Visium data. GSE227651 and additional spatial datasets remain to be incorporated.

### Future directions

The in vitro functional roadmap targets Pirb in primary/BV2 microglia under OGD/LPS stimulation using neutralizing antibodies, siRNA, and Pirb-Fc fusion proteins. Readouts span activation markers, cytokine secretion, NF-κB/STAT signaling, phagocytosis, synapse pruning, and neuronal survival. Validating the causal role of Pirb in microglial activation would support its development as a stroke immunomodulatory target.

---

## Methods

### Data collection and preprocessing

Public datasets were downloaded from GEO: GSE174574, GSE171169, GSE225948, GSE233815, GSE233812, GSE233813, and GSE233814. scRNA-seq/snRNA-seq data were processed with Scanpy (v1.12.1): QC filtering, normalization, log-transformation, highly variable gene selection, scaling, PCA, neighbor graph, UMAP, and Leiden clustering. Cell types were annotated using canonical marker genes.

### Differential expression and pathway analysis

Pirb-positive cells were defined as cells with normalized Pirb expression > 0. Differential expression was performed using Wilcoxon rank-sum tests. Pathway enrichment was performed using gseapy (GSEA/ORA) and custom gene sets for lysosome, complement, and microglia-induced A1 astrocyte programs.

### Cross-dataset microglial integration

Microglia from GSE174574, GSE233812, and GSE233813 were normalized per dataset, concatenated on common genes, and batch-corrected with ComBat (dataset as batch). UMAP, Leiden clustering, and diffusion pseudotime were computed.

### Visium spatial analysis

Spatial coordinates were reconstructed from `json.gz` files by mapping 10x Visium v1 barcodes to fiducial/oligo positions and transform matrices. Spot-level Pirb positivity and distance to tissue boundary were computed. Microglial and inflammatory module scores were calculated using Scanpy `score_genes`.

### Supplementary information

Supplementary Tables 1–7 provide dataset summary, cell-type-specific and cross-time differential expression results, pathway enrichments, and spatial statistics.

---

## Data availability

All data are publicly available via GEO under accession numbers GSE174574, GSE171169, GSE225948, GSE233815, GSE233812, GSE233813, and GSE233814.

## Code availability

Analysis code is available in `D:/Pirb_stroke_project/03_analysis_code/` and is provided as Supplementary Software.

## Acknowledgements

[To be added.]

## Author contributions

[To be added.]

## Competing interests

The authors declare no competing interests.

## Additional information

Supplementary information is available for this paper at [repository path]. Correspondence and requests for materials should be addressed to [corresponding author].

## References

1. [References to be added.]
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Build DOCX
from docx.enum.style import WD_STYLE_TYPE
doc = Document()

# Title
title = doc.add_heading('Pirb-positive cells orchestrate post-ischemic neuroinflammation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('A multi-dataset single-cell and spatial transcriptomics study')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True
subtitle.runs[0].font.size = Pt(14)

# Authors
auth = doc.add_paragraph('[Author names]¹,²,*')
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff = doc.add_paragraph('¹ Department, Institution; ² Department, Institution')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
cor = doc.add_paragraph('*Correspondence: email@institution.edu')
cor.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Abstract
doc.add_heading('Abstract', level=1)
abs_p = doc.add_paragraph()
abs_p.add_run('Background: ').bold = True
abs_p.add_run('Paired immunoglobulin-like receptor B (Pirb) is an immune inhibitory receptor implicated in neuroinflammation and axon regeneration failure after cerebral ischemia. ')
abs_p.add_run('Results: ').bold = True
abs_p.add_run('Integrating seven datasets, we show Pirb peaks in microglia at post-ischemic day 3 (D3), is enriched in A1-like reactive astrocytes, and is highly expressed in acute peripheral monocytes/neutrophils. Visium spatial transcriptomics localizes Pirb+ foci to the ischemic penumbra. ')
abs_p.add_run('Conclusions: ').bold = True
abs_p.add_run('Pirb marks a transient, spatially restricted inflammatory population and represents a candidate therapeutic target.')

# Keywords
kw = doc.add_paragraph()
kw.add_run('Keywords: ').bold = True
kw.add_run('Pirb, LILRB2, stroke, microglia, astrocyte, neuroinflammation, single-cell RNA sequencing, spatial transcriptomics')

# Highlights
doc.add_heading('Highlights', level=1)
for h in [
    'Pirb expression peaks at D3 in brain microglia across three independent datasets.',
    'Pirb-positive reactive astrocytes adopt an A1-like inflammatory state after MCAO.',
    'Peripheral blood monocytes and neutrophils show acute Pirb upregulation after stroke.',
    'Visium spatial transcriptomics places Pirb+ inflammatory foci at the ischemic penumbra boundary.',
    'A functional validation roadmap targeting Pirb in microglia is proposed.'
]:
    doc.add_paragraph(h, style='List Bullet')

# Sections
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'Cerebral ischemia triggers a rapid neuroinflammatory response in which resident microglia and infiltrating peripheral myeloid cells release cytokines, '
    'phagocytose synapses, and shape tissue repair. Pirb (paired immunoglobulin-like receptor B) and its human orthologs LILRB1/2 are immune inhibitory receptors '
    'expressed by myeloid cells. In the CNS, Pirb/LILRB signaling has been linked to myelin-mediated axon growth inhibition, microglial activation, and synapse elimination. '
    'Here, we integrate seven independent mouse stroke datasets to map the cell-type-specific expression of Pirb, define its temporal trajectory, localize Pirb-expressing cells in situ, and propose an in vitro functional validation framework.'
)

# Results
doc.add_heading('Results', level=1)
doc.add_heading('Pirb is induced in A1-like reactive astrocytes after MCAO', level=2)
doc.add_paragraph(
    'Using GSE174574 (Sham vs MCAO 24 h), we identified 5,813 astrocytes among 56,486 high-quality cells. '
    'Pirb-positive astrocytes were significantly enriched after MCAO (3.60% vs 0.50% in sham) and concentrated in the A1-like reactive state (7.04% Pirb+ in MCAO A1-like vs 0.39% in sham). '
    'Differentially expressed genes included Cd14, Fcer1g, Msr1, Lilr4b, and C5ar1 (Supplementary Table 2). Pathway analysis implicated lysosome, complement activation, and microglia-induced A1 astrocyte programs (Supplementary Table 6).'
)

doc.add_heading('Temporal dynamics of Pirb across independent datasets', level=2)
doc.add_paragraph(
    'Multiple datasets converged on a D3 peak for Pirb expression in brain microglia: GSE233812 scRNA-seq (26.6% Pirb+ microglia at D3), '
    'GSE233813 snRNA-seq (6.4%), GSE233815 bulk RNA-seq (CPM 6.69 at 3 d), and cross-dataset integration of GSE174574 + GSE233812 + GSE233813 (14.47% Pirb+ microglia at D3). '
    'GSE174574 at 24 h showed 6.2% Pirb+ microglia, consistent with an ascending phase toward the D3 peak.'
)

doc.add_heading('Peripheral myeloid cells are a major Pirb-expressing compartment', level=2)
doc.add_paragraph(
    'GSE225948 (91,688 cells after QC) showed that brain microglia maintained low Pirb expression across sham, D02, and D14 (~2.4–2.9%). '
    'In contrast, peripheral blood monocytes and neutrophils displayed high acute Pirb expression peaking at D02 (monocytes 51.72%; neutrophils 52.64%; Supplementary Table 7). '
    'These findings suggest that circulating Pirb+ myeloid cells may infiltrate the ischemic brain.'
)

doc.add_heading('Spatial localization of Pirb expression to the ischemic boundary', level=2)
doc.add_paragraph(
    'GSE233814 Visium spatial transcriptomics (11,969 spots) showed that Pirb-positive spot fraction peaked at D3 (5.93%). '
    'Pirb+ spots were spatially clustered and significantly closer to the tissue boundary than Pirb− spots (mean distance ratio = 0.60, p = 3.3 × 10⁻¹⁷), '
    'consistent with localization to the ischemic penumbra (Supplementary Table 7).'
)

doc.add_heading('Discussion', level=1)
doc.add_paragraph(
    'Our multi-dataset analysis provides convergent evidence that Pirb is transiently upregulated in microglia and peripheral myeloid cells after cerebral ischemia, '
    'and that Pirb+ inflammatory foci localize to the ischemic boundary. The D3 microglial peak aligns with the delayed neuroinflammatory phase, while the D02 peripheral peak '
    'identifies a window for early immune intervention. The in vitro functional roadmap targets Pirb in primary/BV2 microglia under OGD/LPS stimulation and will test activation markers, '
    'cytokine secretion, NF-κB/STAT signaling, phagocytosis, synapse pruning, and neuronal survival.'
)

doc.add_heading('Methods', level=1)
for m in [
    'Public datasets were downloaded from GEO and processed with Scanpy v1.12.1.',
    'Pirb-positive cells were defined by normalized Pirb expression > 0.',
    'Differential expression used Wilcoxon rank-sum tests; pathway enrichment used gseapy GSEA/ORA.',
    'Cross-dataset microglial integration used ComBat batch correction.',
    'Visium coordinates were reconstructed from json.gz fiducial/oligo/transform data.'
]:
    doc.add_paragraph(m, style='List Bullet')

# Back matter
doc.add_heading('Data availability', level=1)
doc.add_paragraph('All data are publicly available via GEO: GSE174574, GSE171169, GSE225948, GSE233815, GSE233812, GSE233813, GSE233814.')

doc.add_heading('Code availability', level=1)
doc.add_paragraph('Analysis code is available in D:/Pirb_stroke_project/03_analysis_code/.')

doc.add_heading('Competing interests', level=1)
doc.add_paragraph('The authors declare no competing interests.')

doc.add_heading('Supplementary information', level=1)
doc.add_paragraph('Supplementary Tables 1–7 are provided in Supplementary_Tables_Pirb_Stroke.xlsx.')

# Insert graphical abstract
doc.add_page_break()
doc.add_heading('Graphical Abstract', level=1)
doc.add_picture(GA_PATH, width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Insert supplementary tables note
p = doc.add_paragraph('Fig. S1 | Graphical abstract. Pirb is induced in microglia, A1 astrocytes and infiltrating peripheral myeloid cells after cerebral ischemia, with peak expression at D3 and localization to the ischemic penumbra.')
p.runs[0].italic = True

# Insert key figures
for fname, cap in [
    ('GSE174574/marker_dotplot.png', 'Fig. 1 | GSE174574 astrocyte state markers and Pirb distribution.'),
    ('GSE225948/pirb_Mo_time.png', 'Fig. 2 | GSE225948 peripheral monocyte Pirb dynamics.'),
    ('microglia_cross_time/pirb_fraction_timeline.png', 'Fig. 3 | Cross-dataset microglial Pirb fraction timeline.'),
    ('GSE233814/pirb_spatial_combined_panel.png', 'Fig. 4 | GSE233814 Visium Pirb spatial heatmaps.'),
]:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.runs[0].italic = True

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
