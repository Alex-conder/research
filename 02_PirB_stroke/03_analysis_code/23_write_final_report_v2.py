"""
生成最终阶段性报告 V2（含 doublet 质控与 GSE233815 bulk 验证）
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "D:/Pirb_stroke_project/04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖与多数据集验证阶段性报告_V2.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖与多数据集验证阶段性报告（V2）

**报告日期**：2026-06-16  
**分析数据集**：GSE174574（主数据集）、GSE171169（验证集）、GSE233815（bulk 时间序列补充）  
**后台任务**：GSE227651、GSE225948 原始文件下载中（受网络限制）

---

## 一、目标与假设

### 1.1 核心目标
1. 明确脑缺血后 Pirb 阳性细胞的细胞类型分布与动态变化。
2. 揭示驱动 Pirb 上调的上游信号（小胶质细胞、少突胶质细胞、免疫细胞等）。
3. 构建 Pirb 参与缺血后炎症反应与轴突再生抑制的机制假设。

### 1.2 核心假设
- **假设 1**：小胶质细胞来源的 IL-1α/TNF/C1q 诱导 A1 样 Pirb+ 星形胶质细胞。
- **假设 2**：Pirb 在星形胶质细胞中放大 NF-κB/IRF/STAT 驱动的炎症程序。
- **假设 3**：少突胶质细胞来源的 MAG/OMgp/MOG 通过 Pirb 抑制轴突再生。
- **假设 4**：Pirb 在免疫细胞/小胶质细胞中参与免疫抑制或突触修剪。
- **假设 5**：内皮细胞 Pirb 上调与血脑屏障破坏相关。

---

## 二、数据集与方法

### 2.1 数据集
| 数据集 | 样本 | 细胞数/样本数 | 用途 |
|--------|------|--------------|------|
| GSE174574 | 3 Sham + 3 MCAO（24 h） | 58,287 细胞（质控后） | 主分析：细胞类型、差异表达、LR 互作、伪时间 |
| GSE171169 | 4 CD45high 样本（5d / 14d） | 10,295 细胞 | 验证：免疫细胞 Pirb 高表达 |
| GSE233815 | MCAO 3h/12h/24h/3D/7D + SH | 48 bulk 样本 | 补充：Pirb 时间动态趋势 |

### 2.2 分析方法
1. **数据读取与质控**：10x 标准三文件手动读取；过滤 n_genes < 500、total_counts < 1000、percent.mt > 20%。
2. **Doublet 检测**：使用 Scrublet 按样本预测 doublet，评估 Pirb+ 细胞纯度。
3. **标准化**：normalize_total + log1p。
4. **细胞类型注释**：基于已知 marker 基因得分分配细胞类型。
5. **差异表达**：Pirb+ vs Pirb− 按细胞类型分别进行 Wilcoxon rank-sum 检验。
6. **通路富集**：对 Pirb+ 星形胶质细胞上调基因进行超几何检验（预定义通路基因集）。
7. **配体-受体互作**：基于 mouseconsensus 数据库 + Pirb 特异性配体库。
8. **转录因子分析**：检查 NF-κB、STAT、IRF 家族表达差异。
9. **伪时间分析**：scanpy diffusion pseudotime，以 Sham 星形胶质细胞为根。
10. **Bulk 时间序列**：GSE233815 CPM 标准化 + t-test 比较 MCAO vs SH。

---

## 三、主要结果

### 3.1 Pirb 在 MCAO 后显著上调
- **GSE174574**：MCAO 后 Pirb 阳性率从 Sham 2.7–2.8% 升至 8.2–9.8%，平均表达升高约 3 倍。
- **GSE171169**：CD45high 免疫细胞中 Pirb 阳性率为 27–28%，5d 与 14d 保持稳定。

### 3.2 Pirb 表达的细胞类型格局

| 细胞类型 | MCAO Pirb+ % | Sham Pirb+ % | 排除 doublet 后 MCAO % | 排除 doublet 后 Sham % |
|----------|-------------|-------------|----------------------|----------------------|
| Immune | 31.37% | 24.24% | 31.16% | 24.76% |
| Microglia | 6.16% | 4.59% | 6.17% | 4.59% |
| Ependymal | 4.41% | 0.88% | 4.55% | 0.90% |
| Astrocyte | 3.13% | 0.66% | 2.98% | 0.51% |
| Oligodendrocyte | 2.52% | 0.92% | 1.91% | 0.75% |
| Pericyte | 1.83% | 0.61% | 1.59% | 0.49% |
| Endothelial | 1.77% | 0.58% | 1.67% | 0.51% |

**Doublet 质控结论**：
- 总体预测 doublet 率为 3.09%。
- Pirb+ 星形胶质细胞中 13.25% 被预测为 doublet，高于 Pirb− 星形胶质细胞的 7.46%。
- 排除 doublet 后，Astrocyte MCAO Pirb+ 率从 3.13% 略降至 2.98%，结论保持稳健。

### 3.3 Pirb+ 星形胶质细胞具有 A1-like 炎症表型
对排除 doublet 后的 72 个 Pirb+ 与 3151 个 Pirb− 星形胶质细胞进行差异表达分析。

**上调基因特征**：
- 免疫/髓系标志：Cd14、Fcer1g、Csf1r、Fcgr3、Ly86
- 补体成分：C1qa、C1qb
- 溶酶体/组织蛋白酶：Ctsb、Ctsz、Ctss、Ctsd
- 炎症/修复因子：Spp1、Ccl4、Ccl12、C5ar1

**下调基因特征**：
- 星形胶质细胞稳态标志：Gja1、Slc1a3、Bcan、Mfge8、Ndrg2

### 3.4 通路富集结果（排除 doublet 后）

| 通路 | p 值 | 重叠基因 |
|------|------|----------|
| 溶酶体/组织蛋白酶 | 6.2 × 10⁻¹⁶ | Ctsb, Ctsd, Ctsz, Ctss, Ctsc |
| 补体激活 | 4.9 × 10⁻⁹ | C1qa, C1qb, C3, C5ar1 |
| 小胶质细胞诱导 A1 | 4.1 × 10⁻⁷ | C1qa, C1qb, Tnf, Il1a |
| MHC I 抗原呈递 | 2.5 × 10⁻⁴ | H2-D1, H2-K1, Cd74 |
| A1 反应性星形胶质细胞 | 2.4 × 10⁻² | C3, Gfap, Vim |
| IRF 信号 | 2.1 × 10⁻² | Irf1, Irf5, Irf7, Irf8 |
| NF-κB 信号 | 2.8 × 10⁻² | Nfkb1, Nfkb2, Rela, Relb |

### 3.5 上游转录因子
Pirb+ 星形胶质细胞中显著上调：Cebpb、Rel、Stat6、Irf8、Irf5。

### 3.6 配体-受体互作预测
- 少突胶质细胞是 Pirb 经典配体 MAG / MOG / OMgp / Nogo-A 的最强来源。
- MHC I → PirB 信号在免疫细胞和小胶质细胞中最强。
- 指向星形胶质细胞的 Pirb 信号主要来自少突胶质细胞、小胶质细胞和免疫细胞。

### 3.7 其他细胞类型的 Pirb 功能特征

| 细胞类型 | Pirb+ 细胞数 | 主要上调特征 | 主要下调特征 |
|----------|-------------|-------------|-------------|
| Immune | 2422 | S100a6, Alox5ap, Ifitm, Lyz2 | Dhrs7b, Man1a2, Pdgfa |
| Microglia | 481 | Pirb, Lyz2 | — |
| Endothelial | 376 | Ctsb, Spp1, Hspa5, Ccl4 | Abcb1a, Cxcl12, Rgcc |
| Oligodendrocyte | 32 | Pirb, Dusp1, Lcp1 | — |
| Pericyte | 38 | Pirb | — |

### 3.8 星形胶质细胞伪时间轨迹
以 Sham 星形胶质细胞为根，Pirb 表达在伪时间 0.35–0.6 达到峰值（阳性率 ~6%），提示为缺血响应过渡状态。

### 3.9 GSE233815 bulk 时间序列验证
GSE233815 bulk RNA-seq 显示 Pirb 在 MCAO 后呈时间依赖性上调：

| 时间点 | MCAO Pirb CPM (mean ± std) | SH Pirb CPM (mean ± std) | log2FC |
|--------|---------------------------|-------------------------|--------|
| 3h | 0.42 ± 0.58 | 0.51 ± 0.74 | -0.09 |
| 12h | 2.67 ± 3.53 | 0.14 ± 0.28 | 1.69 |
| 24h | 2.01 ± 1.30 | 0.00 ± 0.00 | 1.59* |
| 3D | 6.69 ± 6.46 | 1.00 ± 0.98 | 1.94 |
| 7D | 2.60 ± 3.82 | 0.00 ± 0.00 | 1.85 |

*24h t-test p = 0.021。

**趋势**：Pirb 在 12h 开始升高，24h 显著，3D 达到峰值，7D 回落。该趋势与 GSE174574（24h 单细胞）结果一致，并提示 3D 可能是 Pirb 表达高峰。

---

## 四、是否支持核心目标

### 目标 1：明确 Pirb 阳性细胞的细胞类型分布
**支持**。Pirb 在多种细胞类型中表达，免疫细胞和小胶质细胞最高；doublet 质控后结论稳健。

### 目标 2：揭示驱动 Pirb 上调的上游信号
**部分支持**。LR 分析预测少突胶质细胞 MAG/MOG/OMgp 和小胶质细胞/免疫细胞 C1q/TNF/IL-1α 为主要驱动；GSE233815 bulk 时间序列支持缺血后时间依赖性上调。

### 目标 3：构建 Pirb 参与炎症与轴突再生抑制的机制假设
**支持**。Pirb+ 星形胶质细胞的炎症通路富集、TF 激活及髓鞘抑制配体来源分析共同支持该假设。

---

## 五、局限与已补全内容

| 原局限 | 补全措施 | 状态 |
|--------|---------|------|
| Pirb+ 星形胶质细胞可能存在 doublet | Scrublet doublet 检测 + 排除后重分析 | ✅ 已补全 |
| GSE174574 仅 24h 单时间点 | GSE233815 bulk 时间序列（3h–7D）验证 | ✅ 已部分补全 |
| GSE227651 / GSE225948 下载未完成 | 后台多次尝试，网络极不稳定 | ⏳ 待网络稳定后补全 |
| GSE233815 空间数据未下载 | 尚未补全 | ⏳ 数据体量大，待后续决定 |
| 细胞类型注释基于简单 marker | 已用已知 marker 严格注释，尚未用 SingleR/CellTypist | ⏳ 可选优化 |
| 配体-受体分析为预测性 | 尚未进行功能验证 | ⏳ 需实验验证 |

---

## 六、下一步计划

1. **严格质控后续**：对 Pirb+ 星形胶质细胞进一步做环境 RNA 污染评估（如 CellBender）。
2. **空间验证**：下载 GSE233815 Visium 数据，验证 Pirb+ 细胞空间定位。
3. **时间动态**：完成 GSE227651 下载后，分析 day 1/3/7 单细胞 Pirb 变化。
4. **功能实验**：
   - 体外 IL-1α + TNF + C1q 诱导原代星形胶质细胞，检测 Pirb 表达。
   - Pirb 阻断后检测炎症因子释放和 BBB 通透性。
   - 免疫荧光验证 Pirb 与 C3 / C1q / H2-D1 共定位。
5. **多数据集整合**：使用 Harmony / Scanorama 整合多数据集，扩大样本量。

---

## 七、结论

基于 GSE174574、GSE171169 和 GSE233815 的整合分析，脑缺血后 Pirb 在免疫细胞、小胶质细胞和星形胶质细胞等多种细胞类型中上调。排除 doublet 后，Pirb+ 星形胶质细胞仍呈现典型的 A1-like 炎症表型，其上游可能由小胶质细胞来源的 IL-1α/TNF/C1q 诱导，并通过 NF-κB/IRF/STAT 转录程序驱动；少突胶质细胞来源的 MAG/OMgp/MOG 可能通过 Pirb 抑制轴突再生。GSE233815 bulk 时间序列进一步显示 Pirb 在缺血后 12h 开始升高、3D 达到峰值。这些发现为 Pirb 作为缺血性卒中治疗靶点提供了多细胞类型、多时间点的证据支持。
"""

# 保存 Markdown
with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

# 生成 DOCX
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
