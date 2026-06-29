"""
生成 Nature Communications 格式完整报告的中文版（DOCX + MD）。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞调控神经炎症：多数据集单细胞与空间转录组研究

**作者**：[作者姓名]¹,²,*  
**单位**：¹ 院系，机构，城市，国家；² 院系，机构，城市，国家  
**通讯作者**：*email@institution.edu  
**日期**：{datetime.now().strftime('%Y-%m-%d')}

---

## 摘要

**背景**：配对免疫球蛋白样受体 B（Pirb，人源同源基因为 LILRB2）是一种免疫抑制性受体，参与神经炎症和轴突再生障碍，但脑缺血后其细胞类型特异性动态仍不清楚。  
**结果**：本研究整合 7 个公共数据集（GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814），发现实验性卒中后 Pirb 在小胶质细胞和浸润性外周髓系细胞中显著上调，并于缺血后第 3 天（D3）达到峰值。在 GSE174574 中，Pirb 阳性星形胶质细胞富集于 A1-like 反应性状态（MCAO 7.04% vs Sham 0.39%）。跨数据集小胶质细胞整合进一步确认 D3 峰值（14.47%）。GSE225948 显示急性期 D02 外周血单核细胞（51.72%）和中性粒细胞（52.64%）Pirb 高表达，而脑内小胶质细胞维持低水平（~2.4–2.9%）。Visium 空间转录组将 D3 Pirb 表达 spots 定位至缺血边界/半暗带（Pirb+ spots 距组织边缘距离为 Pirb− 的 0.60 倍，p = 3.3 × 10⁻¹⁷）。  
**结论**：Pirb 标记卒中后一过性、空间受限的炎症细胞群体。本研究将 Pirb 定位为调控缺血后神经炎症的潜在治疗靶点，并提供了体外功能验证路线图。

**关键词**：Pirb，LILRB2，脑卒中，小胶质细胞，星形胶质细胞，神经炎症，单细胞 RNA 测序，空间转录组

---

## 亮点

- 三个独立数据集一致显示脑缺血后第 3 天小胶质细胞 Pirb 表达达峰。
- MCAO 后 Pirb 阳性反应性星形胶质细胞呈现 A1-like 炎症状态。
- 外周血单核细胞和中性粒细胞在卒中急性期 Pirb 上调，可能浸润受损脑组织。
- Visium 空间转录组将 Pirb 阳性炎症灶定位于缺血半暗带边界。
- 提出了针对小胶质细胞 Pirb 的功能验证路线图。

---

## 引言

脑缺血触发快速的神经炎症反应，其中驻留小胶质细胞和浸润性外周髓系细胞释放细胞因子、吞噬突触并塑造组织修复。配对免疫球蛋白样受体 B（Pirb）及其人源同源基因 LILRB1/2 是表达于髓系细胞和特定神经元亚群的免疫抑制性受体。在中枢神经系统中，PirB/LILRB 信号与髓鞘介导的轴突生长抑制、小胶质细胞激活、突触消除和神经可塑性相关（Atwal et al., 2008; Kim et al., 2013; Adelson et al., 2012）。

尽管已有证据表明 PirB 在缺血性脑损伤后上调（Gou et al., 2013; Gou et al., 2018），但其细胞类型特异性表达轨迹、空间分布及其与反应性胶质增生的机制关系尚未在单细胞及空间分辨率下系统解析。本研究整合 7 个独立小鼠卒中数据集，旨在（i）绘制 Pirb 的细胞类型特异性表达图谱，（ii）定义其时间动态，（iii）利用 Visium 空间转录组进行原位定位，（iv）提出体外功能验证框架。

---

## 结果

### MCAO 后 Pirb 阳性星形胶质细胞富集于 A1-like 状态

利用 GSE174574（Sham vs MCAO 24 h），我们在 56,486 个高质量细胞中鉴定出 5,813 个星形胶质细胞。Pirb 阳性星形胶质细胞虽比例较低，但在 MCAO 后显著增加（3.60% vs Sham 0.50%）。状态评分显示 Pirb+ 星形胶质细胞集中于 A1-like 反应性状态（MCAO A1-like 中 7.04% Pirb+ vs Sham A1-like 0.39%；补充表 2）。Pirb+ 星形胶质细胞的差异基因包括 Cd14、Fcer1g、Msr1、Lilr4b、C5ar1 等免疫/髓系相关基因（补充表 2），通路分析提示溶酶体、补体激活及小胶质细胞诱导 A1 星形胶质细胞程序显著上调（补充表 6）。

### Pirb 的时间动态跨数据集验证

多个独立数据集一致显示脑内小胶质细胞 Pirb 在 D3 达峰：GSE233812 scRNA-seq（D3 小胶质细胞 Pirb+ 26.6%）、GSE233813 snRNA-seq（6.4%）、GSE233815 bulk RNA-seq（3 d 时 CPM 6.69），以及 GSE174574 + GSE233812 + GSE233813 跨数据集整合（14.47%）。GSE174574 在 24 h 显示 6.2% Pirb+ 小胶质细胞，与上升至 D3 峰值的动态一致。

### 外周髓系细胞是 Pirb 表达的重要来源

对完整 26 样本 GSE225948 数据集（质控后 91,688 个细胞）的分析显示，脑内小胶质细胞在 Sham、D02、D14 各时间点 Pirb 表达均维持低水平（~2.4–2.9%）。相反，外周血单核细胞和中性粒细胞在急性期 D02 呈现高 Pirb 表达（单核细胞 51.72%；中性粒细胞 52.64%），并在 D14 回落（补充表 7–8）。D02 外周血 Mo/Neu 中 Pirb+ vs Pirb− 的差异基因包括 Pirb、Actb、Alox5ap、Tyrobp、S100a11 等（补充表 8）。这些结果提示循环 Pirb+ 髓系细胞可能浸润缺血脑组织，构成脑内 Pirb+ 炎症细胞池的一部分。

### Pirb 表达的空间定位：缺血边界

GSE233814 Visium 空间转录组（11,969 spots）显示 Pirb 阳性 spot 比例在 D3 达峰（5.93%），高于对照（0.62%）、D1（2.40%）和 D7（~2%）。Pirb+ spots 呈空间聚集（D3 最近邻距离均值 236 pixels），且显著靠近组织边界（Pirb+/Pirb− 平均距离比 = 0.60，p = 3.3 × 10⁻¹⁷），与缺血半暗带定位一致。Pirb+ spots 同时表现出更高的小胶质细胞和炎症模块评分（补充表 7）。

### Pirb 介导缺血后神经炎症的整合模型

我们提出以下整合模型：（1）缺血急性期动员 Pirb+ 外周单核细胞/中性粒细胞；（2）驻留小胶质细胞上调 Pirb，于 D3 达峰；（3）活化小胶质细胞分泌 IL-1α/TNF/C1q，驱动星形胶质细胞向 A1-like Pirb+ 状态转化；（4）Pirb+ 髓系/星形胶质细胞通过 NF-κB/STAT/IRF 程序放大神经炎症；（5）少突胶质细胞来源的 Pirb 配体（MAG/OMgp）进一步抑制轴突再生。

---

## 讨论

本多数据集分析提供了趋同证据：Pirb 在脑缺血后一过性上调于小胶质细胞和外周髓系细胞，且 Pirb+ 炎症灶定位于缺血边界。小胶质细胞 D3 峰值与已确立的延迟性神经炎症期一致，而外周 D02 峰值提示存在早期免疫干预窗口。

GSE174574 中低丰度但高度炎症性的 Pirb+ A1-like 星形胶质细胞群体提示，Pirb 可能标记参与神经毒性信号的反应性星形胶质细胞亚群。然而，由于 Pirb+ 星形胶质细胞亦表达髓系标志物，需通过严格去 doublet 和原位验证确认其身份。

### 局限性与未来方向

局限性包括依赖异质性公共数据集、GSE233812 小胶质细胞样本量较小、跨数据集整合存在残余批次效应、Visium 缺乏单细胞分辨率。GSE227651 及其他空间数据集仍有待纳入。

体外功能路线图以 OGD/LPS 刺激下的原代/BV2 小胶质细胞为模型，采用中和抗体、siRNA 及 Pirb-Fc 融合蛋白干预，检测激活标志物、细胞因子分泌、NF-κB/STAT 信号、吞噬、突触修剪和神经元存活。验证 Pirb 在小胶质细胞激活中的因果作用将支持其作为卒中免疫调节靶点的开发。

---

## 方法

### 数据收集与预处理

公共数据集下载自 GEO：GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814。scRNA-seq/snRNA-seq 数据使用 Scanpy（v1.12.1）处理：QC 过滤、归一化、log1p 转换、高变基因选择、缩放、PCA、邻居图、UMAP 和 Leiden 聚类。细胞类型基于经典标志基因注释。

### 差异表达与通路分析

Pirb 阳性细胞定义为归一化 Pirb 表达 > 0 的细胞。差异表达采用 Wilcoxon 秩和检验。通路富集使用 gseapy（GSEA/ORA）及手工整理的溶酶体、补体激活、小胶质细胞诱导 A1 星形胶质细胞基因集。

### 跨数据集小胶质细胞整合

GSE174574、GSE233812 和 GSE233813 的小胶质细胞按数据集分别归一化，基于共同基因拼接，并以 ComBat 按数据集批次校正。随后进行 UMAP、Leiden 聚类和扩散拟时分析。

### Visium 空间分析

从 `json.gz` 文件中解析 fiducial/oligo/transform 信息，通过 10x Visium v1 barcode 映射重建各 spot 像素坐标。计算 spot 水平 Pirb 阳性率和到组织边界的距离。使用 Scanpy `score_genes` 计算小胶质细胞和炎症模块评分。

### 补充信息

补充表 1–8 提供数据集汇总、细胞类型及跨时间差异表达结果、通路富集、空间统计和外周髓系差异基因。

---

## 数据可用性

所有数据均可通过 GEO 公开获取：GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814。

## 代码可用性

分析代码存放于 `D:/Pirb_stroke_project/03_analysis_code/`，关键脚本包括：
- `D:/Pirb_stroke_project/03_analysis_code/56_final_assembly_all.py`（一键重新生成最终交付物）
- `D:/Pirb_stroke_project/03_analysis_code/55_gse225948_pb_de.py`（外周血 Mo/Neu 差异表达）
- `D:/Pirb_stroke_project/03_analysis_code/54_nc_report_supp_graphical_abstract.py`（初始 NC 报告、补充表、Graphical Abstract）
- `D:/Pirb_stroke_project/03_analysis_code/57_final_report_original_format.py`（原格式中文最终报告）
- `D:/Pirb_stroke_project/03_analysis_code/58_generate_improvements_pdf.py`（提升点对比 PDF）
- `D:/Pirb_stroke_project/03_analysis_code/59_package_deliverables.py`（打包脚本）
- `D:/Pirb_stroke_project/03_analysis_code/60_nc_report_chinese.py`（本中文版报告生成脚本）

## 致谢

[待补充]

## 作者贡献

[待补充]

## 利益冲突

作者声明无利益冲突。

## 补充信息

补充表 1–8 存放于：
- `D:/Pirb_stroke_project/04_reports/supplementary/Supplementary_Tables_Pirb_Stroke.xlsx`

Graphical Abstract 存放于：
- `D:/Pirb_stroke_project/04_reports/figures/graphical_abstract_draft.png`
- `D:/Pirb_stroke_project/04_reports/figures/graphical_abstract_draft.svg`
- `D:/Pirb_stroke_project/04_reports/figures/graphical_abstract_draft.pdf`

## 参考文献

1. Atwal, J.K., et al. PirB is a functional receptor for myelin inhibitors of axonal regeneration. *Science* 322, 967–970 (2008).
2. Syken, J., et al. PirB restricts ocular-dominance plasticity in visual cortex. *Science* 313, 1795–1800 (2006).
3. Gou, X., et al. Spatio-temporal expression of paired immunoglobulin-like receptor-B after focal cerebral ischaemia. *Brain Inj* 27, 1311–1315 (2013).
4. Gou, X., et al. Neuronal PirB upregulated in cerebral ischemia as a theranostic target. *J Am Heart Assoc* 7, e007197 (2018).
5. Adelson, J.D., et al. Neuroprotection from stroke in the absence of MHCI or PirB. *Neuron* 73, 1100–1107 (2012).
6. Liddelow, S.A., et al. Neurotoxic reactive astrocytes are induced by activated microglia. *Nature* 541, 481–487 (2017).
7. Zamanian, J.L., et al. Genomic analysis of reactive astrogliosis. *J Neurosci* 32, 6391–6410 (2012).
8. Keren-Shaul, H., et al. A unique microglia type associated with restricting development of Alzheimer's disease. *Cell* 169, 1276–1290 (2017).
9. Hammond, T.R., et al. Single-cell RNA sequencing of microglia throughout the mouse lifespan and in the injured brain. *Immunity* 50, 253–271 (2019).
10. Zheng, K., et al. Single-cell RNA-seq reveals the transcriptional landscape in ischemic stroke. *J Cereb Blood Flow Metab* 42, 56–73 (2022).
11. Kim, T., et al. Human LilrB2 is a β-amyloid receptor and its murine homolog PirB regulates synaptic plasticity. *Science* 341, 1399–1404 (2013).
12. Wolf, F.A., et al. SCANPY: large-scale single-cell gene expression data analysis. *Genome Biol* 19, 15 (2018).
13. Hafemeister, C. & Satija, R. Normalization and variance stabilization of single-cell RNA-seq data using regularized negative binomial regression. *Genome Biol* 20, 296 (2019).
14. 10x Genomics. Visium Spatial Gene Expression. https://www.10xgenomics.com/products/spatial-gene-expression (2026).

---

## 图注

**图 1 | GSE174574 中 Pirb 的细胞类型特异性表达。**（a）56,486 个细胞按主要细胞类型着色的 UMAP。（b）各细胞类型中 Pirb 表达情况。（c）按条件和细胞类型划分的 Pirb+ 细胞比例。（d）星形胶质细胞状态 UMAP，显示 Homeostatic、PanReactive、A1-like、A2-like 和 Proliferative 状态。（e）Sham 与 MCAO 各星形胶质细胞状态中 Pirb+ 比例。

**图 2 | 跨数据集的 Pirb 时间动态。**（a）GSE233815 bulk RNA-seq 中 3 h、12 h、24 h、3 d、7 d 的 Pirb CPM。（b）GSE233812 scRNA-seq 中小胶质细胞 Pirb+ 比例随 sham、D1、D3、D7 变化。（c）GSE233813 snRNA-seq 中小胶质细胞 Pirb+ 比例。（d）跨数据集小胶质细胞整合中不同数据集和时间点的 Pirb+ 比例。

**图 3 | GSE225948 外周血 vs 脑组织 Pirb 表达。**（a）91,688 个 PB 和脑组织细胞 UMAP，按组织着色。（b）Sham、D02、D14 脑内小胶质细胞 Pirb+ 比例。（c）Sham、D02、D14 PB 单核细胞和中性粒细胞 Pirb+ 比例。（d）D02 PB Mo/Neu 中 Pirb+ vs Pirb− 的 top 差异基因。

**图 4 | GSE233814 Visium 空间定位 Pirb。**（a）对照、D1、D3、D7、D7_rep 的 H&E 图像叠加 Pirb 表达。（b）各时间点 Pirb+ spot 比例。（c）Pirb+ 与 Pirb− spots 到组织边界的距离。（d）Pirb+ spots 的空间聚集及其在缺血边界的富集。

**图 5 | Pirb 介导缺血后神经炎症的模型。**健康脑、D3 缺血半暗带（Pirb+ 小胶质细胞、A1 星形胶质细胞、浸润 Mo/Neu）、下游分子机制（IL-1β/TNF/C1q、NF-κB/STAT/IRF）以及 Pirb 表达时间轴（小胶质细胞 D3 峰值；PB Mo/Neu D02 峰值）示意图。

**Graphical Abstract |** 脑缺血后 Pirb 在小胶质细胞、A1 星形胶质细胞和浸润性外周髓系细胞中诱导表达，于 D3 在脑内小胶质细胞达峰，定位于缺血半暗带，并放大神经炎症、抑制轴突再生。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Build Chinese DOCX
doc = Document()

# Title
title = doc.add_heading('脑缺血后 Pirb 阳性细胞调控神经炎症', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('多数据集单细胞与空间转录组研究')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].italic = True
subtitle.runs[0].font.size = Pt(14)

# Authors
auth = doc.add_paragraph('[作者姓名]¹,²,*')
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
aff = doc.add_paragraph('¹ 院系，机构；² 院系，机构')
aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
cor = doc.add_paragraph('*通讯作者：email@institution.edu')
cor.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Abstract
doc.add_heading('摘要', level=1)
abs_p = doc.add_paragraph()
abs_p.add_run('背景：').bold = True
abs_p.add_run('配对免疫球蛋白样受体 B（Pirb，人源同源基因为 LILRB2）是一种免疫抑制性受体，参与神经炎症和轴突再生障碍，但脑缺血后其细胞类型特异性动态仍不清楚。')
abs_p.add_run('结果：').bold = True
abs_p.add_run('本研究整合 7 个公共数据集，发现实验性卒中后 Pirb 在小胶质细胞和浸润性外周髓系细胞中显著上调，并于缺血后第 3 天（D3）达到峰值。在 GSE174574 中，Pirb 阳性星形胶质细胞富集于 A1-like 反应性状态（7.04% vs 0.39%）。跨数据集小胶质细胞整合确认 D3 峰值（14.47%）。GSE225948 显示急性期 D02 外周血单核细胞（51.72%）和中性粒细胞（52.64%）Pirb 高表达。Visium 空间转录组将 D3 Pirb 表达 spots 定位至缺血边界/半暗带。')
abs_p.add_run('结论：').bold = True
abs_p.add_run('Pirb 标记卒中后一过性、空间受限的炎症细胞群体，是调控缺血后神经炎症的潜在治疗靶点。')

kw = doc.add_paragraph()
kw.add_run('关键词：').bold = True
kw.add_run('Pirb，LILRB2，脑卒中，小胶质细胞，星形胶质细胞，神经炎症，单细胞 RNA 测序，空间转录组')

# Highlights
doc.add_heading('亮点', level=1)
for h in [
    '三个独立数据集一致显示脑缺血后第 3 天小胶质细胞 Pirb 表达达峰。',
    'MCAO 后 Pirb 阳性反应性星形胶质细胞呈现 A1-like 炎症状态。',
    '外周血单核细胞和中性粒细胞在卒中急性期 Pirb 上调，可能浸润受损脑组织。',
    'Visium 空间转录组将 Pirb 阳性炎症灶定位于缺血半暗带边界。',
    '提出了针对小胶质细胞 Pirb 的功能验证路线图。'
]:
    doc.add_paragraph(h, style='List Bullet')

# Introduction
doc.add_heading('引言', level=1)
doc.add_paragraph(
    '脑缺血触发快速的神经炎症反应，其中驻留小胶质细胞和浸润性外周髓系细胞释放细胞因子、吞噬突触并塑造组织修复。'
    '配对免疫球蛋白样受体 B（Pirb）及其人源同源基因 LILRB1/2 是表达于髓系细胞和特定神经元亚群的免疫抑制性受体。'
    '在中枢神经系统中，PirB/LILRB 信号与髓鞘介导的轴突生长抑制、小胶质细胞激活、突触消除和神经可塑性相关。'
    '本研究整合 7 个独立小鼠卒中数据集，旨在绘制 Pirb 的细胞类型特异性表达图谱、定义其时间动态、利用 Visium 空间转录组进行原位定位，并提出体外功能验证框架。'
)

# Results
doc.add_heading('结果', level=1)
sections = [
    ('MCAO 后 Pirb 阳性星形胶质细胞富集于 A1-like 状态',
     '利用 GSE174574（Sham vs MCAO 24 h），我们在 56,486 个高质量细胞中鉴定出 5,813 个星形胶质细胞。Pirb 阳性星形胶质细胞虽比例较低，但在 MCAO 后显著增加（3.60% vs Sham 0.50%）。状态评分显示 Pirb+ 星形胶质细胞集中于 A1-like 反应性状态（MCAO A1-like 中 7.04% vs Sham A1-like 0.39%；补充表 2）。差异基因包括 Cd14、Fcer1g、Msr1、Lilr4b、C5ar1 等免疫/髓系相关基因（补充表 2），通路分析提示溶酶体、补体激活及小胶质细胞诱导 A1 程序显著上调（补充表 6）。'),
    ('Pirb 的时间动态跨数据集验证',
     '多个独立数据集一致显示脑内小胶质细胞 Pirb 在 D3 达峰：GSE233812 scRNA-seq（26.6%）、GSE233813 snRNA-seq（6.4%）、GSE233815 bulk RNA-seq（3 d 时 CPM 6.69），以及跨数据集整合（14.47%）。GSE174574 在 24 h 显示 6.2% Pirb+ 小胶质细胞，与上升至 D3 峰值的动态一致。'),
    ('外周髓系细胞是 Pirb 表达的重要来源',
     '对完整 26 样本 GSE225948 数据集（91,688 个细胞）的分析显示，脑内小胶质细胞 Pirb 维持低水平（~2.4–2.9%）。相反，外周血单核细胞和中性粒细胞在急性期 D02 呈现高 Pirb 表达（单核细胞 51.72%；中性粒细胞 52.64%），并在 D14 回落（补充表 7–8）。这些结果提示循环 Pirb+ 髓系细胞可能浸润缺血脑组织。'),
    ('Pirb 表达的空间定位：缺血边界',
     'GSE233814 Visium 空间转录组（11,969 spots）显示 Pirb 阳性 spot 比例在 D3 达峰（5.93%）。Pirb+ spots 呈空间聚集且显著靠近组织边界（Pirb+/Pirb− 平均距离比 = 0.60，p = 3.3 × 10⁻¹⁷），与缺血半暗带定位一致。Pirb+ spots 同时表现出更高的小胶质细胞和炎症模块评分（补充表 7）。'),
]
for sub, txt in sections:
    doc.add_heading(sub, level=2)
    doc.add_paragraph(txt)

# Discussion
doc.add_heading('讨论', level=1)
doc.add_paragraph(
    '本多数据集分析提供了趋同证据：Pirb 在脑缺血后一过性上调于小胶质细胞和外周髓系细胞，且 Pirb+ 炎症灶定位于缺血边界。'
    '小胶质细胞 D3 峰值与已确立的延迟性神经炎症期一致，而外周 D02 峰值提示存在早期免疫干预窗口。'
    'GSE174574 中低丰度但高度炎症性的 Pirb+ A1-like 星形胶质细胞群体提示，Pirb 可能标记参与神经毒性信号的反应性星形胶质细胞亚群。'
    '体外功能路线图以 OGD/LPS 刺激下的原代/BV2 小胶质细胞为模型，采用中和抗体、siRNA 及 Pirb-Fc 融合蛋白干预，验证 Pirb 在小胶质细胞激活中的因果作用。'
)

# Methods
doc.add_heading('方法', level=1)
for m in [
    '公共数据集下载自 GEO 并使用 Scanpy v1.12.1 处理。',
    'Pirb 阳性细胞定义为归一化 Pirb 表达 > 0 的细胞。',
    '差异表达采用 Wilcoxon 秩和检验；通路富集使用 gseapy GSEA/ORA。',
    '跨数据集小胶质细胞整合采用 ComBat 批次校正。',
    'Visium 坐标通过解析 json.gz 中 fiducial/oligo/transform 数据重建。'
]:
    doc.add_paragraph(m, style='List Bullet')

# Back matter
doc.add_heading('数据可用性', level=1)
doc.add_paragraph('所有数据均可通过 GEO 公开获取：GSE174574、GSE171169、GSE225948、GSE233815、GSE233812、GSE233813、GSE233814。')

doc.add_heading('代码可用性', level=1)
doc.add_paragraph('分析代码存放于 D:/Pirb_stroke_project/03_analysis_code/，关键脚本包括 56_final_assembly_all.py、55_gse225948_pb_de.py、54_nc_report_supp_graphical_abstract.py、57_final_report_original_format.py、58_generate_improvements_pdf.py、59_package_deliverables.py 和本中文版生成脚本 60_nc_report_chinese.py。')

doc.add_heading('利益冲突', level=1)
doc.add_paragraph('作者声明无利益冲突。')

doc.add_heading('补充信息', level=1)
doc.add_paragraph('补充表 1–8 存放于 D:/Pirb_stroke_project/04_reports/supplementary/Supplementary_Tables_Pirb_Stroke.xlsx；Graphical Abstract 存放于 D:/Pirb_stroke_project/04_reports/figures/graphical_abstract_draft.png/.svg/.pdf。')

doc.add_heading('参考文献', level=1)
ref_list = [
    'Atwal, J.K., et al. PirB is a functional receptor for myelin inhibitors of axonal regeneration. Science 322, 967–970 (2008).',
    'Syken, J., et al. PirB restricts ocular-dominance plasticity in visual cortex. Science 313, 1795–1800 (2006).',
    'Gou, X., et al. Spatio-temporal expression of paired immunoglobulin-like receptor-B after focal cerebral ischaemia. Brain Inj 27, 1311–1315 (2013).',
    'Gou, X., et al. Neuronal PirB upregulated in cerebral ischemia as a theranostic target. J Am Heart Assoc 7, e007197 (2018).',
    'Adelson, J.D., et al. Neuroprotection from stroke in the absence of MHCI or PirB. Neuron 73, 1100–1107 (2012).',
    'Liddelow, S.A., et al. Neurotoxic reactive astrocytes are induced by activated microglia. Nature 541, 481–487 (2017).',
    'Zamanian, J.L., et al. Genomic analysis of reactive astrogliosis. J Neurosci 32, 6391–6410 (2012).',
    'Keren-Shaul, H., et al. A unique microglia type associated with restricting development of Alzheimer\'s disease. Cell 169, 1276–1290 (2017).',
    'Hammond, T.R., et al. Single-cell RNA sequencing of microglia throughout the mouse lifespan and in the injured brain. Immunity 50, 253–271 (2019).',
    'Zheng, K., et al. Single-cell RNA-seq reveals the transcriptional landscape in ischemic stroke. J Cereb Blood Flow Metab 42, 56–73 (2022).',
    'Kim, T., et al. Human LilrB2 is a β-amyloid receptor and its murine homolog PirB regulates synaptic plasticity. Science 341, 1399–1404 (2013).',
    'Wolf, F.A., et al. SCANPY: large-scale single-cell gene expression data analysis. Genome Biol 19, 15 (2018).',
    'Hafemeister, C. & Satija, R. Normalization and variance stabilization of single-cell RNA-seq data using regularized negative binomial regression. Genome Biol 20, 296 (2019).',
    '10x Genomics. Visium Spatial Gene Expression. https://www.10xgenomics.com/products/spatial-gene-expression (2026).',
]
for r in ref_list:
    doc.add_paragraph(r, style='List Number')

# Graphical abstract page
doc.add_page_break()
doc.add_heading('Graphical Abstract', level=1)
doc.add_picture(os.path.join(FIG_DIR, 'graphical_abstract_draft.png'), width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('图 S1 | Graphical Abstract。脑缺血后 Pirb 在小胶质细胞、A1 星形胶质细胞和浸润性外周髓系细胞中诱导表达，于 D3 在脑内小胶质细胞达峰，定位于缺血半暗带，并放大神经炎症、抑制轴突再生。')
p.runs[0].italic = True

# Figure legends
doc.add_page_break()
doc.add_heading('图注', level=1)
legends = [
    ('图 1 | GSE174574 中 Pirb 的细胞类型特异性表达。', '(a) 56,486 个细胞 UMAP 按主要细胞类型着色。(b) 各细胞类型 Pirb 表达。(c) 按条件和细胞类型划分的 Pirb+ 细胞比例。(d) 星形胶质细胞状态 UMAP。(e) Sham 与 MCAO 各星形胶质细胞状态中 Pirb+ 比例。'),
    ('图 2 | 跨数据集的 Pirb 时间动态。', '(a) GSE233815 bulk RNA-seq Pirb CPM。(b–c) GSE233812 和 GSE233813 中小胶质细胞 Pirb+ 比例。(d) 跨数据集小胶质细胞整合。'),
    ('图 3 | GSE225948 外周血 vs 脑组织 Pirb 表达。', '(a) 91,688 个 PB 和脑组织细胞 UMAP。(b) 脑内小胶质细胞 Pirb+ 比例。(c) PB 单核细胞/中性粒细胞 Pirb+ 比例。(d) D02 PB Mo/Neu top DE 基因。'),
    ('图 4 | GSE233814 Visium 空间定位 Pirb。', '(a) Pirb 空间热图。(b) 各时间点 Pirb+ spot 比例。(c) 到组织边界距离。(d) Pirb+ spots 空间聚集。'),
    ('图 5 | Pirb 介导缺血后神经炎症的模型。', '健康脑、D3 缺血半暗带、分子机制及表达时间轴示意图。'),
]
for title, body in legends:
    p = doc.add_paragraph()
    p.add_run(title + ' ').bold = True
    p.add_run(body)

# Key figures
doc.add_page_break()
doc.add_heading('精选图片', level=1)
for fname, cap in [
    ('GSE174574/marker_dotplot.png', '图 1 相关：GSE174574 星形胶质细胞 marker'),
    ('GSE225948/pirb_Mo_time.png', '图 3 相关：GSE225948 PB 单核细胞 Pirb 动态'),
    ('microglia_cross_time/pirb_fraction_timeline.png', '图 2 相关：跨时间小胶质细胞 Pirb 比例'),
    ('GSE233814/pirb_spatial_combined_panel.png', '图 4 相关：GSE233814 Visium Pirb 空间热图'),
]:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.runs[0].italic = True

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
