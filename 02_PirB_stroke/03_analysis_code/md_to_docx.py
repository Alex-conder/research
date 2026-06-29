#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 markdown 报告转换为 docx（保留标题、段落、表格）
用法：
    python md_to_docx.py <input.md> <output.docx>
"""
import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def md_to_docx(md_path: str, docx_path: str):
    doc = Document()
    # 设置中文字体支持
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(10.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    table_lines = []
    in_table = False

    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            if in_table:
                add_table(doc, table_lines)
                table_lines = []
                in_table = False
            doc.add_paragraph()
            continue

        # 表格
        if line.startswith('|'):
            table_lines.append(line)
            in_table = True
            continue
        elif in_table:
            add_table(doc, table_lines)
            table_lines = []
            in_table = False

        # 标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            # 独立加粗行作为小标题
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
        else:
            # 普通段落，简单处理加粗和斜体
            p = doc.add_paragraph()
            add_formatted_text(p, line)

    if in_table:
        add_table(doc, table_lines)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


def add_formatted_text(paragraph, text):
    # 粗体 **text**
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_table(doc, lines):
    # 解析 markdown 表格
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        # 跳过分隔行
        if re.match(r'^[:\-|\s]+$', cells[0]):
            continue
        rows.append(cells)
    if len(rows) < 1:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell_text = row[j] if j < len(row) else ''
            table.rows[i].cells[j].text = cell_text
    doc.add_paragraph()


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
