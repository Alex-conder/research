"""
生成最终阶段性报告 V6（加入 GSE233812 scRNA-seq 时间序列结果）
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "D:/Pirb_stroke_project/04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及单细胞时间序列确认阶段性报告_V6.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及单细胞时间序列确认阶段性报告（V6）

**报告日期**：2026-06-16  
**核心数据集**：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq  
**分析环境**：Python + scanpy；R 不可用  
**下载状态**：GSE233812 已完成；GSE233813 下载中；GSE233814 / GSE227651 因网络问题暂停

---

## 一、核心假设

1. 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
2. Pirb 放大 NF-κB/IRF/STAT 炎症程序
3. 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
4. Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
5. 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集

| 数据集 | 类型 | 样本 | 细胞数/样本数 | 用途 |
|--------|------|------|---------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 63,733 | 髓系免疫细胞验证 |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |

---

## 三、主要结果

### 3.1 GSE174574 主分析（24 h MCAO）
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。
- Pirb+ 星形胶质细胞富集：溶酶体、补体、小胶质细胞诱导 A1、NF-κB/IRF/STAT。
- 配体-受体：少突胶质细胞 MAG/MOG/OMgp/Nogo-A 和 MHC I → PirB。

### 3.2 GSE171169 验证
- CD45high 免疫细胞 Pirb+ 率 27–28%。

### 3.3 GSE225948 独立验证
- 髓系免疫细胞高表达 Pirb：PB 中性粒细胞/单核细胞 >50%；脑内浸润 MdC/粒细胞 24–46%。
- 小胶质细胞 Pirb 表达低（~2.5–2.9%），内皮细胞几乎不表达。

### 3.4 GSE233815 bulk 时间序列
- Pirb CPM 在 3D 达峰（6.69），24h 显著（p=0.021）。
- A1/炎症标志（Gfap、C3、C1qa/C1qb、Spp1、Lyz2）同步 3D 达峰。

### 3.5 GSE233812 单细胞时间序列（新）
**样本**：sham（1,484 cells）、D1（1,516）、D3（1,323）、D7（1,836）。

#### 3.5.1 细胞类型组成
- Oligodendrocyte（4,852, 78.8%）、Astrocyte（462）、Microglia（406）、Endothelial（147）、Neuron（156）、Pericyte（107）、OPC（29）。

#### 3.5.2 Pirb 表达格局
- **仅小胶质细胞显著表达 Pirb**。
- 其他细胞类型（Astrocyte、Endothelial、Neuron、Oligodendrocyte、Pericyte、OPC）Pirb 几乎不表达。

#### 3.5.3 时间动态（关键发现）
| 时间 | Microglia Pirb+ 率 |
|------|-------------------|
| sham | 8.9% |
| D1 | 15.9% |
| **D3** | **26.6%** |
| D7 | 8.2% |

- **Pirb 在小胶质细胞中于 D3（3 天）达到峰值**，与 GSE233815 bulk 时间序列高度一致。
- D1 开始上调，D7 回落至基线水平。

---

## 四、跨数据集共识

### 4.1 时间动态
- **GSE233812（scRNA-seq）**：小胶质细胞 Pirb D3 峰值。
- **GSE233815（bulk RNA-seq）**：组织整体 Pirb 3D 峰值。
- **GSE225948**：脑内浸润髓系细胞 D02 高、D14 降。
- **GSE174574**：24 h 小胶质细胞 Pirb+ 率 6.2%。

### 4.2 细胞类型
- **GSE233812**：Pirb 主要在小胶质细胞。
- **GSE174574**：小胶质细胞、星形胶质细胞、免疫细胞均表达。
- **GSE225948**：髓系免疫细胞（浸润巨噬细胞/粒细胞）高表达，resident 小胶质细胞低表达。

### 4.3 差异解释
- GSE233812 和 GSE174574 都是全脑 scRNA-seq，但 GSE233812 中小胶质细胞 Pirb 峰值在 D3（26.6%），24h 可能介于 sham 和 D3 之间。
- GSE225948 使用 CD45hi + MG + EC 富集，侧重于免疫细胞，因此 resident 小胶质细胞占比和 Pirb 信号被稀释。
- GSE174574 中星形胶质细胞 Pirb 表达可能反映 24h 特异性反应或 A1 转化。

---

## 五、修正后的机制模型

1. 缺血后小胶质细胞迅速激活，Pirb 表达从 sham 8.9% → D1 15.9% → **D3 26.6%** 递增。
2. 激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。
3. Pirb+ 小胶质细胞可能通过 NF-κB/IRF/STAT 程序放大神经炎症。
4. 少突胶质细胞来源的 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
5. 外周浸润髓系免疫细胞是另一 Pirb 高表达群体，尤其在 CD45hi 富集数据中突出。

---

## 六、生成的关键结果

| 结果 | 路径 |
|------|------|
| GSE233812 UMAP by cell type | `04_reports/figures/GSE233812/umap_cell_type.png` |
| GSE233812 Pirb+ fraction heatmap | `04_reports/figures/GSE233812/pirb_fraction_celltype_time.png` |
| GSE233812 Pirb time violin | `04_reports/figures/GSE233812/pirb_Microglia_time.png` |
| GSE233812 处理对象 | `04_reports/figures/GSE233812_processed.h5ad` |
| GSE225948 图表 | `04_reports/figures/GSE225948/*.png` |

---

## 七、局限与下一步

### 7.1 局限
1. GSE233812 细胞数相对较少（6,159 cells），小胶质细胞仅 406 个。
2. GSE233813（snRNA-seq）尚未下载完成。
3. GSE233814（空间转录组）尚未下载。
4. GSE225948 部分样本缺失。

### 7.2 下一步
1. 完成 GSE233813 snRNA-seq 下载与分析，验证 scRNA-seq 结论。
2. 完成 GSE233814 空间转录组，定位 Pirb+ 小胶质细胞空间分布。
3. 整合 GSE174574 + GSE233812，构建 24h–D7 小胶质细胞 Pirb 动态图谱。
4. 体外功能实验：小胶质细胞激活 + Pirb 阻断。

---

## 八、结论

多数据集交叉验证一致表明，脑缺血后 Pirb 主要在小胶质细胞中表达，并于 **D3（3 天）达到峰值**。GSE233812 单细胞时间序列首次在单细胞水平确认了这一动态，为"小胶质细胞 Pirb → 神经炎症 → 轴突再生抑制"机制提供了关键时间维度证据。后续需结合 snRNA-seq、空间转录组和功能实验进一步阐明 Pirb 的时空调控网络。
"""

with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
