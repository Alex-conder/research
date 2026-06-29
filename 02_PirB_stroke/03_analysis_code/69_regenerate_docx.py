"""
从更新后的 Markdown 报告重新生成 DOCX。
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "../04_reports"
FIG_DIR = "../04_reports/figures"

reports = [
    ("脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.md",
     "脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.docx"),
    ("脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.md",
     "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.docx"),
    ("脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.md",
     "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.docx"),
    ("深度机制分析_5优先级结果汇总.md",
     "深度机制分析_5优先级结果汇总.docx"),
]

def parse_md(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # Replace image references with placeholder
    elements = []
    lines = text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        # Title
        if line.startswith('# ') and not line.startswith('##'):
            elements.append(('title', line.lstrip('#').strip()))
        elif line.startswith('##'):
            level = len(line) - len(line.lstrip('#'))
            elements.append(('heading', level, line.lstrip('#').strip()))
        elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|---'):
            # Table
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            # Skip separator
            if len(rows) > 1:
                table_rows = [r.strip('|').split('|') for r in rows]
                elements.append(('table', table_rows))
            continue
        elif line.startswith('!['):
            m = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if m:
                elements.append(('image', m.group(1), m.group(2)))
        elif re.match(r'^\d+\.', line):
            elements.append(('numbered', re.sub(r'^\d+\.', '', line).strip()))
        elif line.startswith('-') or line.startswith('*'):
            elements.append(('bullet', line[1:].strip()))
        elif line.startswith('```'):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
        else:
            elements.append(('paragraph', line))
        i += 1
    return elements

def add_elements_to_doc(doc, elements, figures_to_insert=None):
    inserted_figs = set()
    for elem in elements:
        if elem[0] == 'title':
            h = doc.add_heading(elem[1], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif elem[0] == 'heading':
            doc.add_heading(elem[2], level=min(elem[1], 3))
        elif elem[0] == 'paragraph':
            doc.add_paragraph(elem[1])
        elif elem[0] == 'bullet':
            doc.add_paragraph(elem[1], style='List Bullet')
        elif elem[0] == 'numbered':
            doc.add_paragraph(elem[1], style='List Number')
        elif elem[0] == 'code':
            p = doc.add_paragraph(elem[1])
            p.style = 'Quote'
        elif elem[0] == 'image':
            path = os.path.join(FIG_DIR, elem[2])
            if os.path.exists(path):
                try:
                    doc.add_picture(path, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph(elem[1]).runs[0].italic = True
                except Exception as e:
                    doc.add_paragraph(f"[Image: {elem[1]} ({e})]")
            else:
                doc.add_paragraph(f"[Image: {elem[1]}]")
        elif elem[0] == 'table':
            if len(elem[1]) < 2:
                continue
            rows = [[c.strip() for c in r] for r in elem[1]]
            n_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=n_cols)
            table.style = 'Light Grid Accent 1'
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    if j < n_cols:
                        table.rows[i].cells[j].text = cell
    
    # Insert additional figures at end if specified and not already inserted
    if figures_to_insert:
        for fname, cap in figures_to_insert:
            fpath = os.path.join(FIG_DIR, fname)
            if os.path.exists(fpath):
                try:
                    doc.add_picture(fpath, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph(cap).runs[0].italic = True
                except Exception as e:
                    print(f"[WARN] figure skip {fpath}: {e}")

# Extra figures for deep mechanism report
extra_figs = [
    ('regulatory_brake/boxplot_GSE233812_D3.png', '图 S1 | GSE233812 D3 Pirb+ vs Pirb- 小胶质细胞应激/percent.mt 比较'),
    ('GSE233814/gradient_de/zone_module_scores_D3.png', '图 S2 | GSE233814 D3 空间梯度模块评分'),
    ('peripheral_central_lr/top_lr_pairs.png', '图 S3 | 外周-中枢 Top LR pairs'),
    ('wgcna_microglia/module_corr_with_Pirb.png', '图 S4 | WGCNA 模块与 Pirb 相关性'),
]

# Figure panels for each report
report_figures = {
    "脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.docx": [
        ('GSE174574/marker_dotplot.png', '图 1 | GSE174574 细胞类型与 Pirb 表达'),
        ('microglia_cross_time/pirb_fraction_timeline.png', '图 2 | 跨数据集 Pirb 时间动态'),
        ('GSE225948/pirb_Mo_time.png', '图 3 | GSE225948 外周血 Pirb 表达'),
        ('GSE233814/pirb_spatial_combined_panel.png', '图 4 | GSE233814 空间定位 Pirb'),
        ('graphical_abstract_draft.png', '图 5 | 整合机制模型'),
    ],
    "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.docx": [
        ('GSE174574/marker_dotplot.png', '图 1 | GSE174574 细胞类型与 Pirb 表达'),
        ('microglia_cross_time/pirb_fraction_timeline.png', '图 2 | 跨数据集 Pirb 时间动态'),
        ('GSE225948/pirb_Mo_time.png', '图 3 | GSE225948 外周血 Pirb 表达'),
        ('GSE233814/pirb_spatial_combined_panel.png', '图 4 | GSE233814 空间定位 Pirb'),
        ('graphical_abstract_draft.png', '图 5 | 整合机制模型'),
    ],
    "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.docx": [
        ('GSE174574/marker_dotplot.png', 'Figure 1 | Pirb expression across cell types in GSE174574'),
        ('microglia_cross_time/pirb_fraction_timeline.png', 'Figure 2 | Cross-dataset Pirb temporal dynamics'),
        ('GSE225948/pirb_Mo_time.png', 'Figure 3 | Peripheral blood Pirb expression in GSE225948'),
        ('GSE233814/pirb_spatial_combined_panel.png', 'Figure 4 | Spatial localization of Pirb in GSE233814'),
        ('graphical_abstract_draft.png', 'Figure 5 | Integrated mechanistic model'),
    ],
}

for md_name, docx_name in reports:
    md_path = os.path.join(REPORT_DIR, md_name)
    docx_path = os.path.join(REPORT_DIR, docx_name)
    if not os.path.exists(md_path):
        print(f"[SKIP] {md_path} not found")
        continue
    
    elements = parse_md(md_path)
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    
    figs = extra_figs if '深度机制' in md_name else report_figures.get(docx_name, [])
    add_elements_to_doc(doc, elements, figs)
    doc.save(docx_path)
    print(f"[SAVE] {docx_path}")

print("[DONE]")
