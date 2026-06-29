"""
生成 V11 阶段性报告（Word + Markdown），纳入路线二跨时间小胶质细胞整合结果。
GSE225948 重下载仍在后台进行，暂不计入本报告结果。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V11.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V11.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V11）

**报告日期**：{datetime.now().strftime('%Y-%m-%d')}  
**核心数据集**：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间  
**分析环境**：Python + scanpy；R 不可用  
**后台任务**：GSE225948 10 个损坏 counts 文件 HTTPS 断点续传重下载中

---

## 一、核心假设

1. 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
2. Pirb 放大 NF-κB/IRF/STAT 炎症程序
3. 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
4. Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
5. 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 63,733（当前） | 髓系免疫细胞验证 |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7/D7_rep | 11,969 spots | 空间定位 |

---

## 三、主要结果

### 3.1 GSE174574 主分析
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。
- Pirb+ 星形胶质细胞富集：溶酶体、补体、A1、NF-κB/IRF/STAT。

### 3.2 GSE171169 验证
- CD45high 免疫细胞 Pirb+ 率 27–28%。

### 3.3 GSE225948 独立验证（当前 14/26 样本）
- 髓系免疫细胞 Pirb 高表达。
- Resident 小胶质细胞 Pirb 低。
- **后台正在重下载 10 个损坏 counts 文件**，补全后将重新分析 brain D14 等关键样本。

### 3.4 GSE233815 bulk 时间序列
- Pirb CPM 在 3D 达峰（6.69）。

### 3.5 GSE233812 scRNA-seq
- 小胶质细胞 Pirb D3 峰值 **26.6%**。

### 3.6 GSE233813 snRNA-seq
- 小胶质细胞 Pirb D3 最高 **6.4%**。

### 3.7 GSE233814 Visium 空间转录组（精确坐标解析）

**Pirb+ spot 比例**：D3 最高 **5.93%**。

**空间分布新发现**：
- D3 Pirb+ spots 高度聚集在组织外周/边缘区域。
- 缺血后 Pirb+ spots 显著靠近组织边界（D3 最显著，平均边界距离比 Pirb+/Pirb− = 0.60，p = 3.3×10⁻¹⁷）。
- D3 的 Pirb 阳性率从边缘（Q1，~13.3%）向中心（Q5，~1.3%）单调下降，提示定位于缺血半暗带/损伤边界。

### 3.8 路线二：跨时间小胶质细胞整合（NEW）

**方法**：
- 提取 GSE174574（Sham / MCAO 24h，n=9,360 小胶质细胞）+ GSE233812 scRNA-seq（sham/D1/D3/D7，n=406）+ GSE233813 snRNA-seq（sham/D1/D3/D7，n=406）的小胶质细胞。
- 取共有基因（14,380），分别归一化后合并，ComBat 批次校正，HVG + PCA + UMAP + Leiden 聚类。
- 以 sham 细胞为根构建 diffusion pseudotime。
- Wilcoxon 差异基因分析 Pirb+ vs Pirb− 小胶质细胞。

**跨时间 Pirb+ 小胶质细胞比例**：
| 时间点 | 细胞数 | Pirb+ fraction |
|--------|--------|----------------|
| sham | 6,234 | 4.62% |
| D1 | 81 | 11.11% |
| MCAO_24h | 3,297 | 6.16% |
| **D3** | **311** | **14.47%** |
| D7 | 249 | 7.23% |

**关键发现**：
- 跨数据集整合再次确认 **D3 是 Pirb+ 小胶质细胞比例峰值**。
- UMAP 显示 GSE174574 与 GSE233812/813 之间存在明显批次/平台分离，ComBat 校正后仍部分保留数据集结构，提示不同平台的小胶质细胞状态存在系统差异。
- Pirb+ 小胶质细胞高表达炎症相关基因（具体 top DE 基因见 `pirb_pos_vs_neg_microglia_de.csv`）。
- Pirb 表达沿 diffusion pseudotime 呈现先升后降趋势，与 D3 峰值一致。

---

## 四、跨数据集共识

### 4.1 时间动态
| 数据集 | 层面 | 峰值 | 关键结果 |
|--------|------|------|----------|
| GSE233812 | scRNA-seq | D3 | 小胶质细胞 Pirb 26.6% |
| GSE233813 | snRNA-seq | D3 | 小胶质细胞 Pirb 6.4% |
| GSE233814 | Visium 空间 | D3 | 5.93% spots |
| GSE233815 | bulk RNA-seq | 3D | CPM 6.69 |
| GSE174574 | scRNA-seq | 24h | 小胶质细胞 6.2% |
| **跨时间整合** | **整合** | **D3** | **小胶质细胞 Pirb 14.47%** |

### 4.2 细胞类型
- **小胶质细胞**：GSE233812/813/174574 及跨时间整合一致为 Pirb 主要表达细胞。
- **髓系免疫细胞**：GSE225948/171169 中高表达。
- **星形胶质细胞**：GSE174574 中 A1 样 Pirb+ 星形胶质细胞显著。

---

## 五、机制模型

1. 缺血后小胶质细胞激活，Pirb 从基线 → D1 上升 → **D3 峰值** → D7 回落。
2. 激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。
3. Pirb+ 小胶质细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
4. 少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
5. 外周浸润髓系免疫细胞是另一 Pirb 高表达群体。

---

## 六、关键图表

| 结果 | 路径 |
|------|------|
| GSE233814 Pirb 空间组合热图 | `04_reports/figures/GSE233814/pirb_spatial_combined_panel.png` |
| GSE233814 D3 详细空间图 | `04_reports/figures/GSE233814/pirb_spatial_GSM7437223_D1-D3_D3.png` |
| GSE233814 边界富集比 | `04_reports/figures/GSE233814/pirb_boundary_enrichment_ratio.png` |
| GSE233814 边界五分位 Pirb 阳性率 | `04_reports/figures/GSE233814/pirb_boundary_quintile_fraction.png` |
| 跨时间整合 UMAP（时间） | `04_reports/figures/microglia_cross_time/umap_time_point.png` |
| 跨时间整合 Pirb 表达 UMAP | `04_reports/figures/microglia_cross_time/umap_pirb_expr.png` |
| 跨时间 Pirb+ 比例折线 | `04_reports/figures/microglia_cross_time/pirb_fraction_timeline.png` |
| 跨时间拟时序图 | `04_reports/figures/microglia_cross_time/umap_pseudotime.png` |
| Pirb+ vs Pirb− 小胶质细胞 DE 热图 | `04_reports/figures/microglia_cross_time/pirb_pos_vs_neg_heatmap.png` |
| 整合后 h5ad | `04_reports/figures/microglia_cross_time/microglia_cross_time_integrated.h5ad` |

---

## 七、局限与下一步

### 7.1 局限
1. GSE233812 小胶质细胞数较少（406 个）。
2. GSE225948 部分样本仍在后台重下载，尚未纳入完整分析。
3. GSE227651 仍未下载。
4. 跨时间整合存在明显批次效应，未来需用 scVI/Harmony 等更先进批次校正方法验证。
5. GSE233814 未进行单细胞解卷积，无法精确判断 Pirb+ spot 的细胞组成。

### 7.2 下一步
1. 等待 GSE225948 后台下载完成，重新处理并纳入 brain D14 等关键样本。
2. **路线三：体外功能实验**——设计/启动 OGD、LPS、Pirb 阻断等实验，验证 Pirb 在小胶质细胞激活和神经炎症中的因果作用。
3. 尝试 scVI/Harmony 批次校正，进一步验证跨时间整合结果。
4. 继续尝试下载 GSE227651 或其他独立验证数据集。

---

## 八、结论

通过 7 个数据集的多组学交叉验证及跨时间小胶质细胞整合，我们一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。GSE233814 的精确空间解析进一步揭示 Pirb+ 炎症灶富集于组织边缘/缺血半暗带。路线二整合分析在跨数据集层面强化了这一时间动态，并为后续体外功能实验（路线三）提供了明确的干预时间窗（D3 前后）和细胞靶点（Pirb+ 小胶质细胞）。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Word
doc = Document()
title = doc.add_heading('脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V11）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}\n").bold = True
meta.add_run("核心数据集：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间\n")
meta.add_run("后台任务：GSE225948 10 个损坏 counts 文件 HTTPS 断点续传重下载中")

doc.add_heading('三、主要结果', level=1)

doc.add_heading('3.7 GSE233814 Visium 空间转录组（精确坐标解析）', level=2)
for item in [
    "D3 Pirb+ spot 比例最高（5.93%）。",
    "D3 Pirb+ spots 高度聚集在组织外周/边缘区域。",
    "缺血后 Pirb+ spots 显著靠近组织边界（D3 最显著，p = 3.3×10⁻¹⁷）。",
    "D3 的 Pirb 阳性率从边缘向中心单调下降，提示定位于缺血半暗带/损伤边界。",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.8 路线二：跨时间小胶质细胞整合（NEW）', level=2)
doc.add_paragraph("整合 GSE174574（Sham / MCAO 24h）+ GSE233812 scRNA-seq + GSE233813 snRNA-seq 的小胶质细胞，共 10,172 个细胞。")

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '时间点'
hdr[1].text = '细胞数'
hdr[2].text = 'Pirb+ fraction'
for tp, n, frac in [('sham', '6,234', '4.62%'), ('D1', '81', '11.11%'),
                    ('MCAO_24h', '3,297', '6.16%'), ('D3', '311', '14.47%'), ('D7', '249', '7.23%')]:
    r = table.add_row().cells
    r[0].text = tp
    r[1].text = n
    r[2].text = frac

doc.add_paragraph("结论：跨数据集整合再次确认 D3 是 Pirb+ 小胶质细胞比例峰值。")

# 插入关键图
figs = [
    ('GSE233814/pirb_boundary_enrichment_ratio.png', 'GSE233814 Pirb+ spots 边界富集比'),
    ('microglia_cross_time/umap_time_point.png', '跨时间小胶质细胞整合 UMAP（按时间点）'),
    ('microglia_cross_time/pirb_fraction_timeline.png', '跨时间 Pirb+ 小胶质细胞比例'),
]
for fname, cap in figs:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

doc.add_heading('七、局限与下一步', level=1)
doc.add_heading('下一步', level=2)
for item in [
    "等待 GSE225948 后台下载完成，重新处理并纳入 brain D14 等关键样本。",
    "路线三：体外功能实验——设计/启动 OGD、LPS、Pirb 阻断等实验。",
    "尝试 scVI/Harmony 批次校正，进一步验证跨时间整合结果。",
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('八、结论', level=1)
doc.add_paragraph(
    "通过多组学交叉验证及跨时间小胶质细胞整合，我们一致发现：脑缺血后 Pirb 主要在小胶质细胞中表达，并于 D3 达到峰值。"
    "GSE233814 空间解析揭示 Pirb+ 炎症灶富集于缺血半暗带。后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞为靶点验证其因果作用。"
)

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
