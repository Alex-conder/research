"""
将关键分析图片插入到最终报告 DOCX 中
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCX_PATH = "D:/Pirb_stroke_project/04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖与多数据集验证阶段性报告.docx"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures/GSE174574"

# 图片插入配置：(章节关键词, 图片路径, 图片说明)
images_to_insert = [
    ("3.2 Pirb 表达的细胞类型格局", os.path.join(FIG_DIR, "umap_celltype.png"), "图 1. GSE174574 细胞类型注释与 Pirb 表达 UMAP"),
    ("3.2 Pirb 表达的细胞类型格局", os.path.join(FIG_DIR, "pirb_positive_rate_celltype.png"), "图 2. 各细胞类型 Pirb 阳性率（MCAO vs Sham）"),
    ("3.3 Pirb+ 星形胶质细胞具有 A1-like 炎症表型", os.path.join(FIG_DIR, "de_pirb_all_ct/volcano_Astrocyte.png"), "图 3. 星形胶质细胞 Pirb+ vs Pirb- 火山图"),
    ("3.3 Pirb+ 星形胶质细胞具有 A1-like 炎症表型", os.path.join(FIG_DIR, "de_pirb_all_ct/volcano_Immune.png"), "图 4. 免疫细胞 Pirb+ vs Pirb- 火山图"),
    ("3.3 Pirb+ 星形胶质细胞具有 A1-like 炎症表型", os.path.join(FIG_DIR, "de_pirb_all_ct/volcano_Endothelial.png"), "图 5. 内皮细胞 Pirb+ vs Pirb- 火山图"),
    ("3.5 上游转录因子", os.path.join(FIG_DIR, "tf_analysis/tf_heatmap_log2fc.png"), "图 6. 转录因子 log2FC 热图（Pirb+ vs Pirb-）"),
    ("3.6 配体-受体互作预测", os.path.join(FIG_DIR, "lr_pirb_specific/pirb_signal_heatmap_mcao.png"), "图 7. Pirb 特异性配体-受体信号热图（MCAO）"),
    ("3.8 星形胶质细胞伪时间轨迹", os.path.join(FIG_DIR, "pseudotime/pirb_trend_pseudotime.png"), "图 8. Pirb 表达沿星形胶质细胞伪时间变化趋势"),
]

doc = Document(DOCX_PATH)

# 设置中文字体
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 记录已插入的章节，避免同一章节重复查找
inserted_sections = set()

# 从后向前遍历段落，找到目标章节后在其后插入图片
for section_keyword, img_path, caption in reversed(images_to_insert):
    if not os.path.exists(img_path):
        print(f"[WARN] Image not found: {img_path}")
        continue
    
    # 找到包含该章节标题的段落
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        para = doc.paragraphs[i]
        if section_keyword in para.text:
            # 在该段落后插入图片说明和图片
            # 需要在 paragraph i 之后添加新段落
            # docx 没有直接 insert_paragraph_before/after，但可以用 _element
            new_para = para._element.addnext(doc.add_paragraph()._element)
            # 添加图片说明
            cap_p = doc.add_paragraph(caption)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.runs[0]
            cap_run.bold = True
            cap_run.font.size = Pt(10)
            # 插入图片
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            # 将新段落移动到目标段落后
            para._element.addnext(cap_p._element)
            cap_p._element.addnext(img_p._element)
            print(f"[INSERTED] {caption} -> {img_path}")
            break

doc.save(DOCX_PATH)
print(f"[DONE] Updated: {DOCX_PATH}")
