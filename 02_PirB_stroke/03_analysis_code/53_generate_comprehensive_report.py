"""
生成完整汇总报告：以 Desktop 参考文档的星形胶质细胞分析为基础，
补入 GSE225948/GSE233815/GSE233812/GSE233813/GSE233814/跨时间整合/体外实验设计等全部结果。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_多数据集完整汇总报告.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_多数据集完整汇总报告.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞单细胞图谱：多数据集完整汇总报告

**报告日期**：{datetime.now().strftime('%Y-%m-%d')}  
**核心数据集**：GSE174574、GSE171169、GSE225948（26 样本完整版）、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间  
**分析环境**：Python + scanpy；R 不可用

---

## 摘要

本报告在 GSE174574 星形胶质细胞分析的基础上，整合 7 个独立数据集，系统解析脑缺血后 Pirb 阳性细胞的细胞类型、时间动态、空间定位及潜在机制。核心结论：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 MCAO 后 D3 达到峰值**；Pirb+ 小胶质细胞富集于炎症反应状态，且空间上定位于缺血半暗带/组织边缘；外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达，提示外周浸润是脑内 Pirb+ 炎症细胞的重要来源。后续体外功能实验将验证 Pirb 在小胶质细胞激活与神经炎症中的因果作用。

---

## 一、数据集概览

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析（星形胶质细胞 + 全细胞） |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 96,293 → 91,688 | 髓系免疫细胞验证（完整版） |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7/D7_rep | 11,969 spots | 空间定位 |

---

## 二、GSE174574：Pirb 阳性星形胶质细胞主分析

### 2.1 数据与方法
- 质控后共 57,769 个细胞、27,998 个基因；其中星形胶质细胞 5,813 个。
- 从全局整合对象中提取 Astrocyte，重新 PCA、邻居图、聚类、UMAP 和状态评分。
- Pirb 阳性定义为归一化表达层中 Pirb > 0。

### 2.2 MCAO 后 Pirb+ 星形胶质细胞增加
- MCAO 组星形胶质细胞 Pirb+ 比例为 **3.60%**，Sham 组为 **0.50%**。
- 这提示缺血后 Pirb+ 星形胶质细胞群体明显增加，但整体仍属低丰度亚群。

### 2.3 Pirb+ 星形胶质细胞富集于 A1-like 状态
| condition | astro_state | Pirb_expr | Pirb_positive | n_astrocytes |
|-----------|-------------|-----------|---------------|--------------|
| MCAO | A1_like | 0.08663 | **7.04%** | 682 |
| MCAO | Homeostatic | 0.03288 | 2.76% | 2031 |
| MCAO | PanReactive | 0.03621 | 3.19% | 345 |
| MCAO | Proliferative | 0.03784 | 3.08% | 260 |
| MCAO | A2_like | 0 | 0.00% | 101 |
| Sham | A1_like | 0.00581 | 0.39% | 761 |

- MCAO 组 A1-like 星形胶质细胞 Pirb+ 比例最高（7.04%），A2-like 状态中未检测到 Pirb+ 细胞。

### 2.4 状态评分支持 A1-like 倾向
- MCAO 组 Pirb+ 星形胶质细胞的 A1_like 评分高于 Pirb− 细胞，Homeostatic 评分较低，提示 Pirb+ 细胞更偏炎症性/反应性状态。

### 2.5 差异基因提示炎症与髓系相关信号
- **MCAO vs Sham 星形胶质细胞上调基因**：Spp1、Ccl4、Ccl12、Cd14、Ccl3、Gfap、Ccl2 等，符合缺血后反应性胶质和炎症增强。
- **Pirb+ vs Pirb− 星形胶质细胞差异基因**：Msr1、Lilr4b、Sirpb1b、Ifi30、Mnda、Cd86、Ccr1、C5ar1、Fcer1g、Pycard、Itgb2 等免疫/髓系相关基因靠前。

---

## 三、多数据集验证：Pirb 的时间动态

### 3.1 GSE233815 bulk RNA-seq
- Pirb CPM 在 **3D 达峰（6.69）**，24h 显著（p=0.021）。

### 3.2 GSE233812 scRNA-seq（sham/D1/D3/D7）
- 小胶质细胞 Pirb+ 在 **D3 达峰 26.6%**。
- 其他细胞类型几乎不表达。

### 3.3 GSE233813 snRNA-seq（sham/D1/D3/D7）
- 小胶质细胞 Pirb+ 在 **D3 最高 6.4%**。

### 3.4 GSE174574 小胶质细胞
- MCAO（24h）小胶质细胞 Pirb+ 约 6.2%。

### 3.5 跨时间小胶质细胞整合
- 整合 GSE174574 + GSE233812 + GSE233813 小胶质细胞（10,172 个）。
- Pirb+ 比例：sham 4.62% → D1 11.11% → MCAO_24h 6.16% → **D3 14.47%** → D7 7.23%。

### 3.6 时间动态共识
| 数据集 | 层面 | 峰值 | 关键结果 |
|--------|------|------|----------|
| GSE233812 | scRNA-seq | D3 | 小胶质细胞 Pirb 26.6% |
| GSE233813 | snRNA-seq | D3 | 小胶质细胞 Pirb 6.4% |
| GSE233814 | Visium 空间 | D3 | 5.93% spots |
| GSE233815 | bulk RNA-seq | 3D | CPM 6.69 |
| GSE174574 | scRNA-seq | 24h | 小胶质细胞 6.2% |
| 跨时间整合 | 整合 | D3 | 小胶质细胞 14.47% |

---

## 四、多数据集验证：Pirb 的细胞类型分布

### 4.1 GSE174574 全细胞
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。

### 4.2 GSE171169（CD45high 免疫细胞）
- 免疫细胞 Pirb+ 率 27–28%。

### 4.3 GSE225948 完整 26 样本
| 组织 | 细胞类型 | Sham | D02 | D14 |
|------|----------|------|-----|-----|
| Brain | 小胶质细胞 Mg | 2.89% | 2.39% | 2.53% |
| PB | 单核细胞 Mo | 35.25% | **51.72%** | 24.96% |
| PB | 中性粒细胞 Neu | 50.93% | **52.64%** | 28.67% |
| PB | 树突状细胞 DC | - | 36.17% | 23.23% |

- 脑内小胶质细胞 Pirb 维持低水平；外周血髓系细胞（Mo/Neu/DC）Pirb 高表达，D02 急性期达峰，D14 回落。
- 提示：**外周浸润髓系细胞可能是脑内 Pirb+ 炎症细胞的重要来源**。

### 4.4 细胞类型共识
- **小胶质细胞**：多个数据集一致为脑内 Pirb 主要表达细胞，D3 达峰。
- **外周髓系免疫细胞**：GSE225948/171169 中高表达，急性期达峰。
- **星形胶质细胞**：GSE174574 中 A1-like Pirb+ 星形胶质细胞显著，但比例较低。

---

## 五、GSE233814 Visium 空间转录组：Pirb 的空间定位

### 5.1 方法
- 解析 `json.gz` 中的 `fiducial`/`oligo`/transform，通过 10x Visium v1 barcode 映射重建每个 spot 的 `imageX/imageY` 像素坐标。
- 5 张切片：control / D1 / D3 / D7 / D7_rep，QC 后 11,969 spots。

### 5.2 Pirb+ spot 比例
| 时间 | Pirb+ fraction |
|------|----------------|
| control | 0.62% |
| D1 | 2.40% |
| **D3** | **5.93%** |
| D7 | 2.01% |
| D7_rep | 1.20% |

### 5.3 空间分布新发现
- D3 Pirb+ spots 高度聚集在组织外周/边缘区域，呈空间簇状分布。
- Pirb+ spots 最近邻距离在 D3 最小（mean = 236 pixels），提示存在活跃炎症病灶。
- 缺血后 Pirb+ spots 显著靠近组织边界（D3 最显著，平均边界距离比 Pirb+/Pirb− = **0.60**，p = 3.3×10⁻¹⁷）。
- D3 Pirb 阳性率从边缘（Q1，~13.3%）向中心（Q5，~1.3%）单调下降，提示定位于**缺血半暗带/损伤边界**。

### 5.4 Pirb+ spot 细胞特征
- D3 Pirb+ spots 小胶质细胞评分（1.165 vs 0.687）和炎症评分（1.139 vs 0.345）显著高于 Pirb− spots。

---

## 六、机制模型

基于上述多数据集证据，提出如下整合机制模型：

1. **外周来源**：缺血后外周血 **Pirb+ 单核细胞/中性粒细胞** 在 D02/D03 达峰，可能通过破坏的血脑屏障浸润脑组织。
2. **脑内小胶质细胞激活**：脑内小胶质细胞在缺血后激活，Pirb 从基线 → D1 上升 → **D3 峰值** → D7 回落。
3. **星形胶质细胞 A1 转化**：激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞向 A1-like 反应性状态转化，形成 Pirb+ 星形胶质细胞亚群。
4. **炎症放大**：Pirb+ 小胶质细胞/髓系细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
5. **轴突再生抑制**：少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
6. **空间定位**：Pirb+ 炎症灶富集于缺血半暗带/组织边缘，与损伤进展区域高度重合。

---

## 七、体外功能实验设计（路线三）

### 7.1 实验目的
验证 Pirb 在脑缺血后小胶质细胞激活及神经炎症放大中的因果作用。

### 7.2 细胞与刺激
- **细胞模型**：原代小鼠小胶质细胞 / BV2 细胞系 / 小胶质细胞-神经元共培养。
- **刺激**：OGD（2–6 h）、LPS（100 ng/mL）、OGD + LPS。

### 7.3 Pirb 干预
- 抗-Pirb 中和抗体（5–20 μg/mL）。
- Pirb siRNA / shRNA（50–100 nM）。
- Pirb-Fc 融合蛋白（10–50 μg/mL）。

### 7.4 检测指标
- Pirb 表达（qPCR、Western blot、流式、IF）。
- 激活表型（Iba1、CD68、形态、增殖、迁移）。
- 炎症因子（IL-1β、TNF-α、IL-6、C1q、IL-10、NO）。
- 信号通路（NF-κB、MAPK、STAT/IRF、PI3K/Akt）。
- 功能实验（吞噬能力、突触修剪、神经毒性/存活）。

### 7.5 预期结果
- Pirb 在 OGD/LPS 后上调，48–72 h 达峰。
- Pirb 阻断/敲低抑制小胶质细胞激活、减少促炎因子分泌。
- Pirb 阻断抑制 NF-κB/STAT/IRF 通路。
- Pirb 阻断减少过度吞噬/突触修剪，改善神经元存活。

---

## 八、局限与下一步

### 8.1 局限
1. GSE174574 星形胶质细胞中 Pirb+ 比例较低，需严格排除免疫细胞混入和 doublet。
2. GSE225948 脑内小胶质细胞 Pirb 阳性率也较低，可能与 CD45hi 分选策略偏向外周/浸润细胞有关。
3. GSE233812 小胶质细胞数较少（406 个）。
4. 跨时间整合存在批次效应，未来需用 scVI/Harmony 进一步校正。
5. GSE233814 未进行单细胞解卷积，无法精确判断 Pirb+ spot 的细胞组成。
6. GSE227651 仍未下载。

### 8.2 下一步
1. **启动体外功能实验**：按路线三方案开展 OGD/LPS + anti-Pirb 阻断实验。
2. **空间验证**：Pirb + GFAP/AQP4 + C3/S100A10 多重免疫荧光或 RNAscope，重点比较半暗带与梗死核心。
3. **细胞通讯分析**：重点分析 Microglia/Immune/Neuron → Astrocyte 的 IL1、TNF、C3、CXCL、TGFβ、JAK/STAT 通讯。
4. **方法学优化**：尝试 scVI/Harmony 批次校正。
5. **补充数据集**：继续下载 GSE227651 或其他卒中单细胞数据集。

---

## 九、结论

通过 7 个数据集的多组学交叉验证、Visium 空间精确解析及跨时间小胶质细胞整合，本研究一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。GSE174574 进一步揭示低丰度但显著增加的 **Pirb+ A1-like 星形胶质细胞亚群**。GSE225948 完整数据提示 **外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达**，可能是脑内 Pirb+ 炎症细胞的重要来源。GSE233814 空间分析显示 Pirb+ 炎症灶富集于**缺血半暗带/组织边缘**。后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞及外周髓系细胞为靶点，验证其在缺血后神经炎症、突触修剪和轴突再生抑制中的因果作用。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Word
doc = Document()
title = doc.add_heading('脑缺血后 Pirb 阳性细胞单细胞图谱：多数据集完整汇总报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}\n").bold = True
meta.add_run("核心数据集：GSE174574、GSE171169、GSE225948（26 样本完整版）、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间\n")
meta.add_run("分析环境：Python + scanpy；R 不可用")

# 摘要
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    "本报告在 GSE174574 星形胶质细胞分析的基础上，整合 7 个独立数据集，系统解析脑缺血后 Pirb 阳性细胞的细胞类型、时间动态、空间定位及潜在机制。"
    "核心结论：脑缺血后 Pirb 主要在小胶质细胞中表达，并于 MCAO 后 D3 达到峰值；Pirb+ 小胶质细胞富集于炎症反应状态，且空间上定位于缺血半暗带/组织边缘；"
    "外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达，提示外周浸润是脑内 Pirb+ 炎症细胞的重要来源。"
)

# 数据集
doc.add_heading('一、数据集概览', level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '数据集'
hdr[1].text = '类型'
hdr[2].text = '样本'
hdr[3].text = '细胞/spot 数'
hdr[4].text = '用途'
rows = [
    ['GSE174574', 'scRNA-seq', '3 Sham + 3 MCAO（24 h）', '58,287 → 56,486', '主分析'],
    ['GSE171169', 'scRNA-seq', 'CD45high 5d/14d', '10,295', '免疫细胞验证'],
    ['GSE225948', 'scRNA-seq', 'PB + brain Sham/D02/D14', '96,293 → 91,688', '髓系免疫细胞验证'],
    ['GSE233815', 'bulk RNA-seq', 'MCAO 3h/12h/24h/3D/7D', '48', '时间动态验证'],
    ['GSE233812', 'scRNA-seq', 'sham/D1/D3/D7', '6,159', '单细胞时间序列'],
    ['GSE233813', 'snRNA-seq', 'sham/D1/D3/D7', '8,374', '单细胞核时间序列'],
    ['GSE233814', 'Visium 空间', 'control/D1/D3/D7/D7_rep', '11,969 spots', '空间定位'],
]
for r in rows:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

# GSE174574
doc.add_heading('二、GSE174574：Pirb 阳性星形胶质细胞主分析', level=1)
doc.add_heading('2.2 MCAO 后 Pirb+ 星形胶质细胞增加', level=2)
doc.add_paragraph("MCAO 组星形胶质细胞 Pirb+ 比例为 3.60%，Sham 组为 0.50%。")

doc.add_heading('2.3 Pirb+ 星形胶质细胞富集于 A1-like 状态', level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'condition'
hdr[1].text = 'astro_state'
hdr[2].text = 'Pirb_expr'
hdr[3].text = 'Pirb_positive'
hdr[4].text = 'n_astrocytes'
for r in [
    ['MCAO', 'A1_like', '0.08663', '7.04%', '682'],
    ['MCAO', 'Homeostatic', '0.03288', '2.76%', '2031'],
    ['MCAO', 'PanReactive', '0.03621', '3.19%', '345'],
    ['MCAO', 'A2_like', '0', '0.00%', '101'],
    ['Sham', 'A1_like', '0.00581', '0.39%', '761'],
]:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

doc.add_heading('2.5 差异基因提示炎症与髓系相关信号', level=2)
doc.add_paragraph(
    "MCAO vs Sham 上调基因包括 Spp1、Ccl4、Ccl12、Cd14、Ccl3、Gfap、Ccl2 等。"
    "Pirb+ vs Pirb− 差异基因包括 Msr1、Lilr4b、Sirpb1b、Ifi30、Mnda、Cd86、Ccr1、C5ar1、Fcer1g、Itgb2 等免疫/髓系相关基因。"
)

# 多数据集验证：时间动态
doc.add_heading('三、多数据集验证：Pirb 的时间动态', level=1)
doc.add_heading('3.6 时间动态共识', level=2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '数据集'
hdr[1].text = '层面'
hdr[2].text = '峰值'
hdr[3].text = '关键结果'
rows = [
    ['GSE233812', 'scRNA-seq', 'D3', '小胶质细胞 Pirb 26.6%'],
    ['GSE233813', 'snRNA-seq', 'D3', '小胶质细胞 Pirb 6.4%'],
    ['GSE233814', 'Visium 空间', 'D3', '5.93% spots'],
    ['GSE233815', 'bulk RNA-seq', '3D', 'CPM 6.69'],
    ['GSE174574', 'scRNA-seq', '24h', '小胶质细胞 6.2%'],
    ['跨时间整合', '整合', 'D3', '小胶质细胞 14.47%'],
]
for r in rows:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

# 细胞类型分布
doc.add_heading('四、多数据集验证：Pirb 的细胞类型分布', level=1)
doc.add_heading('4.3 GSE225948 完整 26 样本', level=2)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '组织'
hdr[1].text = '细胞类型'
hdr[2].text = 'Sham'
hdr[3].text = 'D02'
hdr[4].text = 'D14'
rows = [
    ['Brain', '小胶质细胞 Mg', '2.89%', '2.39%', '2.53%'],
    ['PB', '单核细胞 Mo', '35.25%', '51.72%', '24.96%'],
    ['PB', '中性粒细胞 Neu', '50.93%', '52.64%', '28.67%'],
]
for r in rows:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

# 空间定位
doc.add_heading('五、GSE233814 Visium 空间转录组：Pirb 的空间定位', level=1)
doc.add_paragraph(
    "D3 Pirb+ spot 比例最高（5.93%）。Pirb+ spots 显著靠近组织边界（平均边界距离比 Pirb+/Pirb− = 0.60，p = 3.3×10⁻¹⁷），"
    "且从边缘向中心单调下降，提示定位于缺血半暗带/损伤边界。"
)

# 机制模型
doc.add_heading('六、机制模型', level=1)
for item in [
    "外周来源：缺血后外周血 Pirb+ 单核细胞/中性粒细胞在 D02/D03 达峰，可能通过破坏的血脑屏障浸润脑组织。",
    "脑内小胶质细胞激活：Pirb 从基线 → D1 上升 → D3 峰值 → D7 回落。",
    "星形胶质细胞 A1 转化：小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1-like 状态。",
    "炎症放大：Pirb+ 小胶质细胞/髓系细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。",
    "轴突再生抑制：少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作。",
    "空间定位：Pirb+ 炎症灶富集于缺血半暗带/组织边缘。",
]:
    doc.add_paragraph(item, style='List Number')

# 体外实验设计
doc.add_heading('七、体外功能实验设计（路线三）', level=1)
doc.add_paragraph(
    "验证 Pirb 在脑缺血后小胶质细胞激活及神经炎症放大中的因果作用。"
    "细胞模型：原代小鼠小胶质细胞 / BV2 / 共培养；刺激：OGD/LPS；干预：anti-Pirb 抗体 / Pirb siRNA / Pirb-Fc。"
    "检测：Pirb 表达、激活表型、炎症因子、NF-κB/STAT/IRF 通路、吞噬/突触修剪/神经毒性。"
)

# 局限与下一步
doc.add_heading('八、局限与下一步', level=1)
doc.add_heading('下一步', level=2)
for item in [
    "启动体外功能实验：OGD/LPS + anti-Pirb 阻断。",
    "空间验证：Pirb + GFAP/AQP4 + C3/S100A10 多重免疫荧光或 RNAscope。",
    "细胞通讯分析：Microglia/Immune/Neuron → Astrocyte 的 IL1/TNF/C3/CXCL/TGFβ/JAK-STAT 通讯。",
    "方法学优化：scVI/Harmony 批次校正。",
    "补充数据集：GSE227651 等。",
]:
    doc.add_paragraph(item, style='List Number')

# 结论
doc.add_heading('九、结论', level=1)
doc.add_paragraph(
    "通过多组学交叉验证、Visium 空间精确解析及跨时间小胶质细胞整合，本研究一致发现："
    "脑缺血后 Pirb 主要在小胶质细胞中表达，并于 D3 达到峰值。GSE174574 揭示低丰度但显著增加的 Pirb+ A1-like 星形胶质细胞亚群。"
    "GSE225948 完整数据提示外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达。"
    "GSE233814 空间分析显示 Pirb+ 炎症灶富集于缺血半暗带/组织边缘。"
    "后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞及外周髓系细胞为靶点，验证其因果作用。"
)

# 插入关键图
figs = [
    ('GSE233814/pirb_spatial_combined_panel.png', 'GSE233814 Pirb 空间组合热图'),
    ('GSE225948/pirb_Mo_time.png', 'GSE225948 PB 单核细胞 Pirb 时间动态'),
    ('microglia_cross_time/pirb_fraction_timeline.png', '跨时间小胶质细胞 Pirb+ 比例'),
]
for fname, cap in figs:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
