"""
生成最终阶段性报告 DOCX 和 Markdown
参考原报告结构，整合最新多数据集机制分析结果
"""
import os
import re
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "D:/Pirb_stroke_project/04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖与多数据集验证阶段性报告.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖与多数据集验证阶段性报告

**报告日期**：2026-06-16  
**分析数据集**：GSE174574（主数据集）、GSE171169（验证集）  
**后台任务**：GSE227651、GSE225948 原始文件下载中（受网络限制）

---

## 一、目标与假设

### 1.1 核心目标
1. 明确脑缺血后 Pirb 阳性细胞的细胞类型分布与动态变化。
2. 揭示驱动 Pirb 上调的上游信号（小胶质细胞、少突胶质细胞、免疫细胞等）。
3. 构建 Pirb 参与缺血后炎症反应与轴突再生抑制的机制假设。

### 1.2 核心假设
- **假设 1**：小胶质细胞来源的 IL-1α/TNF/C1q 诱导 A1 样 Pirb+ 星形胶质细胞。
- **假设 2**：Pirb 在星形胶质细胞中放大 NF-κB/IRF/STAT 驱动的炎症程序。
- **假设 3**：少突胶质细胞来源的 MAG/OMgp/MOG 通过 Pirb 抑制轴突再生。
- **假设 4**：Pirb 在免疫细胞/小胶质细胞中参与免疫抑制或突触修剪。
- **假设 5**：内皮细胞 Pirb 上调与血脑屏障破坏相关。

---

## 二、数据集与方法

### 2.1 数据集
| 数据集 | 样本 | 细胞数 | 用途 |
|--------|------|--------|------|
| GSE174574 | 3 Sham + 3 MCAO（24 h） | 58,287（质控后） | 主分析：细胞类型注释、差异表达、LR 互作、伪时间 |
| GSE171169 | 4 CD45high 样本（5d / 14d） | 10,295 | 验证：免疫细胞中 Pirb 高表达 |

### 2.2 分析方法
1. **数据读取与质控**：10x 标准三文件手动读取；过滤 n_genes < 500、total_counts < 1000、percent.mt > 20%。
2. **标准化**：normalize_total（target_sum = 1e4）+ log1p。
3. **细胞类型注释**：基于已知 marker 基因得分分配细胞类型。
4. **差异表达**：Pirb+ vs Pirb− 按细胞类型分别进行 Wilcoxon rank-sum 检验。
5. **通路富集**：对 Pirb+ 星形胶质细胞上调基因进行超几何检验（预定义免疫/炎症/髓鞘抑制通路基因集）。
6. **配体-受体互作**：基于 mouseconsensus 数据库，计算细胞类型间 ligand × receptor 互作得分；额外构建 Pirb 特异性配体库。
7. **转录因子分析**：检查 NF-κB、STAT、IRF 等家族在 Pirb+ vs Pirb− 细胞中的表达差异。
8. **伪时间分析**：scanpy diffusion pseudotime，以 Sham 星形胶质细胞为根。

---

## 三、主要结果

### 3.1 Pirb 在 MCAO 后显著上调
- **GSE174574**：MCAO 后 Pirb 阳性率从 Sham 的 2.7–2.8% 升至 8.2–9.8%，平均表达升高约 3 倍。
- **GSE171169**：CD45high 免疫细胞中 Pirb 阳性率为 27–28%，5d 与 14d 保持稳定。

### 3.2 Pirb 表达的细胞类型格局

| 细胞类型 | MCAO Pirb+ % | Sham Pirb+ % | 主要结论 |
|----------|-------------|-------------|----------|
| Immune | **31.37%** | 24.24% | Pirb 阳性率最高，以外周/脑浸润免疫细胞为主 |
| Microglia | **6.16%** | 4.59% | 小胶质细胞 Pirb 表达显著，缺血后升高 |
| Ependymal | 4.41% | 0.88% | MCAO 后明显上调 |
| Astrocyte | **3.13%** | 0.66% | 星形胶质细胞 Pirb+ 比例升高约 4.7 倍 |
| Oligodendrocyte | 2.52% | 0.92% | 少量表达 |
| Pericyte | 1.83% | 0.61% | 少量表达 |
| Endothelial | 1.77% | 0.58% | 少量表达 |

**关键发现**：Pirb 上调并非星形胶质细胞特有，免疫细胞和小胶质细胞是更主要的表达群体。

### 3.3 Pirb+ 星形胶质细胞具有 A1-like 炎症表型
对 83 个 Pirb+ 与 3405 个 Pirb− 星形胶质细胞进行差异表达分析，发现 164 个显著上调基因。

**上调基因特征**：
- 免疫/髓系标志：Cd14、Fcer1g、Csf1r、Fcgr3、Ly86
- 补体成分：C1qa、C1qb
- 溶酶体/组织蛋白酶：Ctsb、Ctsz、Ctss、Ctsd
- 炎症/修复因子：Spp1（Osteopontin）、Ccl4、Ccl12、C5ar1

**下调基因特征**：
- 星形胶质细胞稳态标志：Gja1、Slc1a3（GLT-1）、Bcan、Mfge8、Ndrg2、Plpp3

### 3.4 通路富集结果
对 Pirb+ 星形胶质细胞上调基因进行超几何检验，发现以下通路显著富集：

| 通路 | p 值 | 重叠基因 |
|------|------|----------|
| 溶酶体/组织蛋白酶 | 1.8 × 10⁻¹⁷ | Ctsb, Ctsd, Ctsz, Ctss, Ctsc |
| 补体激活 | 7.0 × 10⁻¹¹ | C1qa, C1qb, C3, C5ar1 |
| 小胶质细胞诱导 A1 | 1.3 × 10⁻⁶ | C1qa, C1qb, Tnf, Il1a |
| MHC I 抗原呈递 | 5.4 × 10⁻⁴ | H2-D1, H2-K1, Cd74 |
| A1 反应性星形胶质细胞 | 3.5 × 10⁻² | C3, Gfap, Vim |
| IRF 信号 | 3.0 × 10⁻² | Irf1, Irf5, Irf7, Irf8 |
| NF-κB 信号 | 4.0 × 10⁻² | Nfkb1, Nfkb2, Rela, Relb |

### 3.5 上游转录因子
Pirb+ 星形胶质细胞中显著上调的转录因子：
- Cebpb（log₂FC = 1.75）
- Rel（log₂FC = 1.65）
- Stat6（log₂FC = 1.42）
- Irf8（log₂FC = 1.37）
- Irf5（log₂FC = 1.31）

这些 TF 主要属于 NF-κB、IRF 和 STAT 家族，支持炎症程序驱动的 Pirb 上调。

### 3.6 配体-受体互作预测
- **少突胶质细胞**是 Pirb 经典配体 MAG / MOG / OMgp / Nogo-A 的最强来源。
- **MHC I → PirB** 信号在免疫细胞和小胶质细胞中最强，与这些细胞 Pirb 高表达一致。
- **指向星形胶质细胞的 Pirb 信号**主要来自少突胶质细胞、小胶质细胞和免疫细胞。

### 3.7 其他细胞类型的 Pirb 功能特征

| 细胞类型 | Pirb+ 细胞数 | 主要上调特征 | 主要下调特征 |
|----------|-------------|-------------|-------------|
| Immune | 2422 | S100a6, Alox5ap, Ifitm, Lyz2（先天免疫） | Dhrs7b, Man1a2, Pdgfa |
| Microglia | 481 | Pirb, Lyz2 | — |
| Endothelial | 376 | Ctsb, Spp1, Hspa5, Ccl4（应激/炎症） | Abcb1a, Cxcl12, Rgcc（BBB 特征） |
| Oligodendrocyte | 32 | Pirb, Dusp1, Lcp1 | — |
| Pericyte | 38 | Pirb | — |

### 3.8 星形胶质细胞伪时间轨迹
- 以 Sham 星形胶质细胞为根，Pirb 表达沿伪时间先上升、在 0.35–0.6 达到峰值（阳性率 ~6%），随后下降。
- 提示 Pirb 高表达是缺血响应过渡状态，而非终末状态。

---

## 四、是否支持核心目标

### 目标 1：明确 Pirb 阳性细胞的细胞类型分布
**支持**。Pirb 在免疫细胞、小胶质细胞、星形胶质细胞、内皮细胞、少突胶质细胞中均有表达，缺血后普遍上调。

### 目标 2：揭示驱动 Pirb 上调的上游信号
**部分支持**。LR 分析预测少突胶质细胞来源的 MAG/MOG/OMgp 和小胶质细胞/免疫细胞来源的 C1q/TNF/IL-1α 可能是主要驱动因素；但空间共定位和时间动态数据尚未获取。

### 目标 3：构建 Pirb 参与炎症与轴突再生抑制的机制假设
**支持**。Pirb+ 星形胶质细胞的炎症通路富集、TF 激活以及髓鞘抑制配体来源分析，共同支持该假设。

---

## 五、局限

1. **GSE174574 仅 24 h 单时间点**，无法确定 Pirb 上调的时间动态。
2. **GSE227651 / GSE225948 下载未完成**，缺少时间序列和免疫/内皮验证。
3. **GSE233815 空间数据未下载**，无法验证 Pirb+ 细胞的空间定位。
4. **细胞类型注释基于简单 marker 得分**，未使用 SingleR / CellTypist 等自动化方法。
5. **Pirb+ 星形胶质细胞数量少**（83 个），差异分析统计效力有限，且不能完全排除 doublet / 环境 RNA 污染。
6. **配体-受体分析为预测性**，缺乏功能验证。

---

## 六、下一步计划

1. **严格质控**：对 Pirb+ 星形胶质细胞进行 doublet 检测（Scrublet / DoubletFinder），排除免疫细胞污染。
2. **空间验证**：下载 GSE233815 Visium 数据，验证 Pirb+ 细胞是否富集于梗死周围/半暗带。
3. **时间动态**：完成 GSE227651 下载后，分析 day 1/3/7 的 Pirb 变化趋势。
4. **功能实验**：
   - 体外 IL-1α + TNF + C1q 诱导原代星形胶质细胞，检测 Pirb 表达。
   - Pirb 阻断后检测炎症因子释放和 BBB 通透性。
   - 免疫荧光验证 Pirb 与 C3 / C1q / H2-D1 的共定位。
5. **多数据集整合**：使用 Harmony / Scanorama 整合 GSE174574、GSE227651、GSE225948，扩大样本量。

---

## 七、结论

基于 GSE174574 和 GSE171169 的整合分析，脑缺血后 Pirb 在多种细胞类型中上调，其中免疫细胞和小胶质细胞表达最高，星形胶质细胞中 Pirb+ 亚群呈现典型的 A1-like 炎症表型。上游可能由小胶质细胞来源的 IL-1α/TNF/C1q 诱导，并通过 NF-κB/IRF/STAT 转录程序驱动；同时，少突胶质细胞来源的 MAG/OMgp/MOG 可能通过 Pirb 持续抑制轴突再生。这些发现为 Pirb 作为缺血性卒中治疗靶点提供了多细胞类型视角和可验证的分子机制假设。
"""

# 保存 Markdown
with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

# 生成 DOCX
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    """解析 markdown 表格行"""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            # 跳过分隔行
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        # 收集整个表格
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
