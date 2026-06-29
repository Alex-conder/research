"""
生成 V10 阶段性报告（Word + Markdown），纳入 GSE233814 空间精确解析结果。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures/GSE233814"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V10.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V10.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V10）

**报告日期**：{datetime.now().strftime('%Y-%m-%d')}  
**核心数据集**：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间  
**分析环境**：Python + scanpy；R 不可用

---

## 一、核心假设

1. 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
2. Pirb 放大 NF-κB/IRF/STAT 炎症程序
3. 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
4. Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
5. 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 63,733 | 髓系免疫细胞验证 |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7/D7_rep | 11,969 spots | 空间定位 |

---

## 三、主要结果

### 3.1 GSE174574 主分析
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。
- Pirb+ 星形胶质细胞富集：溶酶体、补体、A1、NF-κB/IRF/STAT。

### 3.2 GSE171169 验证
- CD45high 免疫细胞 Pirb+ 率 27–28%。

### 3.3 GSE225948 独立验证
- 髓系免疫细胞 Pirb 高表达。
- Resident 小胶质细胞 Pirb 低。

### 3.4 GSE233815 bulk 时间序列
- Pirb CPM 在 3D 达峰（6.69）。

### 3.5 GSE233812 scRNA-seq
- 小胶质细胞 Pirb D3 峰值 **26.6%**。
- 其他细胞类型几乎不表达。

### 3.6 GSE233813 snRNA-seq
- 小胶质细胞 Pirb D3 最高 **6.4%**。

### 3.7 GSE233814 Visium 空间转录组（精确坐标解析）

**解析方法**：
1. 从 GEO 下载每个样本配套的 `json.gz` 空间文件，包含 `fiducial`（标记点）、`oligo`（spot 列表，含 `row/col` 与 `imageX/imageY` 像素坐标）以及 3×3 仿射变换矩阵 `transform`。
2. 利用 10x Genomics Visium v1 官方 barcode 到 `(row, col)` 的固定映射（`visium-v1.txt` + `visium-v1_coordinates.txt`，1-based 转 0-based），将每个样本 `barcodes.tsv.gz` 中的 spot barcode 映射到捕获阵列坐标。
3. 通过 `(row, col)` 匹配 `oligo` 条目，提取每个有组织 spot 在 H&E 图像上的像素坐标 `imageX`、`imageY`。
4. 将 Pirb 标准化表达（log1p CPM）叠加到空间像素坐标，绘制高分辨率空间分布热图，并进一步分析 Pirb+ spot 的空间聚类性与边界富集。

**Pirb+ spot 比例**：
| 时间 | Pirb+ fraction |
|------|----------------|
| control | 0.62% |
| D1 | 2.40% |
| **D3** | **5.93%** |
| D7 | 2.01% |
| D7_rep | 1.20% |

**Pirb+ spot 特征（D3）**：
| 特征 | Pirb+ spots | Pirb- spots |
|------|-------------|-------------|
| Microglia score | 1.165 | 0.687 |
| Inflammation score | 1.139 | 0.345 |

**空间分布新发现**：
- D3 Pirb+ spots 高度聚集在组织外周/边缘区域，呈现明显的空间簇状分布。
- Pirb+ spots 的最近邻距离在 D3 最小（mean = 236 pixels），提示 D3 存在活跃的小胶质细胞/炎症病灶。
- 缺血后（D1/D3/D7/D7_rep）Pirb+ spots 均显著靠近组织边界（Mann-Whitney U 检验 p < 0.01），而 control 无显著边界偏好：
  - control 平均边界距离比（Pirb+ / Pirb−）≈ 1.17（略偏中心，p = 0.33）
  - D1 ≈ 0.69（p = 2.7×10⁻⁵）
  - **D3 ≈ 0.60（p = 3.3×10⁻¹⁷）**
  - D7 ≈ 0.61（p = 3.8×10⁻⁷）
  - D7_rep ≈ 0.66（p = 6.5×10⁻³）
- 边界至中心五分位分析显示，D3 的 Pirb 阳性率从边缘（Q1，~13.3%）向中心（Q5，~1.3%）单调下降，进一步证实 Pirb+ 炎症灶定位于缺血半暗带/损伤边界。

**结论**：空间层面再次确认 D3 是 Pirb 表达峰值，且 Pirb+ 区域与小胶质细胞/炎症评分高度共定位，并倾向于分布在组织边缘——这与缺血半暗带的病理解剖高度一致。

---

## 四、跨数据集共识

### 4.1 时间动态
| 数据集 | 层面 | 峰值 | 关键结果 |
|--------|------|------|----------|
| GSE233812 | scRNA-seq | D3 | 小胶质细胞 Pirb 26.6% |
| GSE233813 | snRNA-seq | D3 | 小胶质细胞 Pirb 6.4% |
| GSE233814 | Visium 空间 | D3 | 5.93% spots |
| GSE233815 | bulk RNA-seq | 3D | CPM 6.69 |
| GSE174574 | scRNA-seq | 24h | 小胶质细胞 6.2% |

### 4.2 细胞类型
- **小胶质细胞**：GSE233812/813/174574 一致为 Pirb 主要表达细胞。
- **髓系免疫细胞**：GSE225948/171169 中高表达。
- **星形胶质细胞**：GSE174574 中 A1 样 Pirb+ 星形胶质细胞显著。

---

## 五、机制模型

1. 缺血后小胶质细胞激活，Pirb 从基线 → D1 上升 → **D3 峰值** → D7 回落。
2. 激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。
3. Pirb+ 小胶质细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
4. 少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。
5. 外周浸润髓系免疫细胞是另一 Pirb 高表达群体。

---

## 六、关键图表

| 结果 | 路径 |
|------|------|
| GSE233812 scRNA UMAP / Pirb 热图 | `04_reports/figures/GSE233812/` |
| GSE233813 snRNA UMAP / Pirb 热图 | `04_reports/figures/GSE233813/` |
| sc vs sn 小胶质细胞 Pirb 动态 | `04_reports/figures/GSE233812_813_comparison/` |
| GSE233814 Pirb 空间组合热图 | `04_reports/figures/GSE233814/pirb_spatial_combined_panel.png` |
| GSE233814 D3 详细空间图 | `04_reports/figures/GSE233814/pirb_spatial_GSM7437223_D1-D3_D3.png` |
| GSE233814 Pirb+ 空间聚类箱线图 | `04_reports/figures/GSE233814/pirb_spatial_clustering_boxplot.png` |
| GSE233814 边界距离小提琴图 | `04_reports/figures/GSE233814/pirb_boundary_distance_violin.png` |
| GSE233814 边界富集比 | `04_reports/figures/GSE233814/pirb_boundary_enrichment_ratio.png` |
| GSE233814 边界五分位 Pirb 阳性率 | `04_reports/figures/GSE233814/pirb_boundary_quintile_fraction.png` |
| GSE233814 spot 像素坐标 + Pirb 表达 | `04_reports/figures/GSE233814/spot_pixel_coordinates_pirb.csv` |
| GSE233814 空间模式摘要 | `04_reports/figures/GSE233814/pirb_spatial_pattern_summary.csv` |
| GSE233814 边界距离摘要 | `04_reports/figures/GSE233814/pirb_boundary_distance_summary.csv` |

---

## 七、局限与下一步

### 7.1 局限
1. GSE233812 小胶质细胞数较少（406 个）。
2. GSE225948 部分样本缺失（12 个 counts 文件损坏），brain D14 等关键样本未纳入。
3. GSE227651 仍未下载。
4. GSE233814 仅有 H&E 图像坐标，未能与单细胞解卷积结果直接叠加（需进一步细胞类型去卷积）。

### 7.2 下一步
1. **路线二：跨时间小胶质细胞整合**——整合 GSE174574（24h）+ GSE233812/813（D3），构建跨时间 Pirb+ 小胶质细胞状态转移图谱。
2. **路线三：体外功能实验**——设计/启动 OGD、LPS、Pirb 阻断等实验，验证 Pirb 在小胶质细胞激活和神经炎症中的因果作用。
3. 修复 GSE225948 损坏文件或寻找替代数据集。

---

## 八、结论

通过 7 个数据集的多组学交叉验证，我们一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。这一结论在单细胞、单细胞核、空间和组织 bulk 四个层面高度一致。GSE233814 的精确空间解析进一步揭示：Pirb+ spots 在 D3 高度聚集并显著富集于组织边缘/缺血半暗带，与小胶质细胞/炎症程序共定位，为 Pirb 作为脑缺血后神经炎症关键调控因子提供了更强有力的空间证据。外周浸润髓系免疫细胞是另一重要 Pirb 载体。后续功能实验将进一步验证 Pirb 在缺血后炎症放大和轴突再生抑制中的因果作用。
"""

# 保存 Markdown
with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# 生成 Word
doc = Document()

# 标题
title = doc.add_heading('脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V10）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 元信息
meta = doc.add_paragraph()
meta.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}\n").bold = True
meta.add_run("核心数据集：GSE174574、GSE171169、GSE225948、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间\n")
meta.add_run("分析环境：Python + scanpy；R 不可用")

# 核心假设
doc.add_heading('一、核心假设', level=1)
for item in [
    "小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞",
    "Pirb 放大 NF-κB/IRF/STAT 炎症程序",
    "少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生",
    "Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪",
    "内皮细胞 Pirb 上调与血脑屏障破坏相关",
]:
    doc.add_paragraph(item, style='List Number')

# 数据集
doc.add_heading('二、数据集', level=1)
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '数据集'
hdr[1].text = '类型'
hdr[2].text = '样本'
hdr[3].text = '细胞/spot 数'
hdr[4].text = '用途'
rows = [
    ['GSE174574', 'scRNA-seq', '3 Sham + 3 MCAO（24 h）', '58,287 → 56,486', '主分析'],
    ['GSE171169', 'scRNA-seq', 'CD45high 5d/14d', '10,295', '免疫细胞验证'],
    ['GSE225948', 'scRNA-seq', 'PB + brain Sham/D02/D14', '63,733', '髓系免疫细胞验证'],
    ['GSE233815', 'bulk RNA-seq', 'MCAO 3h/12h/24h/3D/7D', '48', '时间动态验证'],
    ['GSE233812', 'scRNA-seq', 'sham/D1/D3/D7', '6,159', '单细胞时间序列'],
    ['GSE233813', 'snRNA-seq', 'sham/D1/D3/D7', '8,374', '单细胞核时间序列'],
    ['GSE233814', 'Visium 空间', 'control/D1/D3/D7/D7_rep', '11,969 spots', '空间定位'],
]
for r in rows:
    row_cells = table.add_row().cells
    for i, v in enumerate(r):
        row_cells[i].text = v

# 主要结果
doc.add_heading('三、主要结果', level=1)

# 3.7 空间精确解析
doc.add_heading('3.7 GSE233814 Visium 空间转录组（精确坐标解析）', level=2)
doc.add_heading('解析方法', level=3)
for item in [
    "从 GEO 下载每个样本配套的 json.gz 空间文件，包含 fiducial（标记点）、oligo（spot 列表，含 row/col 与 imageX/imageY 像素坐标）以及 3×3 仿射变换矩阵 transform。",
    "利用 10x Genomics Visium v1 官方 barcode 到 (row, col) 的固定映射（visium-v1.txt + visium-v1_coordinates.txt，1-based 转 0-based），将每个样本 barcodes.tsv.gz 中的 spot barcode 映射到捕获阵列坐标。",
    "通过 (row, col) 匹配 oligo 条目，提取每个有组织 spot 在 H&E 图像上的像素坐标 imageX、imageY。",
    "将 Pirb 标准化表达（log1p CPM）叠加到空间像素坐标，绘制高分辨率空间分布热图，并进一步分析 Pirb+ spot 的空间聚类性与边界富集。",
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Pirb+ spot 比例', level=3)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light Grid Accent 1'
hdr = table2.rows[0].cells
hdr[0].text = '时间'
hdr[1].text = 'Pirb+ fraction'
for tp, frac in [('control', '0.62%'), ('D1', '2.40%'), ('D3', '5.93%'), ('D7', '2.01%'), ('D7_rep', '1.20%')]:
    r = table2.add_row().cells
    r[0].text = tp
    r[1].text = frac

doc.add_heading('空间分布新发现', level=3)
for item in [
    "D3 Pirb+ spots 高度聚集在组织外周/边缘区域，呈现明显的空间簇状分布。",
    "Pirb+ spots 的最近邻距离在 D3 最小（mean = 236 pixels），提示 D3 存在活跃的小胶质细胞/炎症病灶。",
    "缺血后（D1/D3/D7/D7_rep）Pirb+ spots 均显著靠近组织边界（Mann-Whitney U 检验 p < 0.01），而 control 无显著边界偏好。",
    "边界至中心五分位分析显示，D3 的 Pirb 阳性率从边缘（Q1，~13.3%）向中心（Q5，~1.3%）单调下降，进一步证实 Pirb+ 炎症灶定位于缺血半暗带/损伤边界。",
]:
    doc.add_paragraph(item, style='List Bullet')

# 插入关键图
doc.add_heading('关键空间图', level=3)
figs_to_insert = [
    ('pirb_spatial_combined_panel.png', 'GSE233814 Pirb 空间组合热图（5 个时间/重复）'),
    ('pirb_spatial_GSM7437223_D1-D3_D3.png', 'D3 样本 Pirb 空间分布详细图'),
    ('pirb_boundary_enrichment_ratio.png', 'Pirb+ spots 边界富集比（Pirb+ / Pirb− 平均边界距离）'),
    ('pirb_boundary_quintile_fraction.png', '边界至中心五分位的 Pirb 阳性率'),
]
for fname, caption in figs_to_insert:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True

# 跨数据集共识
doc.add_heading('四、跨数据集共识', level=1)
doc.add_paragraph("时间动态与细胞类型结论与 V9 一致，详见 Markdown 完整版。")

# 机制模型
doc.add_heading('五、机制模型', level=1)
for item in [
    "缺血后小胶质细胞激活，Pirb 从基线 → D1 上升 → D3 峰值 → D7 回落。",
    "激活的小胶质细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。",
    "Pirb+ 小胶质细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。",
    "少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。",
    "外周浸润髓系免疫细胞是另一 Pirb 高表达群体。",
]:
    doc.add_paragraph(item, style='List Number')

# 局限与下一步
doc.add_heading('七、局限与下一步', level=1)
doc.add_heading('下一步', level=2)
for item in [
    "路线二：跨时间小胶质细胞整合——整合 GSE174574（24h）+ GSE233812/813（D3），构建跨时间 Pirb+ 小胶质细胞状态转移图谱。",
    "路线三：体外功能实验——设计/启动 OGD、LPS、Pirb 阻断等实验，验证 Pirb 在小胶质细胞激活和神经炎症中的因果作用。",
    "修复 GSE225948 损坏文件或寻找替代数据集。",
]:
    doc.add_paragraph(item, style='List Number')

# 结论
doc.add_heading('八、结论', level=1)
doc.add_paragraph(
    "通过 7 个数据集的多组学交叉验证，我们一致发现：脑缺血后 Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值。"
    "GSE233814 的精确空间解析进一步揭示：Pirb+ spots 在 D3 高度聚集并显著富集于组织边缘/缺血半暗带，与小胶质细胞/炎症程序共定位，"
    "为 Pirb 作为脑缺血后神经炎症关键调控因子提供了更强有力的空间证据。"
)

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
