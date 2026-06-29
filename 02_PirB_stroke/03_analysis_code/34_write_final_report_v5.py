"""
生成最终阶段性报告 V5（加入 GSE225948 sub.celltype 结果）
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "../04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及独立数据集确认阶段性报告_V5.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及独立数据集确认阶段性报告（V5）

**报告日期**：2026-06-16  
**核心数据集**：GSE174574、GSE171169、GSE233815 bulk、GSE225948  
**分析环境**：Python + scanpy；R 不可用  
**网络状态**：NCBI FTP/HTTPS 极不稳定，GSE227651 与 GSE233812/813/814 大文件下载已暂停

---

## 一、目标与核心假设

### 1.1 目标
1. 明确脑缺血后 Pirb 阳性细胞的细胞类型分布与动态变化。
2. 揭示驱动 Pirb 上调的上游信号与下游功能。
3. 通过多数据集交叉验证构建稳健结论。

### 1.2 核心假设
- 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
- Pirb 放大 NF-κB/IRF/STAT 炎症程序
- 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
- Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
- 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集与方法

### 2.1 已分析数据集
| 数据集 | 样本 | 细胞数 | 用途 |
|--------|------|--------|------|
| GSE174574 | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | 4 CD45high 样本（5d / 14d） | 10,295 | 免疫细胞验证 |
| GSE233815 bulk | MCAO 3h/12h/24h/3D/7D + SH | 48 样本 | 时间动态验证 |
| GSE225948 | PB + brain（Sham/D02/D14） | 63,733 | 免疫细胞+内皮+小胶质独立验证 |

### 2.2 已暂停下载
| 数据集 | 原因 |
|--------|------|
| GSE227651 | NCBI FTP/HTTPS 多次失败，matrix 无法完整下载 |
| GSE233812/813/814 | RAW tar / 空间文件过大，网络速度 <50 KB/s 且频繁断线 |

---

## 三、主要结果

### 3.1 GSE174574 主分析（24 h MCAO）
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞 Pirb+ 率最高（31.4%），其次小胶质细胞（6.2%）、星形胶质细胞（3.1%）。
- Doublet 稳健性：Scrublet 总体 doublet 率 3.09%；排除 doublet 后 Astrocyte Pirb+ 从 3.13% 略降至 2.98%，结论稳健。
- Pirb+ 星形胶质细胞显著富集：溶酶体（p=6.2×10⁻¹⁶）、补体（p=4.9×10⁻⁹）、小胶质细胞诱导 A1（p=4.1×10⁻⁷）。
- 转录因子：Cebpb、Rel、Stat6、Irf8、Irf5 上调。
- 配体-受体：少突胶质细胞 MAG/MOG/OMgp/Nogo-A 和 MHC I → PirB 为主要预测通路。

### 3.2 GSE171169 验证（CD45high 免疫细胞）
- 免疫细胞 Pirb+ 率 27–28%，与 GSE174574 一致。

### 3.3 GSE233815 bulk 时间序列验证
- Pirb CPM：3h ~0.4，12h ~2.7，24h ~2.0（p=0.021），**3D ~6.7**，7D ~2.6。
- A1/炎症标志（Gfap、C3、C1qa、C1qb、Spp1、Lyz2）在 3D 同步达峰。

### 3.4 GSE225948 独立验证（缺血后 2d / 14d，CD45hi + MG + EC）
**分析样本**：63,733 个细胞（QC 后），外周血 36,272 细胞，脑梗死同侧半球 29,851 细胞。

#### 3.4.1 总体 Pirb 格局
| 组织 | 高表达细胞类型 | Pirb+ 率范围 |
|------|----------------|--------------|
| PB | Neu（中性粒细胞）| 29–55% |
| PB | Mo（单核细胞）| 25–52% |
| PB | DC | 23–36% |
| Brain | MdC（单核来源细胞/巨噬细胞）| 24–35% |
| Brain | Gran（粒细胞）| 21–46% |
| Brain | DC | 16–23% |
| Brain | Mg（小胶质细胞）| ~2.5–2.9% |
| Brain | EC（内皮细胞）| <0.3% |

#### 3.4.2 时间动态
- **外周血**：中性粒细胞和单核细胞 Pirb 在 Sham 已处高位，D02 维持/升高，D14 明显下降。
- **脑内浸润髓系细胞**：MdC 和 Gran 在 D02 高表达，D14 回落。
- **小胶质细胞**：Pirb 表达始终很低（<3%）。
- **内皮细胞**：Pirb 几乎不表达。

#### 3.4.3 Sub.celltype 水平高表达亚群
- **PB D02**：Mo2（62.7%）、Neu6（62.6%）、Neu5（57.6%）、Neu1（53.3%）、Neu4（52.9%）。
- **Brain D02**：MdC3（37.2%）、Gran1（37.0%）、DC3（31.8%）、MdC2（31.3%）、MdC4（29.7%）。
- **小胶质细胞亚群**：所有 Mg 亚群 Pirb+ 率均 <7%。

---

## 四、跨数据集共识与机制修正

### 4.1 共识
- **髓系免疫细胞 Pirb 高表达**：GSE174574（免疫细胞 31.4%）、GSE171169（27–28%）、GSE225948（PB Neu/Mo >50%）一致支持。
- **时间动态**：GSE233815 bulk（3D 峰值）与 GSE225948（D02 高、D14 降）共同提示 Pirb 在急性-亚急性期上调。
- **炎症关联**：Pirb 与 A1 星形胶质细胞反应、补体/溶酶体通路、NF-κB/IRF/STAT 程序共调控。

### 4.2 数据集特异性差异
| 观察 | GSE174574 | GSE225948 |
|------|-----------|-----------|
| 小胶质细胞 Pirb | 6.2% | ~2.5–2.9% |
| 星形胶质细胞 Pirb | 3.1% | 未捕获 |
| 内皮细胞 Pirb | 可见 | <0.3% |
| 主要 Pirb+ 细胞 | 免疫/小胶质/星形胶质 | 髓系免疫细胞 |
| 时间分辨率 | 24 h | 2d / 14d |

### 4.3 修正后的综合模型
1. 缺血后外周髓系免疫细胞（中性粒细胞、单核细胞）迅速浸润脑组织并高表达 Pirb。
2. 部分浸润巨噬细胞/小胶质细胞通过 IL-1α/TNF/C1q 诱导反应性星形胶质细胞 A1 转化，同时上调 Pirb。
3. Pirb 通过 NF-κB/IRF/STAT 程序放大炎症。
4. 少突胶质细胞来源的 MAG/OMgp/MOG 与 Pirb 相互作用，抑制轴突再生。
5. 内皮细胞 Pirb 表达在整体水平上较低，可能仅在特定亚群或更晚时间点发挥作用。

---

## 五、生成的关键图表

| 图表 | 说明 | 路径 |
|------|------|------|
| UMAP by parent/sub cell type | 细胞类型注释 | `04_reports/figures/GSE225948/umap_parent_celltype.png` |
| UMAP by tissue/time | 样本分布 | `04_reports/figures/GSE225948/umap_tissue_time.png` |
| Pirb expression violin | 各细胞类型/组织 Pirb 表达 | `04_reports/figures/GSE225948/pirb_expression_violin.png` |
| Pirb+ fraction heatmap | 组织×时间×细胞类型 | `04_reports/figures/GSE225948/pirb_fraction_heatmap.png` |
| Pirb+ fraction point plot | PB vs brain 时间动态 | `04_reports/figures/GSE225948/pirb_fraction_pointplot.png` |
| Key cell types time violin | Brain/PB 关键细胞类型动态 | `04_reports/figures/GSE225948/pirb_*_time.png` |
| Sub.celltype heatmap | 亚型水平 Pirb 表达 | `04_reports/figures/GSE225948/pirb_subcelltype_heatmap.png` |
| UMAP Pirb expression | 空间表达分布 | `04_reports/figures/GSE225948/umap_Pirb_expr.png` |

---

## 六、局限

1. GSE174574 仅 24 h 单时间点。
2. GSE225948 部分样本因 FTP 损坏缺失（13/26 counts 文件 corrupt），brain-Sham 样本量较小（n=1,009）。
3. GSE233812/813/814 尚未下载完成，缺少单细胞时间序列和空间定位。
4. GSE225948 数据为 normalized expression（非原始 UMI counts）。
5. 所有分析为计算预测，需体外功能实验验证。

---

## 七、下一步计划

1. 在网络稳定时继续尝试 GSE233812/813/814 下载。
2. 整合 GSE174574 + GSE225948 免疫细胞构建跨数据集图谱。
3. 体外实验：IL-1α/TNF/C1q 诱导 + Pirb 阻断。
4. 深入 GSE225948 亚型分析：年轻 vs 年老、PB→brain 细胞状态转换。

---

## 八、结论

多数据集交叉验证表明，脑缺血后 Pirb 主要富集于**髓系免疫细胞**（外周血中性粒细胞、单核细胞；脑内浸润巨噬细胞/粒细胞），并在急性期（2d–3d）达到峰值。GSE174574 进一步提示 Pirb+ 星形胶质细胞具有 A1-like 炎症表型，且与少突胶质细胞来源的抑制性配体存在互作。GSE225948 独立验证支持 Pirb 作为缺血后炎症反应的关键调控因子，并揭示其可能以外周浸润免疫细胞为主要载体进入损伤脑组织。后续需结合单细胞时间序列、空间转录组和功能实验进一步阐明 Pirb 的时空作用机制。
"""

with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
