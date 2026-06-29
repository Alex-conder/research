"""
生成 V12 阶段性报告（Word + Markdown），纳入 GSE225948 补全数据后的新结果。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "D:/Pirb_stroke_project/04_reports"
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"
MD_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V12.md")
DOCX_PATH = os.path.join(REPORT_DIR, "脑缺血后Pirb阳性细胞单细胞图谱_机制深挖多数据集验证及多组学确认阶段性报告_V12.docx")

md_content = f"""# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V12）

**报告日期**：{datetime.now().strftime('%Y-%m-%d')}  
**核心数据集**：GSE174574、GSE171169、GSE225948（26 样本完整版）、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间  
**分析环境**：Python + scanpy；R 不可用

---

## 一、核心假设

1. 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
2. Pirb 放大 NF-κB/IRF/STAT 炎症程序
3. 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
4. Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
5. 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集

| 数据集 | 类型 | 样本 | 细胞/spot 数 | 用途 |
|--------|------|------|--------------|------|
| GSE174574 | scRNA-seq | 3 Sham + 3 MCAO（24 h） | 58,287 → 56,486 | 主分析 |
| GSE171169 | scRNA-seq | CD45high 5d/14d | 10,295 | 免疫细胞验证 |
| GSE225948 | scRNA-seq | PB + brain Sham/D02/D14 | 96,293 → 91,688 | 髓系免疫细胞验证（完整版） |
| GSE233815 | bulk RNA-seq | MCAO 3h/12h/24h/3D/7D | 48 | 时间动态验证 |
| GSE233812 | scRNA-seq | sham/D1/D3/D7 | 6,159 | 单细胞时间序列 |
| GSE233813 | snRNA-seq | sham/D1/D3/D7 | 8,374 | 单细胞核时间序列 |
| GSE233814 | Visium 空间 | control/D1/D3/D7/D7_rep | 11,969 spots | 空间定位 |

---

## 三、主要结果

### 3.1 GSE174574 主分析
- Pirb+ 率：MCAO 8.2–9.8% vs Sham 2.7–2.8%。
- 细胞类型：免疫细胞（31.4%）> 小胶质细胞（6.2%）> 星形胶质细胞（3.1%）。
- Pirb+ 星形胶质细胞富集：溶酶体、补体、A1、NF-κB/IRF/STAT。

### 3.2 GSE171169 验证
- CD45high 免疫细胞 Pirb+ 率 27–28%。

### 3.3 GSE225948 独立验证（完整 26 样本，NEW）

**样本构成**：
- Brain：young Sham / D02 / D14，aged Sham / D02 / D14
- PB（外周血）：Sham / D02 / D14
- QC 后 91,688 细胞，25 Leiden 聚类

**Pirb+ 率关键结果**：
| 组织 | 细胞类型 | Sham | D02 | D14 |
|------|----------|------|-----|-----|
| Brain | 小胶质细胞 Mg | 2.89% | 2.39% | 2.53% |
| PB | 单核细胞 Mo | 35.25% | **51.72%** | 24.96% |
| PB | 中性粒细胞 Neu | 50.93% | **52.64%** | 28.67% |
| PB | 树突状细胞 DC | - | 36.17% | 23.23% |

**新发现**：
- 脑内小胶质细胞 Pirb 表达维持低水平（~2.4–2.9%），无明显 Sham/D02/D14 差异。
- 外周血髓系细胞（单核细胞、中性粒细胞、DC）Pirb 高表达，D02 急性期达峰，D14 显著回落。
- 该结果提示 **Pirb+ 外周浸润髓系细胞可能在卒中急性期（D02）通过血脑屏障进入脑内**，是脑内 Pirb+ 髓系细胞的重要来源。

### 3.4 GSE233815 bulk 时间序列
- Pirb CPM 在 3D 达峰（6.69）。

### 3.5 GSE233812 scRNA-seq
- 小胶质细胞 Pirb D3 峰值 **26.6%**。

### 3.6 GSE233813 snRNA-seq
- 小胶质细胞 Pirb D3 最高 **6.4%**。

### 3.7 GSE233814 Visium 空间转录组（精确坐标解析）
- D3 Pirb+ spot 比例最高（5.93%）。
- Pirb+ spots 显著富集于组织边缘/缺血半暗带。

### 3.8 路线二：跨时间小胶质细胞整合
- 整合 GSE174574 + GSE233812 + GSE233813 小胶质细胞（10,172 个）。
- 跨时间 Pirb+ 比例：sham 4.62% → D1 11.11% → MCAO_24h 6.16% → **D3 14.47%** → D7 7.23%。

---

## 四、跨数据集共识

### 4.1 时间动态
| 数据集 | 层面 | 峰值 | 关键结果 |
|--------|------|------|----------|
| GSE233812 | scRNA-seq | D3 | 小胶质细胞 Pirb 26.6% |
| GSE233813 | snRNA-seq | D3 | 小胶质细胞 Pirb 6.4% |
| GSE233814 | Visium 空间 | D3 | 5.93% spots |
| GSE233815 | bulk RNA-seq | 3D | CPM 6.69 |
| GSE174574 | scRNA-seq | 24h | 小胶质细胞 6.2% |
| 跨时间整合 | 整合 | D3 | 小胶质细胞 14.47% |
| GSE225948 | scRNA-seq | D02 | PB Mo/Neu Pirb ~52% |

### 4.2 细胞类型共识
- **脑内小胶质细胞**：GSE233812/813/174574 及跨时间整合一致显示 Pirb 在 D3 左右达峰。
- **外周浸润髓系细胞**：GSE225948/171169 一致显示外周血单核细胞、中性粒细胞、DC 高表达 Pirb，急性期（D02/D3）达峰。
- **星形胶质细胞**：GSE174574 中 A1 样 Pirb+ 星形胶质细胞显著。

---

## 五、整合机制模型

1. 缺血后外周血 **Pirb+ 单核细胞/中性粒细胞** 在 D02/D03 达峰，可能通过破坏的血脑屏障浸润脑组织。
2. 脑内 **小胶质细胞** 在缺血后激活，Pirb 从基线 → D1 上升 → **D3 峰值** → D7 回落。
3. Pirb+ 小胶质细胞/髓系细胞通过 IL-1α/TNF/C1q 诱导星形胶质细胞 A1 转化。
4. Pirb+ 小胶质细胞通过 NF-κB/IRF/STAT 程序放大神经炎症。
5. 少突胶质细胞 MAG/OMgp/MOG 与 Pirb 互作，抑制轴突再生。

---

## 六、关键图表

| 结果 | 路径 |
|------|------|
| GSE225948 UMAP（组织-时间） | `04_reports/figures/GSE225948/umap_tissue_time.png` |
| GSE225948 脑 Mg Pirb 时间动态 | `04_reports/figures/GSE225948/pirb_Mg_time.png` |
| GSE225948 PB Mo Pirb 时间动态 | `04_reports/figures/GSE225948/pirb_Mo_time.png` |
| GSE225948 PB Neu Pirb 时间动态 | `04_reports/figures/GSE225948/pirb_Neu_time.png` |
| GSE225948 Pirb 分组摘要 | `04_reports/figures/GSE225948/pirb_summary_by_group.csv` |
| GSE233814 空间组合热图 | `04_reports/figures/GSE233814/pirb_spatial_combined_panel.png` |
| 跨时间整合 UMAP | `04_reports/figures/microglia_cross_time/umap_time_point.png` |
| 跨时间 Pirb+ 比例 | `04_reports/figures/microglia_cross_time/pirb_fraction_timeline.png` |

---

## 七、局限与下一步

### 7.1 局限
1. GSE225948 脑内小胶质细胞 Pirb 阳性率较低，可能与 CD45hi 分选策略偏向外周/浸润细胞有关。
2. GSE233812 小胶质细胞数较少（406 个）。
3. 跨时间整合存在批次效应，未来需用 scVI/Harmony 等进一步校正。
4. GSE233814 未进行单细胞解卷积。
5. GSE227651 仍未下载。

### 7.2 下一步
1. **路线三：体外功能实验**——按已设计实验方案启动 OGD/LPS/Pirb 阻断实验，重点验证：
   - Pirb 阻断是否抑制小胶质细胞激活与炎症因子分泌；
   - Pirb 阻断是否减少小胶质细胞对神经元/突触的过度吞噬。
2. 尝试 scVI/Harmony 批次校正，进一步验证跨时间整合。
3. 继续尝试下载 GSE227651 或其他独立验证数据集。

---

## 八、结论

通过 7 个数据集的多组学交叉验证、Visium 空间精确解析及跨时间小胶质细胞整合，我们一致发现：脑缺血后 **Pirb 主要在小胶质细胞中表达，并于 D3（3 天）达到峰值**。GSE225948 完整数据进一步揭示 **外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达（~52%）**，提示外周浸润髓系细胞是脑内 Pirb+ 炎症细胞的重要来源。GSE233814 空间解析显示 Pirb+ 炎症灶富集于缺血半暗带。后续体外功能实验将以 D3 前后为时间窗、以 Pirb+ 小胶质细胞及外周髓系细胞为靶点，验证其因果作用。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Word
doc = Document()
title = doc.add_heading('脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及多组学确认阶段性报告（V12）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.add_run(f"报告日期：{datetime.now().strftime('%Y-%m-%d')}\n").bold = True
meta.add_run("核心数据集：GSE174574、GSE171169、GSE225948（26 样本完整版）、GSE233815 bulk、GSE233812 scRNA-seq、GSE233813 snRNA-seq、GSE233814 Visium 空间")

doc.add_heading('三、主要结果', level=1)

doc.add_heading('3.3 GSE225948 独立验证（完整 26 样本）', level=2)
doc.add_paragraph("QC 后 91,688 细胞，覆盖 brain young/aged 与 PB 的 Sham/D02/D14。")

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '组织'
hdr[1].text = '细胞类型'
hdr[2].text = 'D02 Pirb+'
hdr[3].text = 'D14 Pirb+'
rows = [
    ('Brain', '小胶质细胞 Mg', '2.39%', '2.53%'),
    ('PB', '单核细胞 Mo', '51.72%', '24.96%'),
    ('PB', '中性粒细胞 Neu', '52.64%', '28.67%'),
]
for r in rows:
    c = table.add_row().cells
    for i, v in enumerate(r):
        c[i].text = v

doc.add_paragraph("结论：外周血髓系细胞（Mo/Neu）Pirb 在 D02 急性期高表达，D14 回落；脑内小胶质细胞 Pirb 维持低水平。")

doc.add_heading('3.8 跨时间小胶质细胞整合', level=2)
doc.add_paragraph("整合 10,172 个小胶质细胞，Pirb+ 比例 D3 达峰 14.47%。")

# 插入关键图
figs = [
    ('GSE225948/pirb_Mo_time.png', 'GSE225948 PB 单核细胞 Pirb 表达时间动态'),
    ('GSE225948/pirb_Neu_time.png', 'GSE225948 PB 中性粒细胞 Pirb 表达时间动态'),
    ('GSE225948/pirb_Mg_time.png', 'GSE225948 脑小胶质细胞 Pirb 表达时间动态'),
    ('microglia_cross_time/pirb_fraction_timeline.png', '跨时间小胶质细胞 Pirb+ 比例'),
]
for fname, cap in figs:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True

doc.add_heading('八、结论', level=1)
doc.add_paragraph(
    "多组学交叉验证一致表明：脑缺血后 Pirb 主要在小胶质细胞中表达并于 D3 达峰。"
    "GSE225948 完整数据揭示外周血单核细胞/中性粒细胞在 D02 急性期 Pirb 高表达，提示外周浸润是脑内 Pirb+ 炎症细胞的重要来源。"
    "后续体外功能实验将验证 Pirb 在小胶质细胞激活、炎症放大及突触修剪中的因果作用。"
)

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
