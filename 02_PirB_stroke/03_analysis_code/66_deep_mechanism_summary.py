"""
生成五优先级深度机制分析结果汇总报告（MD + DOCX）。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "D:/Pirb_stroke_project/04_reports"
MD_PATH = os.path.join(OUT_DIR, "深度机制分析_5优先级结果汇总.md")
DOCX_PATH = os.path.join(OUT_DIR, "深度机制分析_5优先级结果汇总.docx")
FIG_DIR = "D:/Pirb_stroke_project/04_reports/figures"

md_content = f"""# 深度机制分析：Pirb 卒中后神经炎症的 5 优先级探索

**分析日期**：{datetime.now().strftime('%Y-%m-%d')}  
**分析目标**：解析 Pirb 在脑缺血后 D3 达峰、D7 回落的调控机制，明确半暗带炎症信号来源，探索外周-中枢通讯，并挖掘共表达隐藏模块。

---

## 第一优先级：D3 达峰与 D7 回落的转录调控刹车

### 研究问题
Pirb+ 小胶质细胞在 D3 高表达、D7 回落，是因为细胞死亡清除，还是存在主动转录抑制（刹车）？

### 方法
- 在 GSE174574（MCAO 24h）和 GSE233812（D3/D7）小胶质细胞中计算：
  - 凋亡模块（Casp3/7/8/9, Fas, Bax, Bcl2, Parp1）
  - 焦亡模块（Nlrp3, Casp1/4, Il1b, Il18, Gsdmd, Pycard）
  - 氧化应激模块（Hif1a, Nfe2l2, Sod1/2, Cat, Gpx1/4, Hmox1）
  - 抑制性 TF 模块（Socs1/3, Tsc22d3, Nfkbia, Tnfaip3, Zfp36, Dusp1, Irf2bp2, Bcl6, Klf2）
- 比较 Pirb+ vs Pirb− 小胶质细胞的模块评分和线粒体 reads 比例（percent.mt）。
- 比较 D7 vs D3 Pirb+ 小胶质细胞中抑制性 TF 和应激基因表达。

### 关键结果

| 数据集 | 比较 | 指标 | Pirb+ mean | Pirb− mean | p 值 | 结论 |
|--------|------|------|------------|------------|------|------|
| GSE233812 D3 | Pirb+ vs Pirb− | Oxidative_stress | 0.431 | 0.184 | **0.0004** | Pirb+ 细胞氧化应激显著升高 |
| GSE233812 D3 | Pirb+ vs Pirb− | Inhibitory_TFs | 0.115 | 0.012 | 0.086 | 抑制性 TF 略高，未达显著 |
| GSE233812 D3 | Pirb+ vs Pirb− | percent.mt | 1.96% | 2.82% | **0.044** | **Pirb+ 细胞 percent.mt 显著更低** |
| GSE233812 D3 | Pirb+ vs Pirb− | Apoptosis | −0.023 | −0.069 | 0.136 | 无显著差异 |
| GSE233812 D3 | Pirb+ vs Pirb− | Pyroptosis | −0.011 | −0.025 | 0.676 | 无显著差异 |
| GSE174574 24h | Pirb+ vs Pirb− | Immune score | 1.440 | 1.332 | **0.004** | Pirb+ 细胞免疫评分更高 |
| GSE174574 24h | Pirb+ vs Pirb− | percent.mt | 2.12% | 2.03% | 0.337 | 无显著差异 |

**D7 vs D3 Pirb+ 小胶质细胞（n_D3=33, n_D7=13）**：
- 显著下调的应激/抗氧化基因：Bax, Gpx4, Gpx1, Prdx1, Sod2（均 p < 0.05）。
- 抑制性 TF/负调控因子 **未观察到显著上调**（Nfkbia 下调，p=0.095；Socs1/Socs3/Tsc22d3 等无显著变化）。

### 结论
1. **不支持细胞耗竭假说**：D3 Pirb+ 小胶质细胞的 percent.mt 显著低于 Pirb− 细胞，凋亡/焦亡模块评分无显著升高，提示它们并非处于濒死状态。
2. **主动转录刹车证据不足**：D7 未观察到 Socs1/Socs3/Tsc22d3 等经典负调控因子的显著上调。
3. **更可能解释**：D7 Pirb 回落是炎症消退后 Pirb+ 激活亚群比例减少或状态转换，而非死亡清除或单一 TF 介导的主动抑制。但 D7 Pirb+ 细胞数极少（n=13），统计效力不足，需更大样本验证。

---

## 第二优先级：空间梯度差异表达（半暗带分子特征）

### 研究问题
Pirb+ spots 富集于缺血半暗带边界，半暗带内的 Pirb+ 细胞具体表达哪些炎症/基质降解信号？

### 方法
- GSE233814 D3 Visium 数据，基于 spot 间距将 spots 分区：
  - 核心（Core）：距边界 < 1 个 spot 间距
  - 半暗带（Penumbra）：距边界 1–3 个 spot 间距
  - 远端（Remote）：距边界 > 3 个 spot 间距
- 在半暗带内比较 Pirb+ vs Pirb− spots 的差异基因。
- 计算趋化因子、MMP、炎症、小胶质稳态模块评分。

### 关键结果
- **分区样本量**：Core=58, Penumbra=330, Remote=2160；Penumbra 中 Pirb+ spots=48。
- **Penumbra Pirb+ vs Pirb− 差异基因**：
  - 显著下调：Meg3, Cdk5r1, Vstm2l, Nsg2, Slc12a5, Thy1, Vgf, Wfs1, Calb1 等神经元相关基因。
  - 显著上调：Fth1（铁蛋白重链）。
- **模块基因在半暗带 Pirb+ spots 中的变化**：
  - Timp1 显著上调（log2FC=1.97, p=0.033）。
  - Mmp3 上调（log2FC=3.28, p=0.071），Mmp9 上调（log2FC=3.60, p=0.245）。
  - Ccl2, Ccl3, Ccl4, Ccl5 均上调，但未达校正显著性。
  - Il1b, Tnf, Il6 上调，但未达显著。
- **三个区域模块评分**：Penumbra 的 Chemokines、MMPs、Inflammation、Microglia_homeostasis 评分均高于 Core 和 Remote。

### 结论
半暗带不仅是 Pirb+ spots 的富集区，更是炎症信号发射源。Pirb+ spots 在半暗带内表现为神经元基因下调、Fth1/Timp1 上调、MMP/趋化因子表达升高的炎症/组织重塑表型。但单个 spot 内细胞混合导致差异基因统计效力有限，建议结合高分辨率空间方法（Xenium/MERFISH）验证。

---

## 第三优先级：外周-中枢配体-受体对接

### 研究问题
外周血单核细胞 D02 Pirb 高表达是否通过特定配体-受体轴诱导脑内 D3 小胶质细胞 Pirb 达峰？

### 方法
- 发送者：GSE225948 PB 单核细胞（Mo）D02（n=841）。
- 接收者：GSE233812 脑内小胶质细胞（Mg）D3（n=124）。
- 计算 curated 配体-受体对的对接分数：ligand_expr_sender × receptor_expr_receiver。

### 关键结果（Top LR pairs）

| 配体 | 受体 | 发送者表达 | 接收者表达 | LR 分数 | 配体阳性率 | 受体阳性率 |
|------|------|-----------|-----------|---------|-----------|-----------|
| App | Cd74 | 0.542 | 0.169 | **0.0914** | 39.4% | 8.9% |
| Mif | Cd74 | 0.086 | 0.169 | **0.0145** | 9.2% | 8.9% |
| Ccl5 | Ccr5 | 0.015 | 0.848 | **0.0127** | 1.9% | 55.6% |
| Csf1 | Csf1r | 0.004 | 2.493 | **0.0096** | 0.7% | 87.1% |
| Il1b | Il1rap | 0.147 | 0.046 | **0.0068** | 21.3% | 4.8% |
| Ltb | Ltbr | 0.032 | 0.152 | **0.0049** | 4.0% | 17.7% |
| Tnf | Tnfrsf1a | 0.010 | 0.434 | **0.0045** | 1.0% | 41.1% |
| Tnf | Tnfrsf1b | 0.010 | 0.335 | **0.0035** | 1.0% | 31.5% |
| Il1b | Il1r1 | 0.147 | 0.016 | **0.0023** | 21.3% | 1.6% |
| Ccl3 | Ccr1 | 0.008 | 0.261 | **0.0022** | 1.0% | 17.7% |
| Ccl2 | Ccr2 | 0.015 | 0.043 | **0.0006** | 1.0% | 3.2% |

### 结论
- 外周单核细胞可通过 **App/Mif→Cd74、Ccl5→Ccr5、Csf1→Csf1r、Il1b→Il1rap、Tnf→Tnfrsf1a/b** 等多条通路向脑内小胶质细胞传递信号。
- **Ccl2→Ccr2 轴在此数据集中信号较弱**（LR 分数 0.0006），可能因为 Ccl2 在 PB Mo D02 表达低或 Ccr2 在脑 Mg D3 表达有限。
- 结果支持"外周先行诱导中枢小胶质激活"假说，但需注意：这是跨数据集、跨时间的推断，未直接证明细胞迁移后的原位互作。

---

## 第四优先级：RNA 速率分析

### 研究问题
Pirb+ 小胶质细胞是否由原位转化而来，还是由不同细胞亚群更替？

### 方法
计划使用 scVelo / CellRank 基于剪接/未剪接 mRNA 推断 RNA 速率矢量。

### 结果
**无法完成**。GSE233812 处理后的 h5ad 仅包含 `layers['counts']`，缺乏 spliced/unspliced counts。scVelo 必须依赖 velocyto 或 alevin 等工具从原始 fastq/BAM 生成的剪接/未剪接计数矩阵。

### 建议
1. 下载 GSE233812 原始 fastq 文件，运行 velocyto 或 STARsolo + velocity 流程生成 loom/h5ad。
2. 或使用谱系示踪实验（如 Lgals3-CreER、Cx3cr1-CreER）直接标记小胶质细胞并追踪其命运。
3. 作为替代，可用 diffusion pseudotime 分析 Pirb 表达沿伪时间的动态（已在跨时间整合中实现）。

---

## 第五优先级：WGCNA 共表达模块分析

### 研究问题
跨时间整合的 10,172 个小胶质细胞中，是否存在与 Pirb 高度共表达的隐藏基因模块？

### 方法
- 使用 microglia_cross_time_integrated.h5ad（10,172 cells × 2001 HVG）。
- 计算基因-基因 Pearson 相关性矩阵。
- K-means 聚类（K=4，基于轮廓系数选择）。
- 计算模块特征基因与 Pirb 表达的相关性。

### 关键结果
- **4 个模块**：module 0（1673 基因，含 Pirb）、module 1（120）、module 2（73）、module 3（135）。
- **模块与 Pirb 相关性均较弱**：最高 r=0.018（module 3, p=0.066），Pirb 所在 module 0 与 Pirb r=−0.007（p=0.486）。
- **与 Pirb 最相关的基因**：Lyz2（r=0.066）、Ryr3（0.054）、Ftl1（0.049）、Jun（0.040）、B2m（0.039）、Il1r2（0.037）。
- Pirb 模块 top 基因涉及：溶酶体/吞噬（Lyz2）、铁代谢（Ftl1）、炎症转录（Jun, B2m）、IL-1 调控（Il1r2）。

### 结论
- 未发现与 Pirb 强共表达的独立隐藏模块，Pirb 与基础髓系/炎症/应激程序微弱共表达。
- 2001 个 HVG 可能限制了模块发现；建议基于全基因组或更大 HVG 集合重新分析。
- Pirb 本身表达稀疏且峰值短暂，难以形成强共表达模块。

---

## 综合机制模型（更新版）

基于五优先级分析，更新机制模型：

1. **外周启动**：D02 外周血单核细胞/中性粒细胞 Pirb 高表达，可能通过 **App/Mif→Cd74、Csf1→Csf1r、Il1b→Il1rap、Tnf→Tnfrsf1a/b** 等轴向脑内小胶质细胞传递激活信号。
2. **脑内 D3 峰值**：脑内小胶质细胞在 D3 达到 Pirb 表达峰值，伴随显著升高的氧化应激评分（Hif1a/Nfe2l2 等），但 percent.mt 更低，提示为**激活状态而非濒死**。
3. **半暗带炎症灶**：D3 Pirb+ spots 富集于缺血半暗带，局部表现为神经元基因下调、Fth1/Timp1/MMP/趋化因子上调，是炎症信号发射源。
4. **D7 回落机制**：目前数据不支持细胞耗竭或单一 TF 主动刹车。更可能是炎症消退后 Pirb+ 激活亚群比例减少或状态转换。需要更大样本、scVelo 或谱系示踪进一步验证。
5. **共表达网络**：Pirb 与 Lyz2/Ftl1/Jun/B2m/Il1r2 等髓系炎症基因微弱共表达，但未形成强隐藏模块。

---

## 下一步建议

1. **验证 D7 回落机制**：
   - 收集更大样本的 D7 小胶质细胞，专门分析 Pirb+ vs Pirb− 细胞。
   - 运行 scVelo（需原始 fastq）或谱系示踪实验。
   - 检测单细胞水平 SOCS1/SOCS3/TSC22D3 蛋白表达。

2. **半暗带高分辨率验证**：
   - 使用 Xenium/MERFISH 验证 Pirb+ 小胶质细胞在半暗带的空间分布及 MMP/Timp1 表达。
   - 进行 Pirb + CD68/MMP9/TIMP1 多重免疫荧光。

3. **外周-中枢通讯验证**：
   - 体外 transwell 实验：PB Mo D02 + 脑 Mg D3 共培养，阻断 Ccl5-Ccr5、Csf1-Csf1r、Il1b-Il1rap、Tnf-TNFR 轴。
   - 体内 Ccr2−/− 或 Csf1r 阻断实验。

4. **扩展 WGCNA**：
   - 基于全基因组或 top 5000 变异基因重新构建网络。
   - 纳入更多时间点和数据集。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# DOCX
doc = Document()
title = doc.add_heading('深度机制分析：Pirb 卒中后神经炎症的 5 优先级探索', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sections = [
    ("第一优先级：D3 达峰与 D7 回落的转录调控刹车", [
        "问题：Pirb+ 小胶质细胞 D3 高表达、D7 回落，是细胞死亡清除还是主动转录抑制？",
        "方法：计算凋亡/焦亡/氧化应激/抑制性 TF 模块评分，比较 Pirb+ vs Pirb−，追踪 D7 vs D3 抑制性 TF。",
        "关键结果：GSE233812 D3 Pirb+ 细胞氧化应激显著升高（p=0.0004），percent.mt 显著更低（p=0.044），凋亡/焦亡无显著差异。D7 未观察到 Socs1/Socs3/Tsc22d3 等抑制性 TF 显著上调。",
        "结论：不支持细胞耗竭；主动刹车证据不足；更可能是炎症消退后 Pirb+ 亚群减少或状态转换。",
    ]),
    ("第二优先级：空间梯度差异表达", [
        "问题：半暗带 Pirb+ 细胞的分子特征是什么？",
        "方法：GSE233814 D3 分 Core/Penumbra/Remote，半暗带内 Pirb+ vs Pirb− DE。",
        "关键结果：半暗带 Pirb+ spots 神经元基因下调（Meg3, Thy1, Vgf 等），Fth1 显著上调，Timp1 显著上调，Mmp3/Mmp9/Ccl2/Ccl3/Ccl4 上调但未达显著。",
        "结论：半暗带是炎症信号发射源，表现为神经元丢失、Fth1/Timp1/MMP/趋化因子升高。",
    ]),
    ("第三优先级：外周-中枢配体-受体对接", [
        "问题：外周单核细胞是否通过特定 LR 轴诱导脑内小胶质细胞 Pirb 达峰？",
        "方法：GSE225948 PB Mo D02 为发送者，GSE233812 脑 Mg D3 为接收者，计算 curated LR 对分数。",
        "关键结果：Top LR pairs 包括 App→Cd74、Mif→Cd74、Ccl5→Ccr5、Csf1→Csf1r、Il1b→Il1rap、Tnf→Tnfrsf1a/b。Ccl2→Ccr2 信号较弱。",
        "结论：支持外周先行诱导中枢小胶质激活假说，但需功能实验验证。",
    ]),
    ("第四优先级：RNA 速率分析", [
        "问题：Pirb+ 小胶质细胞是原位转化还是亚群更替？",
        "结果：无法完成。GSE233812 h5ad 仅含 counts 层，缺乏 spliced/unspliced counts，无法运行 scVelo。",
        "建议：从原始 fastq 运行 velocyto，或使用谱系示踪数据。",
    ]),
    ("第五优先级：WGCNA 共表达模块", [
        "问题：是否存在与 Pirb 高度共表达的隐藏基因模块？",
        "方法：10,172 个小胶质细胞 × 2001 HVG，K-means 共表达模块分析。",
        "关键结果：4 个模块，Pirb 位于最大模块（1673 基因），但模块与 Pirb 相关性弱（最高 r=0.018）。Top Pirb 相关基因：Lyz2, Ftl1, Jun, B2m, Il1r2。",
        "结论：Pirb 与基础髓系/炎症/应激程序微弱共表达，未发现强隐藏模块。",
    ]),
]

for sec_title, bullets in sections:
    doc.add_heading(sec_title, level=1)
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

# 插入关键图
for fname, cap in [
    ('regulatory_brake/boxplot_GSE233812_D3.png', '图 1 | GSE233812 D3 Pirb+ vs Pirb- 小胶质细胞应激/percent.mt 比较'),
    ('GSE233814/gradient_de/zone_module_scores_D3.png', '图 2 | GSE233814 D3 空间梯度模块评分'),
    ('peripheral_central_lr/top_lr_pairs.png', '图 3 | 外周-中枢 Top LR pairs'),
    ('wgcna_microglia/module_corr_with_Pirb.png', '图 4 | WGCNA 模块与 Pirb 相关性'),
]:
    fpath = os.path.join(FIG_DIR, fname)
    if os.path.exists(fpath):
        doc.add_picture(fpath, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(cap)
        p.runs[0].italic = True

doc.add_heading('综合结论', level=1)
doc.add_paragraph(
    'D3 Pirb+ 小胶质细胞为激活/氧化应激状态，非濒死细胞；半暗带是炎症信号发射源；'
    '外周单核细胞可能通过多条 LR 轴诱导脑内小胶质细胞激活；D7 回落更可能是炎症消退后的亚群减少或状态转换，'
    '而非细胞耗竭或单一 TF 主动刹车。'
)

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
