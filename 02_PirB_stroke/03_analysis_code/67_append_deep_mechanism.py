"""
把深度机制分析章节追加到最终报告和 NC 中英文报告的 MD/DOCX。
"""
import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORT_DIR = "../04_reports"
FIG_DIR = "../04_reports/figures"

DEEP_MD = os.path.join(REPORT_DIR, "深度机制分析_5优先级结果汇总.md")

def read_deep_section():
    with open(DEEP_MD, 'r', encoding='utf-8') as f:
        text = f.read()
    # 提取"## 第一优先级"到"## 综合机制模型"之后
    start = text.find("## 第一优先级")
    end = text.find("## 下一步建议")
    return text[start:end].strip()

def md_to_docx_append(docx_path, md_section, figures_to_insert=None):
    doc = Document(docx_path)
    doc.add_page_break()
    doc.add_heading('深度机制探索（5 项优先级任务）', level=1)
    # 简单解析 markdown 标题与段落
    lines = md_section.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('##'):
            # 按 # 数量确定级别
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            # 跳过总标题级别 1
            if level == 1:
                continue
            doc.add_heading(text, level=min(level, 3))
        elif line.startswith('|') or line.startswith('---'):
            continue
        elif line.startswith('-') or line.startswith('*'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
    # 插入关键图
    if figures_to_insert:
        for fname, cap in figures_to_insert:
            fpath = os.path.join(FIG_DIR, fname)
            if os.path.exists(fpath):
                try:
                    doc.add_picture(fpath, width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p = doc.add_paragraph(cap)
                    p.runs[0].italic = True
                except Exception as e:
                    print(f"[WARN] figure skip {fpath}: {e}")
    doc.save(docx_path)
    print(f"[UPDATE DOCX] {docx_path}")

def append_md_before_figures(md_path, md_section):
    with open(md_path, 'r', encoding='utf-8') as f:
        text = f.read()
    insert_pos = text.find("## 图注")
    if insert_pos == -1:
        insert_pos = text.find("## 利益冲突")
    if insert_pos == -1:
        insert_pos = len(text)
    new_text = text[:insert_pos] + "\n\n---\n\n" + md_section + "\n\n" + text[insert_pos:]
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(new_text)
    print(f"[UPDATE MD] {md_path}")

if __name__ == "__main__":
    section = read_deep_section()
    
    # 需要追加的报告
    targets = [
        ("脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.md",
         "脑缺血后Pirb阳性细胞单细胞图谱_多数据集验证最终报告_原格式.docx"),
        ("脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.md",
         "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告_中文版.docx"),
        ("脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.md",
         "脑缺血后Pirb阳性细胞单细胞图谱_NatureCommunications_完整报告.docx"),
    ]
    
    figures = [
        ('regulatory_brake/boxplot_GSE233812_D3.png', '图 S1 | GSE233812 D3 Pirb+ vs Pirb- 小胶质细胞应激/percent.mt 比较'),
        ('GSE233814/gradient_de/zone_module_scores_D3.png', '图 S2 | GSE233814 D3 空间梯度模块评分'),
        ('peripheral_central_lr/top_lr_pairs.png', '图 S3 | 外周-中枢 Top LR pairs'),
        ('wgcna_microglia/module_corr_with_Pirb.png', '图 S4 | WGCNA 模块与 Pirb 相关性'),
    ]
    
    for md_name, docx_name in targets:
        md_path = os.path.join(REPORT_DIR, md_name)
        docx_path = os.path.join(REPORT_DIR, docx_name)
        if not os.path.exists(md_path):
            print(f"[SKIP] {md_path} not found")
            continue
        if not os.path.exists(docx_path):
            print(f"[SKIP] {docx_path} not found")
            continue
        # 避免重复追加：检查是否已有"深度机制探索"
        with open(md_path, 'r', encoding='utf-8') as f:
            txt = f.read()
        if "深度机制探索" in txt:
            print(f"[SKIP] {md_name} already contains deep mechanism section")
            continue
        append_md_before_figures(md_path, section)
        md_to_docx_append(docx_path, section, figures)
    
    print("[DONE]")
