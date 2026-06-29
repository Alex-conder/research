"""
生成最终阶段性报告 V3（含横向扩展计划与 GSE233815 bulk marker 趋势）
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_MD = "D:/Pirb_stroke_project/04_reports/脑缺血后Pirb阳性细胞单细胞图谱_机制深挖与多数据集验证及横向扩展阶段性报告_V3.md"
OUT_DOCX = OUT_MD.replace(".md", ".docx")

report_md = """# 脑缺血后 Pirb 阳性细胞单细胞图谱：机制深挖、多数据集验证及横向扩展阶段性报告（V3）

**报告日期**：2026-06-16  
**分析数据集**：GSE174574、GSE171169、GSE233815 bulk  
**后台下载中**：GSE225948、GSE233812（scRNA-seq）、GSE233813（snRNA-seq）、GSE233814（Visium 空间）  
**已停止**：GSE227651（FTP 多次受阻）

---

## 一、目标与假设

### 1.1 核心目标
1. 明确脑缺血后 Pirb 阳性细胞的细胞类型分布与动态变化。
2. 揭示驱动 Pirb 上调的上游信号。
3. 构建 Pirb 参与缺血后炎症反应与轴突再生抑制的机制假设。

### 1.2 核心假设
- 小胶质细胞 IL-1α/TNF/C1q → A1 样 Pirb+ 星形胶质细胞
- Pirb 放大 NF-κB/IRF/STAT 炎症程序
- 少突胶质细胞 MAG/OMgp/MOG → Pirb 抑制轴突再生
- Pirb 参与免疫细胞/小胶质细胞免疫抑制或突触修剪
- 内皮细胞 Pirb 上调与血脑屏障破坏相关

---

## 二、数据集与方法

### 2.1 已分析数据集
| 数据集 | 样本 | 用途 |
|--------|------|------|
| GSE174574 | 3 Sham + 3 MCAO（24 h），58,287 细胞 | 主分析 |
| GSE171169 | 4 CD45high 样本（5d / 14d），10,295 细胞 | 免疫细胞验证 |
| GSE233815 bulk | MCAO 3h/12h/24h/3D/7D + SH，48 样本 | 时间动态验证 |

### 2.2 后台下载中
| 数据集 | 内容 | 预期用途 |
|--------|------|----------|
| GSE225948 | 29 样本 Brain + Blood CSV | 免疫/内皮互作 |
| GSE233812 | scRNA-seq（sham/D1/D3/D7） | 单细胞时间序列 |
| GSE233813 | snRNA-seq（sham/D1/D3/D7） | 单细胞核时间序列 |
| GSE233814 | Visium 空间转录组（control/D1/D3/D7） | 空间定位 |

### 2.3 已停止下载
| 数据集 | 原因 |
|--------|------|
| GSE227651 | NCBI FTP 连接极不稳定，多次尝试（urllib/curl/ftplib）均无法完成 3d/7d/sham matrix 下载 |

---

## 三、主要结果

### 3.1 Pirb 在 MCAO 后显著上调
- **GSE174574**：MCAO Pirb+ 率 8.2–9.8% vs Sham 2.7–2.8%。
- **GSE171169**：免疫细胞 Pirb+ 率 27–28%。

### 3.2 细胞类型格局与 Doublet 稳健性
- 免疫细胞 Pirb+ 率最高（31.4%），其次小胶质细胞（6.2%）、星形胶质细胞（3.1%）。
- Scrublet doublet 检测：总体 doublet 率 3.09%；Pirb+ 星形胶质细胞中 13.25% 为预测 doublet。
- 排除 doublet 后，Astrocyte MCAO Pirb+ 率从 3.13% 略降至 2.98%，结论稳健。

### 3.3 Pirb+ 星形胶质细胞 A1-like 炎症表型
排除 doublet 后仍显著富集：溶酶体（p=6.2×10⁻¹⁶）、补体（p=4.9×10⁻⁹）、小胶质细胞诱导 A1（p=4.1×10⁻⁷）。

### 3.4 上游转录因子
Cebpb、Rel、Stat6、Irf8、Irf5 在 Pirb+ 星形胶质细胞中上调。

### 3.5 配体-受体互作
- 少突胶质细胞是 MAG/MOG/OMgp/Nogo-A 主要来源。
- MHC I → PirB 信号最强。

### 3.6 GSE233815 bulk 时间序列验证关键发现
Pirb 及 A1/炎症相关基因在 MCAO 后呈时间依赖性变化：
- **Pirb**：12h 开始升高，24h 显著，**3D 达到峰值**，7D 回落。
- **Gfap / C3 / Vim / Serping1（A1 反应性）**：3D 达到峰值。
- **C1qa / C1qb（小胶质细胞诱导 A1）**：3D 达到峰值。
- **Spp1 / Lyz2 / Ccl2（炎症）**：持续升高，7D 仍高。
- **Abcb1a / Cxcl12（内皮 BBB）**：MCAO 中相对 SH 降低。

该趋势与"小胶质细胞 C1q → A1 星形胶质细胞 → Pirb 上调"假设高度一致。

---

## 四、横向扩展计划

由于 GSE227651 下载受阻，已转向 GSE233815 子系列进行横向扩展：

| 子系列 | 数据类型 | 价值 |
|--------|----------|------|
| GSE233812 | scRNA-seq（sham/D1/D3/D7） | 单细胞时间序列，替代 GSE227651 |
| GSE233813 | snRNA-seq（sham/D1/D3/D7） | 细胞核转录组，补充 scRNA-seq |
| GSE233814 | Visium 空间转录组 | 空间定位，替代空间实验 |

**整合策略**：
1. 对 GSE233812/813 进行单细胞质控、注释、Pirb 动态分析。
2. 与 GSE174574 整合（Harmony），扩大样本量并验证跨数据集结论。
3. 对 GSE233814 进行空间分析，验证 Pirb+ 细胞在梗死周围/半暗带的富集。
4. 结合 GSE225948（Brain + Blood）分析免疫细胞与内皮细胞的 Pirb 表达。

---

## 五、局限

1. GSE174574 为 24h 单时间点（已由 GSE233815 bulk 部分补全时间趋势）。
2. GSE227651 下载失败，缺少独立单细胞时间序列（拟由 GSE233812/813 补全）。
3. GSE233815 空间原始数据未下载（拟由 GSE233814 补全）。
4. 细胞类型注释基于 marker 得分，未用 SingleR/CellTypist。
5. 配体-受体分析为预测性，需功能验证。

---

## 六、下一步计划

1. 等待 GSE233812/813/814、GSE225948 后台下载完成。
2. 分析 GSE233812/813 单细胞时间序列，验证 Pirb 动态。
3. 分析 GSE233814 空间数据，定位 Pirb+ 细胞空间分布。
4. 整合 GSE174574 + GSE233812/813，构建跨数据集共识。
5. 体外功能实验：IL-1α/TNF/C1q 诱导 + Pirb 阻断。

---

## 七、结论

脑缺血后 Pirb 在免疫细胞、小胶质细胞和星形胶质细胞中显著上调。Pirb+ 星形胶质细胞具有 A1-like 炎症表型，上游可能由小胶质细胞来源的 IL-1α/TNF/C1q 诱导，并通过 NF-κB/IRF/STAT 程序驱动。GSE233815 bulk 时间序列显示 Pirb 与 A1/炎症标志在 3D 同步达到峰值，支持该机制的时效性。少突胶质细胞来源的 MAG/OMgp/MOG 可能通过 Pirb 抑制轴突再生。横向扩展至 GSE233812/813/814 将进一步提供单细胞时间序列和空间定位证据。
"""

# 保存 Markdown
with open(OUT_MD, 'w', encoding='utf-8') as f:
    f.write(report_md)
print(f"[MD] Saved: {OUT_MD}")

# 生成 DOCX
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def parse_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            cells = [c.strip() for c in line[1:-1].split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
    return rows

i = 0
lines = report_md.split('\n')
while i < len(lines):
    line = lines[i].strip()
    if not line:
        i += 1
        continue
    if line.startswith('# '):
        doc.add_heading(line[2:], level=0)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=1)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=2)
    elif line.startswith('---'):
        pass
    elif line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        table_data = parse_table(table_lines)
        if table_data:
            tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            tbl.style = 'Light Grid Accent 1'
            for r_idx, row in enumerate(table_data):
                for c_idx, cell_text in enumerate(row):
                    tbl.rows[r_idx].cells[c_idx].text = cell_text
        continue
    else:
        doc.add_paragraph(line)
    i += 1

doc.save(OUT_DOCX)
print(f"[DOCX] Saved: {OUT_DOCX}")
