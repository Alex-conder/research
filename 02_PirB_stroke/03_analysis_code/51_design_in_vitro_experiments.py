"""
路线三：体外功能实验设计
基于 Pirb 在脑缺血后小胶质细胞 D3 峰值的证据，设计 OGD/LPS/Pirb 阻断实验方案。
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "../04_reports"
MD_PATH = os.path.join(OUT_DIR, "Pirb体外功能实验设计_路线三.md")
DOCX_PATH = os.path.join(OUT_DIR, "Pirb体外功能实验设计_路线三.docx")

md_content = f"""# Pirb 体外功能实验设计方案（路线三）

**设计日期**：{datetime.now().strftime('%Y-%m-%d')}  
**实验目的**：验证 Pirb 在脑缺血后小胶质细胞激活及神经炎症放大中的因果作用  
**理论基础**：多数据集交叉验证显示 Pirb 主要在小胶质细胞中表达，并于 MCAO 后 D3 达到峰值；GSE233814 空间分析提示 Pirb+ 炎症灶定位于缺血半暗带/组织边缘。

---

## 一、实验模型选择

### 1.1 细胞模型
| 模型 | 优点 | 缺点 | 用途 |
|------|------|------|------|
| 原代小鼠小胶质细胞（C57BL/6，P0-P3） | 最接近生理状态 | 获取量有限、批次差异 | 主实验 |
| BV2 小鼠小胶质细胞系 | 易培养、可重复 | 与原代细胞表型有差异 | 机制初筛 |
| 小胶质细胞-神经元共培养 | 可观察突触修剪/神经毒性 | 复杂、变量多 | 功能验证 |

**推荐方案**：先用 BV2 进行条件优化和初筛，再用原代小胶质细胞验证关键结果。

### 1.2 缺血/炎症刺激
| 刺激 | 模拟病理 | 参数建议 |
|------|----------|----------|
| OGD（氧糖剥夺） | 缺血核心/半暗带 | 无糖 DMEM + 1% O₂ / 5% CO₂ / 94% N₂，2–6 h |
| LPS（100 ng/mL） | 系统性炎症/Toll 样受体激活 | 处理 6–24 h |
| OGD + LPS | 缺血后炎症放大 | OGD 4 h → 复氧复糖 + LPS 100 ng/mL 24 h |

---

## 二、Pirb 干预策略

### 2.1 干预手段
| 干预 | 作用机制 | 建议浓度/剂量 |
|------|----------|---------------|
| 抗-Pirb 中和抗体（anti-PIRB） | 阻断 Pirb 与其配体（MAG/OMgp/MOG）结合 | 5–20 μg/mL |
| 同型对照 IgG | 排除非特异性 Fc 效应 | 匹配浓度 |
| Pirb siRNA / shRNA | 敲低 Pirb 表达 | 50–100 nM siRNA，验证敲低效率 ≥70% |
| 乱序 siRNA | 阴性对照 | 等浓度 |
| Pirb-Fc 融合蛋白（竞争性抑制） | 竞争性结合配体 | 10–50 μg/mL |

### 2.2 给药时间窗
- **预防性**：刺激前 30 min 给药。
- **治疗性**：OGD 后 0 h、6 h、24 h 给药，模拟临床再灌注后干预。
- **基于 D3 峰值**：重点观察刺激后 48–72 h。

---

## 三、实验分组

### 3.1 主实验：OGD + anti-Pirb
| 组别 | 处理 | 目的 |
|------|------|------|
| Control | 正常培养基，常氧 | 基线 |
| OGD | OGD 4 h → 复氧 24/48/72 h | 建立缺血模型 |
| OGD + IgG | OGD + 同型对照抗体 | 排除抗体非特异性效应 |
| OGD + anti-Pirb | OGD + 抗-Pirb 抗体（预防性或治疗性） | 验证 Pirb 阻断效应 |
| anti-Pirb alone | 仅抗-Pirb，无 OGD | 评估药物本身影响 |

### 3.2 机制实验：Pirb 敲低
| 组别 | 处理 |
|------|------|
| Control siRNA + OGD | 阴性对照 |
| Pirb siRNA + OGD | 基因敲低验证 |
| Pirb siRNA + Control | 敲低本身影响 |

---

## 四、检测指标与时间点

### 4.1 Pirb 表达验证
| 方法 | 指标 | 时间点 |
|------|------|--------|
| qRT-PCR | Pirb mRNA | 0, 6, 24, 48, 72 h |
| Western blot | PIRB 蛋白 | 24, 48, 72 h |
| 流式细胞术 | 细胞表面 PIRB+ 比例 | 24, 48, 72 h |
| 免疫荧光 | PIRB / Iba1 双标 | 48, 72 h |

### 4.2 小胶质细胞激活表型
| 指标 | 方法 | 预期结果（Pirb 阻断后） |
|------|------|------------------------|
| Iba1 表达 | IF / WB | 降低 |
| CD68 表达 | 流式 / IF | 降低 |
| 细胞形态 | 免疫荧光（骨架染色） | 阿米巴样形态减少 |
| 增殖 | EdU / Ki67 | 降低 |
| 迁移 | Transwell / 划痕实验 | 降低 |

### 4.3 炎症因子分泌
| 因子 | 检测方法 | 预期结果（Pirb 阻断后） |
|------|----------|------------------------|
| IL-1β | ELISA / CBA | 降低 |
| TNF-α | ELISA / CBA | 降低 |
| IL-6 | ELISA / CBA | 降低 |
| C1q | ELISA / WB | 降低 |
| IL-10 | ELISA / CBA | 可能升高（抗炎表型） |
| NO | Griess 试剂 | 降低 |

### 4.4 信号通路
| 通路 | 指标 | 方法 |
|------|------|------|
| NF-κB | p-p65 / p65, IκBα 降解 | WB / IF |
| MAPK | p-ERK, p-p38, p-JNK | WB |
| STAT/IRF | p-STAT1, p-STAT3, IRF8 | WB |
| PI3K/Akt | p-Akt, mTOR | WB |

### 4.5 吞噬与突触修剪功能
| 功能 | 方法 | 预期结果 |
|------|------|----------|
| 吞噬能力 | 荧光微球 / 荧光标记突触小体 | 降低 |
| 突触修剪 | 小胶质细胞-神经元共培养 + PSD-95/synapsin 染色 | 减少突触丢失 |
| 神经毒性 | LDH 释放、神经元存活（MAP2/NeuN） | 神经元存活增加 |

---

## 五、实验时间安排

| 天数 | 任务 |
|------|------|
| Day -2 | 铺板 BV2 / 原代小胶质细胞 |
| Day -1 | 换液、准备 siRNA 转染（如需要） |
| Day 0 | OGD / LPS 刺激开始（t = 0） |
| Day 0.25 | 预防性给药（anti-Pirb / IgG） |
| Day 1 | 收取 24 h 样本（RNA / 蛋白 / 上清） |
| Day 2 | 收取 48 h 样本；免疫荧光固定 |
| Day 3 | 收取 72 h 样本；共培养实验终点 |
| Day 4-7 | qPCR、WB、ELISA、成像、统计分析 |

---

## 六、样本量与统计

- **生物学重复**：每组 n = 3–6（原代细胞建议 n ≥ 4，来自不同动物）。
- **技术重复**：每个生物学重复内 2–3 次技术重复。
- **统计方法**：
  - 多组比较：单因素或双因素 ANOVA + Tukey/Sidak 校正。
  - 两组比较：t 检验或 Mann-Whitney U 检验。
  - 时间序列：重复测量 ANOVA。
- **显著性标准**：*p* < 0.05 为显著，*p* < 0.01 为极显著。

---

## 七、预期结果

1. **Pirb 在 OGD/LPS 刺激后上调**，48–72 h 达峰，与原位 D3 峰值一致。
2. **Pirb 阻断/敲低显著抑制小胶质细胞激活**：Iba1/CD68 降低、阿米巴样形态减少。
3. **Pirb 阻断减少促炎因子 IL-1β/TNF-α/C1q 分泌**，可能促进抗炎表型。
4. **Pirb 阻断抑制 NF-κB / STAT / IRF 炎症信号通路**。
5. **Pirb 阻断减少小胶质细胞过度吞噬/突触修剪**，改善神经元存活。

---

## 八、关键对照与质控

1. **抗体特异性**：使用 PIRB KO 细胞或 siRNA 敲低验证抗体特异性。
2. **细胞活力**：CCK-8 / LDH 排除细胞死亡对结果的干扰。
3. **内毒素**：使用低内毒素抗体；LPS 实验设置未处理对照。
4. **批次控制**：原代细胞每次实验记录动物批次；BV2 细胞定期验证支原体。
5. **盲法分析**：图像定量由不知分组的研究者完成。

---

## 九、可交付成果

1. 实验方案 SOP（本方案）。
2. 原始数据记录表（Excel）。
3. 统计分析报告与图表。
4. 更新 V12 报告：纳入体外功能实验结果。

---

## 十、与前期分析的衔接

- **干预时间窗**：基于 D3 峰值，重点观察 48–72 h。
- **细胞靶点**：Pirb+ 小胶质细胞。
- **下游通路**：聚焦 NF-κB / IRF / STAT（GSE174574 中 Pirb+ 星形胶质细胞/免疫细胞富集通路）。
- **空间定位**：GSE233814 提示 Pirb+ 灶位于缺血半暗带，体外 OGD 模型可模拟半暗带低氧/低糖微环境。
"""

with open(MD_PATH, 'w', encoding='utf-8') as f:
    f.write(md_content)
print(f"[SAVE] {MD_PATH}")

# Word
doc = Document()
title = doc.add_heading('Pirb 体外功能实验设计方案（路线三）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.add_run(f"设计日期：{datetime.now().strftime('%Y-%m-%d')}\n").bold = True
meta.add_run("实验目的：验证 Pirb 在脑缺血后小胶质细胞激活及神经炎症放大中的因果作用\n")
meta.add_run("理论基础：Pirb 主要在小胶质细胞中表达，于 MCAO 后 D3 达到峰值")

doc.add_heading('一、实验模型选择', level=1)
doc.add_paragraph("推荐方案：先用 BV2 进行条件优化和初筛，再用原代小胶质细胞验证关键结果。")

doc.add_heading('二、Pirb 干预策略', level=1)
for item in [
    "抗-Pirb 中和抗体：5–20 μg/mL，阻断 Pirb 与 MAG/OMgp/MOG 配体结合。",
    "Pirb siRNA / shRNA：50–100 nM，验证敲低效率 ≥70%。",
    "Pirb-Fc 融合蛋白：10–50 μg/mL，竞争性抑制。",
    "给药时间窗：预防性（刺激前 30 min）或治疗性（OGD 后 0/6/24 h）。",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('三、主实验分组', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '组别'
hdr[1].text = '处理'
rows = [
    ('Control', '正常培养基，常氧'),
    ('OGD', 'OGD 4 h → 复氧 24/48/72 h'),
    ('OGD + IgG', 'OGD + 同型对照抗体'),
    ('OGD + anti-Pirb', 'OGD + 抗-Pirb 抗体'),
    ('anti-Pirb alone', '仅抗-Pirb，无 OGD'),
]
for r in rows:
    c = table.add_row().cells
    c[0].text = r[0]
    c[1].text = r[1]

doc.add_heading('四、核心检测指标', level=1)
for item in [
    "Pirb 表达：qRT-PCR、Western blot、流式、免疫荧光。",
    "激活表型：Iba1、CD68、细胞形态、增殖、迁移。",
    "炎症因子：IL-1β、TNF-α、IL-6、C1q、IL-10、NO。",
    "信号通路：NF-κB、MAPK、STAT/IRF、PI3K/Akt。",
    "功能实验：吞噬能力、突触修剪、神经毒性/存活。",
]:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('五、预期结果', level=1)
for item in [
    "Pirb 在 OGD/LPS 后上调，48–72 h 达峰。",
    "Pirb 阻断/敲低抑制小胶质细胞激活。",
    "Pirb 阻断减少促炎因子分泌。",
    "Pirb 阻断抑制 NF-κB / STAT / IRF 通路。",
    "Pirb 阻断减少过度吞噬/突触修剪，改善神经元存活。",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('六、与前期分析的衔接', level=1)
doc.add_paragraph(
    "基于单细胞/空间分析确定的 D3 峰值、小胶质细胞主要表达细胞类型及 NF-κB/IRF/STAT 通路，"
    "本实验以 48–72 h 为关键时间窗，以 Pirb+ 小胶质细胞为靶点，验证其在缺血后神经炎症中的因果作用。"
)

doc.save(DOCX_PATH)
print(f"[SAVE] {DOCX_PATH}")
print("[DONE]")
