#!/usr/bin/env python3
"""Generate a Word (.docx) report with embedded figures (MM-GBSA version)."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

REPORT_DIR = Path(__file__).parent
FIG_DIR = REPORT_DIR / 'figures'

doc = Document()

# Title
title = doc.add_heading('6UJN + Meropenem + Vancomycin 三元复合物对称 MD 验证报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Abstract
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '本研究对 6UJN 受体进行两种顺序的序贯对接（Vancomycin-first → Meropenem 与 Meropenem-first → Vancomycin），'
    '并对两种三元复合物均执行相同的 MD 约束松弛（100 ps NPT），以消除“松弛结构 vs 未松弛结构”的方法学不对称。'
    '随后对松弛轨迹进行 MM-GBSA 计算，获得可比较的结合自由能（ΔG）。'
)

# Comparison table in abstract
abs_table = doc.add_table(rows=3, cols=8)
abs_table.style = 'Light Grid Accent 1'
abs_headers = ['顺序', '第二步配体', '对接结合能 E (kcal/mol)', 'apo 参考能 (kcal/mol)', '对接后 COM (Å)', '松弛后 COM (Å)', '松弛后最小距离 (Å)', 'MM-GBSA ΔG (kcal/mol)']
for j, h in enumerate(abs_headers):
    abs_table.rows[0].cells[j].text = h
abs_table.rows[1].cells[0].text = 'Vancomycin-first → Meropenem'
abs_table.rows[1].cells[1].text = 'Meropenem'
abs_table.rows[1].cells[2].text = '−6.738'
abs_table.rows[1].cells[3].text = '−6.788'
abs_table.rows[1].cells[4].text = '51.05'
abs_table.rows[1].cells[5].text = '46.76'
abs_table.rows[1].cells[6].text = '31.25'
abs_table.rows[1].cells[7].text = '−23.03 ± 2.48'
abs_table.rows[2].cells[0].text = 'Meropenem-first → Vancomycin'
abs_table.rows[2].cells[1].text = 'Vancomycin'
abs_table.rows[2].cells[2].text = '−8.502'
abs_table.rows[2].cells[3].text = '−8.743'
abs_table.rows[2].cells[4].text = '7.38'
abs_table.rows[2].cells[5].text = '7.55'
abs_table.rows[2].cells[6].text = '3.14'
abs_table.rows[2].cells[7].text = '−52.85 ± 4.20'

doc.add_paragraph(
    '核心结论：对称 MD + MM-GBSA 表明，Meropenem-first 能形成拥挤且能量显著更有利的三元复合物（对接 COM 7.38 Å，ΔG = −52.85 ± 4.20 kcal/mol），'
    '而 Vancomycin-first 中 Meropenem 远离 Vancomycin（对接 COM 51.05 Å，ΔG = −23.03 ± 2.48 kcal/mol）。'
    '交叉对比还显示，Vancomycin-first 中 Meropenem 的最佳对接亲和力（−6.738 kcal/mol）与其 apo 参考（−6.788 kcal/mol）几乎相同，'
    '说明 Vancomycin 预结合未促进 Meropenem 的结合。'
    '这提示 Meropenem 的预结合可能通过局部口袋预组织促进 Vancomycin 的后续结合，而非简单的能量中性共存。'
)

# Background
doc.add_heading('1. 研究背景与方法学修正', level=1)
doc.add_heading('1.1 原始方法学缺陷', level=2)
doc.add_paragraph(
    '在先前的报告中，Meropenem-first 结构经过 MD 松弛，而 Vancomycin-first 结构仅使用对接能量 −6.738 kcal/mol，未做松弛。'
    '这相当于用“精修后的照片”对比“原始底片”，结论在方法论上不严谨。'
)
doc.add_heading('1.2 对称验证策略', level=2)
strategy_items = [
    '取 Vancomycin-first 三元复合物 PDB。',
    '使用与 Meropenem-first 完全相同的脚本和参数：动态口袋提取、封顶、水盒、100 ps 约束升温。',
    '比较松弛后的配体–配体距离。',
    '对两个松弛体系运行 MM-GBSA（igb=2, saltcon=0.15 M），量化第二步配体与（蛋白 + 第一步配体）的结合自由能。'
]
for item in strategy_items:
    doc.add_paragraph(item, style='List Bullet')

# Methods
doc.add_heading('2. 材料与方法', level=1)
doc.add_heading('2.1 软件与环境', level=2)
methods_table = doc.add_table(rows=8, cols=2)
methods_table.style = 'Light Grid Accent 1'
data = [
    ('AutoDock Vina', '1.2.5'),
    ('Open Babel', '3.1.1'),
    ('AmberTools', 'conda 环境 md_env'),
    ('OpenMM', '8.5.2'),
    ('mdtraj / matplotlib', '轨迹分析与作图'),
    ('MMPBSA.py', 'AmberTools 14.0'),
    ('Python', '3.13'),
    ('运行平台', 'WSL Ubuntu + Windows Git Bash'),
]
for i, (k, v) in enumerate(data):
    methods_table.rows[i].cells[0].text = k
    methods_table.rows[i].cells[1].text = v

doc.add_heading('2.2 MD 协议（两种顺序完全一致）', level=2)
md_items = [
    '口袋提取：配体重原子周围 10 Å',
    '力场：ff14SB / GAFF2 / TIP3P',
    '非键方法：PME，1.0 nm 截断',
    '最小化：5 阶段约束最小化（k = 5000 → 0），每阶段 5,000 步',
    '平衡：100 ps NPT，蛋白骨架 CA/C/N/O 约束 100 kJ/mol/nm²，配体自由',
    '温度/压力：300 K / 1 bar，2 fs 步长'
]
for item in md_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 MM-GBSA 协议', level=2)
mmgbsa_items = [
    '方法：单轨迹 MM-GBSA（MMPBSA.py）',
    '溶剂模型：GB^OBC II（igb=2）',
    '盐浓度：0.15 M',
    '分析帧：100 ps 轨迹全部 100 帧（interval=1）',
    '受体定义：蛋白 + 第一步配体',
    '配体定义：第二步配体',
    '拓扑处理：ante-MMPBSA.py 去除水/离子；cpptraj parmstrip 拆分受体/配体拓扑',
]
for item in mmgbsa_items:
    doc.add_paragraph(item, style='List Bullet')

# Results
doc.add_heading('3. 结果', level=1)

doc.add_heading('3.1 对接能量（未松弛）', level=2)
dock_table = doc.add_table(rows=3, cols=5)
dock_table.style = 'Light Grid Accent 1'
dock_data = [
    ('顺序', '第二步配体', '对接结合能', 'apo 参考能', 'ΔE'),
    ('Vancomycin-first → Meropenem', 'Meropenem', '−6.738', '−6.788', '+0.050'),
    ('Meropenem-first → Vancomycin', 'Vancomycin', '−8.502', '−8.743', '+0.241'),
]
for i, row in enumerate(dock_data):
    for j, val in enumerate(row):
        dock_table.rows[i].cells[j].text = val

doc.add_paragraph('未松弛时，两种顺序的 ΔE 都接近 0，似乎均无显著协同/竞争。单独的对接能量无法反映几何差异，需要结合图 1 的交叉对比分析。')

doc.add_heading('3.2 交叉对比：对接能量与三元几何', level=2)
doc.add_paragraph(
    '为对齐两种对接顺序的结论，补充了独立的交叉对比分析（图 1）。该图同时展示：'
    '（左）四种对接能量的并列比较——Meropenem apo（−6.788）、Vancomycin-first 中的 Meropenem（−6.738）、'
    'Vancomycin apo（−8.743）、Meropenem-first 中的 Vancomycin（−8.502）；'
    '（右）对接所得三元复合物的几何差异——Meropenem-first 中 VAN-MER 质心距仅 7.38 Å，而 Vancomycin-first 中高达 51.05 Å。'
)
doc.add_picture(str(FIG_DIR / 'cross_comparison_summary.png'), width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 1. 两种对接顺序的交叉对比：对接亲和力（左）与对接三元几何（右）。Mero-first 同时具有强亲和力与近距离几何；Van-first 亲和力接近 apo 且两配体完全分离。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    '该图与 MD/MM-GBSA 结论高度一致：Mero-first 形成拥挤且能量显著有利的三元复合物；'
    'Van-first 中第二步配体实际上未进入三元相互作用区域。'
)

doc.add_heading('3.3 MD 松弛后的关键距离', level=2)
dist_table = doc.add_table(rows=3, cols=5)
dist_table.style = 'Light Grid Accent 1'
dist_data = [
    ('顺序', '对接后 COM (Å)', '100 ps 后 COM (Å)', '对接后最小距离 (Å)', '100 ps 后最小距离 (Å)'),
    ('Meropenem-first → Vancomycin', '7.38', '7.55', '2.84', '3.14'),
    ('Vancomycin-first → Meropenem', '51.05', '46.76', '34.85', '31.25'),
]
for i, row in enumerate(dist_data):
    for j, val in enumerate(row):
        dist_table.rows[i].cells[j].text = val

# Figures
doc.add_heading('3.4 Meropenem-first 距离轨迹', level=2)
doc.add_picture(str(FIG_DIR / 'distances_100ps.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 2. Meropenem-first：100 ps NPT 中配体–配体距离。COM 稳定在 ~7.5 Å，最小距离多数 > 3.0 Å，但数次短暂跌破 3.0 Å。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.5 Vancomycin-first 距离轨迹', level=2)
doc.add_picture(str(FIG_DIR / 'distances_van_first_100ps.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 3. Vancomycin-first：100 ps NPT 中配体–配体距离。COM 始终 ~47 Å，最小距离始终 > 30 Å，说明 Meropenem 未进入 Vancomycin 位点。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('3.6 两种顺序的直接对比', level=2)
doc.add_picture(str(FIG_DIR / 'comparison_distances_100ps.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 4. 两种顺序松弛后的配体距离对比。Meropenem-first 保持近距离，Vancomycin-first 则完全分离。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# MM-GBSA results
doc.add_heading('3.7 MM-GBSA 结合自由能', level=2)
mmgbsa_table = doc.add_table(rows=3, cols=5)
mmgbsa_table.style = 'Light Grid Accent 1'
mmgbsa_data = [
    ('顺序', '第二步配体', '受体定义', 'ΔG (kcal/mol)', '标准误'),
    ('Meropenem-first → Vancomycin', 'Vancomycin', 'protein + Meropenem', '−52.85 ± 4.20', '0.42'),
    ('Vancomycin-first → Meropenem', 'Meropenem', 'protein + Vancomycin', '−23.03 ± 2.48', '0.25'),
]
for i, row in enumerate(mmgbsa_data):
    for j, val in enumerate(row):
        mmgbsa_table.rows[i].cells[j].text = val

doc.add_paragraph('ΔΔG（Meropenem-first − Vancomycin-first）≈ −29.8 kcal/mol，表明 Meropenem 先占据时，第二步 Vancomycin 的结合自由能显著更有利。')

doc.add_heading('3.8 MM-GBSA 能量分解（Meropenem-first）', level=2)
decomp1_table = doc.add_table(rows=8, cols=3)
decomp1_table.style = 'Light Grid Accent 1'
decomp1_data = [
    ('能量项', '平均值 (kcal/mol)', '标准误'),
    ('ΔVDWAALS', '−84.98 ± 5.17', '0.52'),
    ('ΔEEL', '−30.10 ± 6.40', '0.64'),
    ('ΔEGB', '+71.98 ± 6.88', '0.69'),
    ('ΔESURF', '−9.74 ± 0.43', '0.04'),
    ('ΔG gas', '−115.08 ± 8.84', '0.88'),
    ('ΔG solv', '+62.24 ± 6.74', '0.67'),
    ('ΔTOTAL', '−52.85 ± 4.20', '0.42'),
]
for i, row in enumerate(decomp1_data):
    for j, val in enumerate(row):
        decomp1_table.rows[i].cells[j].text = val

doc.add_paragraph('范德华贡献主导了有利结合，极性溶剂化部分抵消了静电吸引。')

doc.add_heading('3.9 MM-GBSA 能量分解（Vancomycin-first）', level=2)
decomp2_table = doc.add_table(rows=8, cols=3)
decomp2_table.style = 'Light Grid Accent 1'
decomp2_data = [
    ('能量项', '平均值 (kcal/mol)', '标准误'),
    ('ΔVDWAALS', '−29.47 ± 2.79', '0.28'),
    ('ΔEEL', '+343.87 ± 14.76', '1.48'),
    ('ΔEGB', '−334.12 ± 12.10', '1.21'),
    ('ΔESURF', '−3.30 ± 0.27', '0.03'),
    ('ΔG gas', '+314.40 ± 13.17', '1.32'),
    ('ΔG solv', '−337.42 ± 12.27', '1.23'),
    ('ΔTOTAL', '−23.03 ± 2.48', '0.25'),
]
for i, row in enumerate(decomp2_data):
    for j, val in enumerate(row):
        decomp2_table.rows[i].cells[j].text = val

doc.add_paragraph('与 Meropenem-first 相比，Vancomycin-first 的范德华吸引明显较弱；气相静电呈现强烈斥力（+343.87 kcal/mol），被 GB 溶剂化项（−334.12 kcal/mol）近乎完全抵消。这种“大数相消”导致净 ΔG 虽仍为负，但数值上明显弱于 Meropenem-first，且对 GB 参数/采样较为敏感。')

doc.add_heading('3.10 体系稳定性', level=2)
doc.add_paragraph('两种 MD 均稳定收敛：温度 ~300 K，密度 ~1.0 g/mL，势能平稳，无 NaN。')

# Discussion
doc.add_heading('4. 讨论', level=1)

doc.add_heading('4.1 方法学修正后的新结论', level=2)
doc.add_paragraph(
    '原始“无显著 paving/竞争”结论建立在不公平的对比上。对称 MD + MM-GBSA 揭示：'
    'Meropenem-first 能形成并保持拥挤的三元复合物，对接 COM 仅 7.38 Å，MM-GBSA ΔG = −52.85 ± 4.20 kcal/mol；'
    '而 Vancomycin-first 中 Meropenem 远离 Vancomycin（对接 COM 51.05 Å），MM-GBSA ΔG = −23.03 ± 2.48 kcal/mol。'
    '交叉对比显示 Vancomycin-first 中 Meropenem 的亲和力（−6.738 kcal/mol）与其 apo 状态（−6.788 kcal/mol）几乎相同，'
    '进一步证明 Vancomycin 预结合并未促进 Meropenem 的结合。'
)

doc.add_heading('4.2 为什么两种顺序结果不同？', level=2)
discussion_points = [
    '刚性受体近似：Meropenem-first 对接中，Vancomycin 被强制在含 Meropenem 的固定受体上搜索，Grid 中心又定在 Vancomycin 原生位点，因此必须靠近 Meropenem。',
    'Vancomycin-first 更真实：当 Vancomycin 先占据原生位点后，Meropenem 在自由搜索中发现更优位点远离 Vancomycin；其最佳亲和力与 apo 参考几乎相同。',
    '方向依赖性：paving/预组织效应具有方向性。Meropenem 先结合可能诱导口袋局部变化，有利于 Vancomycin 随后结合；反之则不存在这种诱导。',
    '能量学与几何一致：MM-GBSA ΔΔG ≈ −29.8 kcal/mol 与交叉对比图（近距离 vs 远距离）共同支持“Meropenem-first 更有利”，排除了仅由对接能量得出的“无差异”结论。'
]
for point in discussion_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('4.3 对原始 ΔΔG 的重新评估', level=2)
doc.add_paragraph(
    '未松弛 ΔΔG = −1.764 kcal/mol 不能直接比较，因为 Vancomycin-first 的 Meropenem 未在 Vancomycin 位点。'
    'MM-GBSA 给出的 ΔΔG ≈ −29.8 kcal/mol 定量表明 Meropenem-first 顺序显著更有利。'
)

doc.add_heading('4.4 局限与未来工作', level=2)
limitations = [
    '模拟时长仅 100 ps，无法观察长期稳定性或解离动力学。',
    '两种顺序口袋不同（124 vs 131 残基），但均包含各自配体。',
    'MM-GBSA 使用 igb=2 和 100 帧采样，仅为快速估算；更长轨迹、igb=5/8 或 PB 可进一步提高精度。',
    '完整蛋白 MD（需 GPU/HPC）可进一步验证长程效应。'
]
for lim in limitations:
    doc.add_paragraph(lim, style='List Bullet')

# Conclusion
doc.add_heading('5. 结论', level=1)
conclusion = doc.add_paragraph()
conclusion_run = conclusion.add_run(
    '对称 MD 松弛 + MM-GBSA 能量学量化后，Meropenem-first 与 Vancomycin-first 给出定性且定量不同的结果。'
    'Meropenem-first 形成可松弛的拥挤三元复合物（对接 COM 7.38 Å，松弛后 7.55 Å；最小距离 ~3.14 Å），MM-GBSA ΔG = −52.85 ± 4.20 kcal/mol；'
    'Vancomycin-first 中 Meropenem 远离 Vancomycin（对接 COM 51.05 Å，松弛后 46.76 Å；最小距离 31.25 Å），MM-GBSA ΔG = −23.03 ± 2.48 kcal/mol。'
    'ΔΔG ≈ −29.8 kcal/mol 定量支持 Meropenem-first 顺序的能量优势。'
    '这一方向依赖性提示 Meropenem 的预结合可能通过局部口袋预组织促进 Vancomycin 的后续结合，而非简单的能量中性共存。'
    '原始基于未松弛能量的“无 paving”结论需要修正。'
)
conclusion_run.bold = True

# Appendix A: Vancomycin-first docking details
doc.add_heading('6. 附录 A：Vancomycin-first → Meropenem 分子对接详细结果', level=1)

doc.add_heading('A.1 实验目的', level=2)
doc.add_paragraph(
    '在 6UJN 受体上先对接 Vancomycin，再在其形成的 holo 受体上对接 Meropenem，'
    '以获取 Vancomycin-first 序贯对接的详细结果与可视化。该部分作为“交叉 first”对接的原始记录，'
    '与正文 MM-GBSA/MD 结论相互印证。'
)

doc.add_heading('A.2 对接参数', level=2)
params_table = doc.add_table(rows=7, cols=2)
params_table.style = 'Light Grid Accent 1'
params_data = [
    ('对接配体 poses', 'input/Meropenem_out.pdbqt'),
    ('对接结果 CSV', 'input/6UJN_Vancomycin_Meropenem.csv'),
    ('对接中心 (Å)', '(56.0328, 17.1666, 25.8216)'),
    ('盒子大小 (Å)', '60.68 × 60.68 × 60.68'),
    ('Exhaustiveness', '8'),
    ('num_modes', '10'),
    ('seed', '42'),
]
for i, (k, v) in enumerate(params_data):
    params_table.rows[i].cells[0].text = k
    params_table.rows[i].cells[1].text = v

doc.add_heading('A.3 对接结果', level=2)
van_dock_table = doc.add_table(rows=10, cols=6)
van_dock_table.style = 'Light Grid Accent 1'
van_dock_data = [
    ('Mode', 'Binding Affinity', 'RMSD/lb', 'RMSD/ub', '到 VAN 质心距离 (Å)', 'vs apo MER RMSD (Å)'),
    ('1', '−6.738', '0.000', '0.000', '51.05', '52.809'),
    ('2', '−6.713', '30.680', '33.330', '51.71', '40.233'),
    ('3', '−6.685', '36.790', '40.410', '53.70', '42.406'),
    ('4', '−6.630', '3.079', '4.411', '18.39', '4.720'),
    ('5', '−6.459', '4.192', '6.486', '16.71', '7.188'),
    ('6', '−6.459', '36.480', '40.060', '44.02', '46.083'),
    ('7', '−6.418', '29.430', '32.080', '51.16', '40.228'),
    ('8', '−6.415', '3.711', '8.083', '26.83', '16.456'),
    ('9', '−6.376', '50.870', '55.590', '53.24', '41.737'),
]
for i, row in enumerate(van_dock_data):
    for j, val in enumerate(row):
        van_dock_table.rows[i].cells[j].text = val

doc.add_paragraph('最佳亲和力：−6.738 kcal/mol（Mode 1）。')

doc.add_heading('A.4 关键观察', level=2)
observations = [
    'Meropenem 远离 Vancomycin：最佳 pose（Mode 1）的质心距 Vancomycin 质心约 51.05 Å。',
    '多模式分散：各 pose 距 Vancomycin 的范围约 16.7–53.7 Å，说明 Meropenem 未聚集于 Vancomycin 位点。',
    '与 apo Meropenem 差异大：最佳 pose 与 apo Meropenem 的重原子 RMSD 为 52.81 Å，结合模式显著不同。',
    '能量意义：最佳亲和力 −6.738 kcal/mol 与 apo Meropenem 参考能 −6.788 kcal/mol 接近（ΔE ≈ +0.050 kcal/mol），表明 Vancomycin 预结合几乎未提升 Meropenem 在其位点的结合。'
]
for obs in observations:
    doc.add_paragraph(obs, style='List Number')

doc.add_heading('A.5 可视化', level=2)
doc.add_picture(str(FIG_DIR / 'van_first_complex_overview.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 A1. 蛋白 CA（灰）、Vancomycin（蓝）与最佳 Meropenem pose（红）。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(str(FIG_DIR / 'van_first_docking_poses.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 A2. 9 个 Meropenem pose 按亲和力着色；Vancomycin（蓝）。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(str(FIG_DIR / 'van_first_distance_vs_affinity.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 A3. 各 pose 的 Meropenem 质心到 Vancomycin 质心距离与其对接亲和力。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_picture(str(FIG_DIR / 'van_first_xy_projection.png'), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph('图 A4. XY 平面投影显示 Meropenem pose 分布于蛋白表面，远离 Vancomycin。')
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('A.6 小结', level=2)
doc.add_paragraph(
    'Vancomycin-first 序贯对接中，Meropenem 未占据 Vancomycin 附近的口袋，而是结合在远离 Vancomycin 的位点。'
    '最佳亲和力 −6.738 kcal/mol 与 apo 状态几乎相同，说明 Vancomycin 预结合对 Meropenem 结合无显著促进作用。'
    '这与正文对称 MD 松弛结果（COM ~46.76 Å，MM-GBSA ΔG = −23.03 kcal/mol）一致。'
)

# File list
doc.add_heading('7. 生成文件清单', level=1)
files = [
    ('output/6UJN_Meropenem_Vancomycin.csv', 'Meropenem-first 对接结果'),
    ('input/6UJN_Vancomycin_Meropenem.csv', 'Vancomycin-first 对接结果'),
    ('md/input_van_first/protein.pdb', 'Vancomycin-first 清理后蛋白'),
    ('md/output/complex_short.prmtop / .inpcrd', 'Meropenem-first 体系'),
    ('md/output_van_first/complex_van_first_short.prmtop / .inpcrd', 'Vancomycin-first 体系'),
    ('md/output/short_equil.dcd', 'Meropenem-first 100 ps 轨迹'),
    ('md/output_van_first/short_equil_van_first.dcd', 'Vancomycin-first 100 ps 轨迹'),
    ('md/output/mmgbsa_meropenem_first_RESULTS.dat', 'Meropenem-first MM-GBSA 结果'),
    ('md/output_van_first/mmgbsa_vancomycin_first_RESULTS.dat', 'Vancomycin-first MM-GBSA 结果'),
    ('md/scripts/run_mmgbsa.sh', 'MM-GBSA 批量脚本'),
    ('04_reports/figures/distances_100ps.png', 'Meropenem-first 距离图'),
    ('04_reports/figures/distances_van_first_100ps.png', 'Vancomycin-first 距离图'),
    ('04_reports/figures/comparison_distances_100ps.png', '对称比较图'),
    ('04_reports/figures/cross_comparison_summary.png', '交叉对比汇总图'),
    ('04_reports/figures/van_first_*.png', '附录 A Vancomycin-first 对接可视化图'),
    ('04_reports/6UJN_Meropenem_Vancomycin_validation_report.md', 'Markdown 报告（含附录 A）'),
    ('04_reports/6UJN_Meropenem_Vancomycin_validation_report.docx', 'Word 报告（含附录 A，本文件）'),
]
file_table = doc.add_table(rows=len(files), cols=2)
file_table.style = 'Light Grid Accent 1'
for i, (path, desc) in enumerate(files):
    file_table.rows[i].cells[0].text = path
    file_table.rows[i].cells[1].text = desc

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('报告更新时间：2026-06-26（MM-GBSA 能量学补全 + Vancomycin-first 对接交叉对比整合后）')
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Save (use v2/v3 suffix if original is locked/open in Word)
output_path = REPORT_DIR / '6UJN_Meropenem_Vancomycin_validation_report.docx'
try:
    doc.save(output_path)
except PermissionError:
    output_path = REPORT_DIR / '6UJN_Meropenem_Vancomycin_validation_report_v3.docx'
    doc.save(output_path)
print(f"Saved {output_path}")
