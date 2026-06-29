#!/usr/bin/env python3
"""Generate Vancomycin-first -> Meropenem docking report and figures."""

import csv
import math
from pathlib import Path
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from mpl_toolkits.mplot3d import Axes3D

WORK = Path(r"D:\6UJN_Meropenem_Vancomycin")
REPORT_DIR = WORK / "04_reports"
FIG_DIR = REPORT_DIR / "figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)

# Source files
VAN_FIRST_DIR = Path(r"D:\Tools\Docking\docking_project\data\receptors\6UJN_Vancomycin")
MERO_DOCKED = VAN_FIRST_DIR / "Meropenem_out.pdbqt"
VAN_FIRST_CSV = WORK / "input" / "6UJN_Vancomycin_Meropenem.csv"
APO_MERO = WORK / "input" / "Meropenem_out.pdbqt"
# Prepared ternary complex PDB contains protein + Vancomycin
COMPLEX_PDB = WORK / "md" / "input_van_first" / "complex.pdb"


def parse_pdbqt_atoms(path, model=None):
    """Parse ATOM/HETATM records from PDBQT. If model is None, parse all models sequentially."""
    atoms = []
    current_model = 0
    in_model = False
    has_models = False
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            if line.startswith("MODEL"):
                has_models = True
                current_model += 1
                in_model = True
                continue
            if line.startswith("ENDMDL"):
                in_model = False
                continue
            if has_models and not in_model:
                continue
            record = line[:6]
            if record in ("ATOM  ", "HETATM"):
                atoms.append({
                    "atom_name": line[12:16].strip(),
                    "res_name": line[17:20].strip(),
                    "chain": line[21].strip(),
                    "res_seq": int(line[22:26]),
                    "x": float(line[30:38]),
                    "y": float(line[38:46]),
                    "z": float(line[46:54]),
                    "atom_type": line[77:79].strip(),
                    "model": current_model if has_models else 1,
                })
    return atoms


def parse_pdb_atoms(path):
    """Parse ATOM/HETATM records from a standard PDB file."""
    atoms = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            record = line[:6]
            if record in ("ATOM  ", "HETATM"):
                elem = line[76:78].strip()
                if not elem:
                    elem = line[12:16].strip()[0]
                atoms.append({
                    "atom_name": line[12:16].strip(),
                    "res_name": line[17:20].strip(),
                    "chain": line[21].strip(),
                    "res_seq": int(line[22:26]),
                    "x": float(line[30:38]),
                    "y": float(line[38:46]),
                    "z": float(line[46:54]),
                    "atom_type": elem,
                    "model": 1,
                })
    return atoms


def parse_vina_csv(path):
    modes = []
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            modes.append({
                "mode": int(row.get("Mode", row.get("mode", 0))),
                "affinity": float(row["Binding Affinity"]),
                "rmsd_lb": float(row.get("rmsd/lb", 0)),
                "rmsd_ub": float(row.get("rmsd/ub", 0)),
            })
    return modes


def centroid(atoms):
    n = len(atoms)
    return (sum(a["x"] for a in atoms) / n,
            sum(a["y"] for a in atoms) / n,
            sum(a["z"] for a in atoms) / n)


def is_heavy(atom):
    return atom["atom_type"] not in ("H", "HD", "HS")


def distance(a, b):
    return math.sqrt((a["x"] - b["x"]) ** 2 + (a["y"] - b["y"]) ** 2 + (a["z"] - b["z"]) ** 2)


def extract_model(atoms, model):
    return [a for a in atoms if a["model"] == model]


def rmsd(atoms1, atoms2):
    m1 = {a["atom_name"]: a for a in atoms1}
    m2 = {a["atom_name"]: a for a in atoms2}
    common = set(m1.keys()) & set(m2.keys())
    if not common:
        return None
    s = sum(distance(m1[n], m2[n]) ** 2 for n in common)
    return math.sqrt(s / len(common))


def main():
    # Parse receptor (protein + Vancomycin) from prepared complex PDB
    complex_atoms = parse_pdb_atoms(COMPLEX_PDB)
    protein_ca = [a for a in complex_atoms if a["atom_name"] == "CA"]
    van_atoms = [a for a in complex_atoms if a["res_name"] == "VAN"]
    van_heavy = [a for a in van_atoms if is_heavy(a)]

    # Parse all Meropenem docking models
    mero_atoms = parse_pdbqt_atoms(MERO_DOCKED)
    models = sorted(set(a["model"] for a in mero_atoms))

    # Parse apo Meropenem pose (model 1)
    apo_mero_atoms = parse_pdbqt_atoms(APO_MERO, model=1)
    apo_mero_heavy = [a for a in apo_mero_atoms if is_heavy(a)]

    # CSV data
    csv_modes = parse_vina_csv(VAN_FIRST_CSV)

    # Compute centroids and distances
    van_centroid = centroid(van_heavy)
    mero_data = []
    for m in models:
        m_atoms = extract_model(mero_atoms, m)
        m_heavy = [a for a in m_atoms if is_heavy(a)]
        c = centroid(m_heavy)
        d = math.sqrt((c[0] - van_centroid[0]) ** 2 +
                      (c[1] - van_centroid[1]) ** 2 +
                      (c[2] - van_centroid[2]) ** 2)
        rmsd_apo = rmsd(m_heavy, apo_mero_heavy)
        aff = csv_modes[m - 1]["affinity"] if m - 1 < len(csv_modes) else None
        mero_data.append({
            "model": m,
            "centroid": c,
            "atoms": m_atoms,
            "heavy": m_heavy,
            "dist_to_van": d,
            "rmsd_apo": rmsd_apo,
            "affinity": aff,
        })

    # Figure 1: 3D overview of complex with top pose
    fig = plt.figure(figsize=(10, 8))
    ax = fig.add_subplot(111, projection='3d')
    px = [a["x"] for a in protein_ca]
    py = [a["y"] for a in protein_ca]
    pz = [a["z"] for a in protein_ca]
    ax.scatter(px, py, pz, c='lightgray', s=5, alpha=0.5, label='Protein CA')
    vx = [a["x"] for a in van_heavy]
    vy = [a["y"] for a in van_heavy]
    vz = [a["z"] for a in van_heavy]
    ax.scatter(vx, vy, vz, c='blue', s=30, alpha=0.8, label='Vancomycin')
    top = mero_data[0]
    mx = [a["x"] for a in top["heavy"]]
    my = [a["y"] for a in top["heavy"]]
    mz = [a["z"] for a in top["heavy"]]
    ax.scatter(mx, my, mz, c='red', s=40, alpha=0.9,
               label=f"Meropenem (top, {top['affinity']:.2f} kcal/mol)")
    ax.set_xlabel('X (A)')
    ax.set_ylabel('Y (A)')
    ax.set_zlabel('Z (A)')
    ax.set_title('Vancomycin-first complex: top Meropenem docking pose')
    ax.legend()
    fig.tight_layout()
    fig.savefig(FIG_DIR / 'van_first_complex_overview.png', dpi=300)
    plt.close(fig)

    # Figure 2: All Meropenem poses colored by affinity
    fig = plt.figure(figsize=(10, 8))
    ax = fig.add_subplot(111, projection='3d')
    ax.scatter(px, py, pz, c='lightgray', s=3, alpha=0.3, label='Protein CA')
    ax.scatter(vx, vy, vz, c='blue', s=25, alpha=0.7, label='Vancomycin')
    cmap = plt.cm.RdYlGn_r
    affs = [d["affinity"] for d in mero_data]
    norm = plt.Normalize(vmin=min(affs), vmax=max(affs))
    for d in mero_data:
        mx = [a["x"] for a in d["heavy"]]
        my = [a["y"] for a in d["heavy"]]
        mz = [a["z"] for a in d["heavy"]]
        color = cmap(norm(d["affinity"]))
        ax.scatter(mx, my, mz, c=[color], s=25, alpha=0.7)
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=ax, shrink=0.6)
    cbar.set_label('Binding affinity (kcal/mol)')
    ax.set_xlabel('X (A)')
    ax.set_ylabel('Y (A)')
    ax.set_zlabel('Z (A)')
    ax.set_title('All 10 Meropenem docking poses (Vancomycin-first)')
    fig.tight_layout()
    fig.savefig(FIG_DIR / 'van_first_docking_poses.png', dpi=300)
    plt.close(fig)

    # Figure 3: Distance to Vancomycin vs affinity
    fig, ax = plt.subplots(figsize=(8, 6))
    mode_nums = [d["model"] for d in mero_data]
    dists = [d["dist_to_van"] for d in mero_data]
    colors = ['red' if d == 1 else 'steelblue' for d in mode_nums]
    ax.scatter(dists, affs, c=colors, s=80, alpha=0.8)
    for i, m in enumerate(mode_nums):
        ax.annotate(f'M{m}', (dists[i], affs[i]), textcoords="offset points",
                    xytext=(5, 5), fontsize=8)
    ax.set_xlabel('Meropenem centroid - Vancomycin centroid distance (A)')
    ax.set_ylabel('Binding affinity (kcal/mol)')
    ax.set_title('Vancomycin-first: Meropenem pose distance vs affinity')
    ax.axhline(y=min(affs), color='gray', linestyle='--', alpha=0.3)
    ax.grid(True, alpha=0.3)
    fig.tight_layout()
    fig.savefig(FIG_DIR / 'van_first_distance_vs_affinity.png', dpi=300)
    plt.close(fig)

    # Figure 4: 2D projection (XY)
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.scatter([a["x"] for a in protein_ca], [a["y"] for a in protein_ca],
               c='lightgray', s=3, alpha=0.4, label='Protein CA')
    ax.scatter([a["x"] for a in van_heavy], [a["y"] for a in van_heavy],
               c='blue', s=25, alpha=0.7, label='Vancomycin')
    for i, d in enumerate(mero_data):
        mx = [a["x"] for a in d["heavy"]]
        my = [a["y"] for a in d["heavy"]]
        c = 'red' if i == 0 else cmap(norm(d["affinity"]))
        ax.scatter(mx, my, c=[c], s=15, alpha=0.6)
    ax.set_xlabel('X (A)')
    ax.set_ylabel('Y (A)')
    ax.set_title('XY projection: Meropenem poses relative to Vancomycin')
    ax.legend()
    ax.set_aspect('equal', adjustable='box')
    fig.tight_layout()
    fig.savefig(FIG_DIR / 'van_first_xy_projection.png', dpi=300)
    plt.close(fig)

    print(f"Saved figures to {FIG_DIR}")

    # Generate Markdown report
    md_path = REPORT_DIR / "Vancomycin_first_docking_report.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("# Vancomycin-first -> Meropenem Molecular Docking Supplement Report\n\n")
        f.write("## 1. Objective\n\n")
        f.write("Dock Vancomycin onto 6UJN first, then dock Meropenem onto the resulting holo receptor. ")
        f.write("This report documents the Vancomycin-first sequential docking results and visualization.\n\n")
        f.write("## 2. Input Files and Parameters\n\n")
        f.write("| Item | Path / Value |\n")
        f.write("|---|---|\n")
        f.write(f"| Receptor (holo, visualization) | `{COMPLEX_PDB}` |\n")
        f.write(f"| Docked ligand poses | `{MERO_DOCKED}` |\n")
        f.write(f"| Docking result CSV | `{VAN_FIRST_CSV}` |\n")
        f.write("| Docking center (A) | (56.0328, 17.1666, 25.8216) |\n")
        f.write("| Box size (A) | 60.68 x 60.68 x 60.68 |\n")
        f.write("| Exhaustiveness | 8 |\n")
        f.write("| num_modes | 10 |\n")
        f.write("| seed | 42 |\n\n")
        f.write("## 3. Docking Results\n\n")
        f.write("| Mode | Binding Affinity (kcal/mol) | RMSD/lb | RMSD/ub | Distance to VAN centroid (A) | vs apo Meropenem RMSD (A) |\n")
        f.write("|---|---|---|---|---|---|\n")
        for d in mero_data:
            rmsd_str = f"{d['rmsd_apo']:.3f}" if d['rmsd_apo'] is not None else "N/A"
            f.write(f"| {d['model']} | {d['affinity']:.3f} | {csv_modes[d['model']-1]['rmsd_lb']:.3f} | {csv_modes[d['model']-1]['rmsd_ub']:.3f} | {d['dist_to_van']:.2f} | {rmsd_str} |\n")
        f.write("\n")
        f.write(f"**Best affinity**: {mero_data[0]['affinity']:.3f} kcal/mol (Mode 1).\n\n")
        f.write("## 4. Key Observations\n\n")
        f.write(f"1. **Meropenem stays far from Vancomycin**: the best pose (Mode 1) centroid is **{mero_data[0]['dist_to_van']:.2f} A** from the Vancomycin centroid.\n")
        f.write(f"2. **Dispersed poses**: the 10 poses span **{min(d['dist_to_van'] for d in mero_data):.1f}-{max(d['dist_to_van'] for d in mero_data):.1f} A** from Vancomycin, indicating Meropenem does not cluster near the Vancomycin site.\n")
        f.write(f"3. **Large difference from apo pose**: heavy-atom RMSD vs apo Meropenem is **{mero_data[0]['rmsd_apo']:.2f} A**, showing a distinct binding mode.\n")
        f.write("4. **Energetic implication**: the best affinity (-6.738 kcal/mol) is close to the apo Meropenem reference (-6.788 kcal/mol), ")
        f.write("so Vancomycin pre-binding does not promote Meropenem binding to its own site.\n\n")
        f.write("## 5. Visualizations\n\n")
        f.write("### 5.1 Complex Overview (best pose)\n\n")
        f.write("![Vancomycin-first complex overview](figures/van_first_complex_overview.png)\n\n")
        f.write("*Figure 1. Protein CA (gray), Vancomycin (blue), and best Meropenem pose (red).*\n\n")
        f.write("### 5.2 All 10 Meropenem Docking Poses\n\n")
        f.write("![Vancomycin-first docking poses](figures/van_first_docking_poses.png)\n\n")
        f.write("*Figure 2. Ten Meropenem poses colored by affinity; Vancomycin in blue.*\n\n")
        f.write("### 5.3 Meropenem-Vancomycin Distance vs Affinity\n\n")
        f.write("![Distance vs affinity](figures/van_first_distance_vs_affinity.png)\n\n")
        f.write("*Figure 3. Distance from each Meropenem pose centroid to Vancomycin centroid versus docking affinity.*\n\n")
        f.write("### 5.4 XY Plane Projection\n\n")
        f.write("![XY projection](figures/van_first_xy_projection.png)\n\n")
        f.write("*Figure 4. XY projection showing Meropenem poses distributed on the protein surface, away from Vancomycin.*\n\n")
        f.write("## 6. Conclusion\n\n")
        f.write("> In the Vancomycin-first sequential docking, Meropenem does not occupy the pocket near Vancomycin; ")
        f.write("instead it binds at sites remote from Vancomycin. The best affinity (-6.738 kcal/mol) is almost identical to the apo reference, ")
        f.write("indicating no significant promoting effect of Vancomycin pre-binding on Meropenem. ")
        f.write("This is consistent with the subsequent symmetric MD relaxation (COM ~46.76 A, MM-GBSA dG = -23.03 kcal/mol).\n\n")
        f.write("---\n\n")
        f.write("*Report generated: 2026-06-26*\n")

    print(f"Saved Markdown report: {md_path}")

    # Generate Chinese Markdown report
    md_cn_path = REPORT_DIR / "Vancomycin_first_docking_report_CN.md"
    with open(md_cn_path, "w", encoding="utf-8") as f:
        f.write("# 万古霉素先结合（Vancomycin-first）→ 美罗培南分子对接补充报告\n\n")
        f.write("## 1. 实验目的\n\n")
        f.write("在 6UJN 受体上先对接 Vancomycin，再在其形成的 holo 受体上对接 Meropenem，")
        f.write("以获取 Vancomycin-first 序贯对接的详细结果与可视化。\n\n")
        f.write("## 2. 输入文件与参数\n\n")
        f.write("| 项目 | 路径 / 数值 |\n")
        f.write("|---|---|\n")
        f.write(f"| 受体（holo，可视化用） | `{COMPLEX_PDB}` |\n")
        f.write(f"| 对接配体 poses | `{MERO_DOCKED}` |\n")
        f.write(f"| 对接结果 CSV | `{VAN_FIRST_CSV}` |\n")
        f.write("| 对接中心 (Å) | (56.0328, 17.1666, 25.8216) |\n")
        f.write("| 盒子大小 (Å) | 60.68 × 60.68 × 60.68 |\n")
        f.write("| Exhaustiveness | 8 |\n")
        f.write("| num_modes | 10 |\n")
        f.write("| seed | 42 |\n\n")
        f.write("## 3. 对接结果\n\n")
        f.write("| Mode | Binding Affinity (kcal/mol) | RMSD/lb | RMSD/ub | 到 VAN 质心距离 (Å) | vs apo Meropenem RMSD (Å) |\n")
        f.write("|---|---|---|---|---|---|\n")
        for d in mero_data:
            rmsd_str = f'{d["rmsd_apo"]:.3f}' if d['rmsd_apo'] is not None else 'N/A'
            f.write(f'| {d["model"]} | {d["affinity"]:.3f} | {csv_modes[d["model"]-1]["rmsd_lb"]:.3f} | {csv_modes[d["model"]-1]["rmsd_ub"]:.3f} | {d["dist_to_van"]:.2f} | {rmsd_str} |\n')
        f.write("\n")
        f.write(f'**最佳亲和力**: {mero_data[0]["affinity"]:.3f} kcal/mol（Mode 1）。\n\n')
        f.write("## 4. 关键观察\n\n")
        f.write(f'1. **Meropenem 远离 Vancomycin**：最佳 pose（Mode 1）的质心距 Vancomycin 质心约 **{mero_data[0]["dist_to_van"]:.2f} Å**。\n')
        f.write(f'2. **多模式分散**：10 个 pose 距 Vancomycin 的范围约 **{min(d["dist_to_van"] for d in mero_data):.1f}–{max(d["dist_to_van"] for d in mero_data):.1f} Å**，说明 Meropenem 未聚集于 Vancomycin 位点。\n')
        f.write(f'3. **与 apo Meropenem 差异大**：最佳 pose 与 apo Meropenem（6UJN 直接对接）的重原子 RMSD 为 **{mero_data[0]["rmsd_apo"]:.2f} Å**，结合模式显著不同。\n')
        f.write("4. **能量意义**：最佳亲和力 −6.738 kcal/mol 与 apo Meropenem 参考能 −6.788 kcal/mol 接近（ΔE ≈ +0.050 kcal/mol），")
        f.write("表明 Vancomycin 预结合几乎未提升 Meropenem 在其位点的结合。\n\n")
        f.write("## 5. 可视化\n\n")
        f.write("### 5.1 复合物总览（最佳 pose）\n\n")
        f.write("![Vancomycin-first complex overview](figures/van_first_complex_overview.png)\n\n")
        f.write("*图 1. 蛋白 CA（灰）、Vancomycin（蓝）与最佳 Meropenem pose（红）。*\n\n")
        f.write("### 5.2 全部 10 个 Meropenem 对接 pose\n\n")
        f.write("![Vancomycin-first docking poses](figures/van_first_docking_poses.png)\n\n")
        f.write("*图 2. 10 个 Meropenem pose 按亲和力着色；Vancomycin（蓝）。*\n\n")
        f.write("### 5.3 Meropenem–Vancomycin 距离 vs 亲和力\n\n")
        f.write("![Distance vs affinity](figures/van_first_distance_vs_affinity.png)\n\n")
        f.write("*图 3. 各 pose 的 Meropenem 质心到 Vancomycin 质心距离与其对接亲和力。*\n\n")
        f.write("### 5.4 XY 平面投影\n\n")
        f.write("![XY projection](figures/van_first_xy_projection.png)\n\n")
        f.write("*图 4. XY 平面投影显示 Meropenem pose 分布于蛋白表面，远离 Vancomycin。*\n\n")
        f.write("## 6. 结论\n\n")
        f.write("> Vancomycin-first 序贯对接中，Meropenem 未占据 Vancomycin 附近的口袋，")
        f.write("而是结合在远离 Vancomycin 的位点。最佳亲和力 −6.738 kcal/mol 与 apo 状态几乎相同，")
        f.write("说明 Vancomycin 预结合对 Meropenem 结合无显著促进作用。")
        f.write("这与后续对称 MD 松弛（COM ~46.76 Å，MM-GBSA ΔG = −23.03 kcal/mol）一致。\n\n")
        f.write("---\n\n")
        f.write("*报告生成时间：2026-06-26*\n")

    print(f"Saved Chinese Markdown report: {md_cn_path}")

    # Generate Word report
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        title = doc.add_heading('Vancomycin-first -> Meropenem Molecular Docking Supplement Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('1. Objective', level=1)
        doc.add_paragraph('Dock Vancomycin onto 6UJN first, then dock Meropenem onto the resulting holo receptor. '
                          'This report documents the Vancomycin-first sequential docking results and visualization.')

        doc.add_heading('2. Input Files and Parameters', level=1)
        params = [
            ('Receptor (holo, visualization)', str(COMPLEX_PDB)),
            ('Docked ligand poses', str(MERO_DOCKED)),
            ('Docking result CSV', str(VAN_FIRST_CSV)),
            ('Docking center (A)', '(56.0328, 17.1666, 25.8216)'),
            ('Box size (A)', '60.68 x 60.68 x 60.68'),
            ('Exhaustiveness', '8'),
            ('num_modes', '10'),
            ('seed', '42'),
        ]
        table = doc.add_table(rows=len(params), cols=2)
        table.style = 'Light Grid Accent 1'
        for i, (k, v) in enumerate(params):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[1].text = v

        doc.add_heading('3. Docking Results', level=1)
        res_table = doc.add_table(rows=len(mero_data) + 1, cols=6)
        res_table.style = 'Light Grid Accent 1'
        headers = ['Mode', 'Affinity', 'RMSD/lb', 'RMSD/ub', 'Dist. to VAN (A)', 'vs apo RMSD']
        for j, h in enumerate(headers):
            res_table.rows[0].cells[j].text = h
        for i, d in enumerate(mero_data):
            rmsd_str = f"{d['rmsd_apo']:.3f}" if d['rmsd_apo'] is not None else 'N/A'
            res_table.rows[i + 1].cells[0].text = str(d['model'])
            res_table.rows[i + 1].cells[1].text = f"{d['affinity']:.3f}"
            res_table.rows[i + 1].cells[2].text = f"{csv_modes[d['model']-1]['rmsd_lb']:.3f}"
            res_table.rows[i + 1].cells[3].text = f"{csv_modes[d['model']-1]['rmsd_ub']:.3f}"
            res_table.rows[i + 1].cells[4].text = f"{d['dist_to_van']:.2f}"
            res_table.rows[i + 1].cells[5].text = rmsd_str

        doc.add_paragraph(f"Best affinity: {mero_data[0]['affinity']:.3f} kcal/mol (Mode 1).")

        doc.add_heading('4. Key Observations', level=1)
        obs = [
            f"Meropenem stays far from Vancomycin: the best pose centroid is {mero_data[0]['dist_to_van']:.2f} A from the Vancomycin centroid.",
            f"Dispersed poses: the 10 poses span {min(d['dist_to_van'] for d in mero_data):.1f}-{max(d['dist_to_van'] for d in mero_data):.1f} A from Vancomycin.",
            f"Large difference from apo pose: heavy-atom RMSD vs apo Meropenem is {mero_data[0]['rmsd_apo']:.2f} A.",
            "Energetic implication: best affinity -6.738 kcal/mol is close to apo reference -6.788 kcal/mol, indicating Vancomycin pre-binding does not promote Meropenem binding."
        ]
        for o in obs:
            doc.add_paragraph(o, style='List Bullet')

        doc.add_heading('5. Visualizations', level=1)
        figs = [
            ('van_first_complex_overview.png', 'Figure 1. Protein CA (gray), Vancomycin (blue), and best Meropenem pose (red).'),
            ('van_first_docking_poses.png', 'Figure 2. Ten Meropenem poses colored by affinity.'),
            ('van_first_distance_vs_affinity.png', 'Figure 3. Meropenem-Vancomycin distance vs affinity.'),
            ('van_first_xy_projection.png', 'Figure 4. XY plane projection.'),
        ]
        for fig_name, caption in figs:
            doc.add_picture(str(FIG_DIR / fig_name), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('6. Conclusion', level=1)
        conclusion = doc.add_paragraph(
            'In the Vancomycin-first sequential docking, Meropenem does not occupy the pocket near Vancomycin; '
            'instead it binds at sites remote from Vancomycin. The best affinity (-6.738 kcal/mol) is almost identical to the apo reference, '
            'indicating no significant promoting effect of Vancomycin pre-binding on Meropenem. '
            'This is consistent with the subsequent symmetric MD relaxation (COM ~46.76 A, MM-GBSA dG = -23.03 kcal/mol).'
        )
        conclusion.runs[0].bold = True

        docx_path = REPORT_DIR / "Vancomycin_first_docking_report.docx"
        doc.save(docx_path)
        print(f"Saved Word report: {docx_path}")

        # Chinese Word report
        doc_cn = Document()
        title_cn = doc_cn.add_heading('万古霉素先结合（Vancomycin-first）→ 美罗培南分子对接补充报告', 0)
        title_cn.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_cn.add_heading('1. 实验目的', level=1)
        doc_cn.add_paragraph('在 6UJN 受体上先对接 Vancomycin，再在其形成的 holo 受体上对接 Meropenem，以获取 Vancomycin-first 序贯对接的详细结果与可视化。')

        doc_cn.add_heading('2. 输入文件与参数', level=1)
        params_cn = [
            ('受体（holo，可视化用）', str(COMPLEX_PDB)),
            ('对接配体 poses', str(MERO_DOCKED)),
            ('对接结果 CSV', str(VAN_FIRST_CSV)),
            ('对接中心 (Å)', '(56.0328, 17.1666, 25.8216)'),
            ('盒子大小 (Å)', '60.68 × 60.68 × 60.68'),
            ('Exhaustiveness', '8'),
            ('num_modes', '10'),
            ('seed', '42'),
        ]
        table_cn = doc_cn.add_table(rows=len(params_cn), cols=2)
        table_cn.style = 'Light Grid Accent 1'
        for i, (k, v) in enumerate(params_cn):
            table_cn.rows[i].cells[0].text = k
            table_cn.rows[i].cells[1].text = v

        doc_cn.add_heading('3. 对接结果', level=1)
        res_table_cn = doc_cn.add_table(rows=len(mero_data) + 1, cols=6)
        res_table_cn.style = 'Light Grid Accent 1'
        headers_cn = ['Mode', '亲和力', 'RMSD/lb', 'RMSD/ub', '到 VAN 距离 (Å)', 'vs apo RMSD']
        for j, h in enumerate(headers_cn):
            res_table_cn.rows[0].cells[j].text = h
        for i, d in enumerate(mero_data):
            rmsd_str = f'{d["rmsd_apo"]:.3f}' if d['rmsd_apo'] is not None else 'N/A'
            res_table_cn.rows[i + 1].cells[0].text = str(d['model'])
            res_table_cn.rows[i + 1].cells[1].text = f'{d["affinity"]:.3f}'
            res_table_cn.rows[i + 1].cells[2].text = f'{csv_modes[d["model"]-1]["rmsd_lb"]:.3f}'
            res_table_cn.rows[i + 1].cells[3].text = f'{csv_modes[d["model"]-1]["rmsd_ub"]:.3f}'
            res_table_cn.rows[i + 1].cells[4].text = f'{d["dist_to_van"]:.2f}'
            res_table_cn.rows[i + 1].cells[5].text = rmsd_str

        doc_cn.add_paragraph(f'最佳亲和力: {mero_data[0]["affinity"]:.3f} kcal/mol（Mode 1）。')

        doc_cn.add_heading('4. 关键观察', level=1)
        obs_cn = [
            f'Meropenem 远离 Vancomycin：最佳 pose 质心距 Vancomycin 质心约 {mero_data[0]["dist_to_van"]:.2f} Å。',
            f'多模式分散：10 个 pose 距 Vancomycin 的范围约 {min(d["dist_to_van"] for d in mero_data):.1f}-{max(d["dist_to_van"] for d in mero_data):.1f} Å。',
            f'与 apo Meropenem 差异大：最佳 pose 与 apo Meropenem 的重原子 RMSD 为 {mero_data[0]["rmsd_apo"]:.2f} Å。',
            '能量意义：最佳亲和力 −6.738 kcal/mol 与 apo 参考 −6.788 kcal/mol 接近，说明 Vancomycin 预结合未促进 Meropenem 结合。'
        ]
        for o in obs_cn:
            doc_cn.add_paragraph(o, style='List Bullet')

        doc_cn.add_heading('5. 可视化', level=1)
        figs_cn = [
            ('van_first_complex_overview.png', '图 1. 蛋白 CA（灰）、Vancomycin（蓝）与最佳 Meropenem pose（红）。'),
            ('van_first_docking_poses.png', '图 2. 10 个 Meropenem pose 按亲和力着色。'),
            ('van_first_distance_vs_affinity.png', '图 3. Meropenem–Vancomycin 距离 vs 亲和力。'),
            ('van_first_xy_projection.png', '图 4. XY 平面投影。'),
        ]
        for fig_name, caption in figs_cn:
            doc_cn.add_picture(str(FIG_DIR / fig_name), width=Inches(5.5))
            doc_cn.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc_cn.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_cn.add_heading('6. 结论', level=1)
        conclusion_cn = doc_cn.add_paragraph(
            'Vancomycin-first 序贯对接中，Meropenem 未占据 Vancomycin 附近的口袋，'
            '而是结合在远离 Vancomycin 的位点。最佳亲和力 −6.738 kcal/mol 与 apo 状态几乎相同，'
            '说明 Vancomycin 预结合对 Meropenem 结合无显著促进作用。'
            '这与后续对称 MD 松弛（COM ~46.76 Å，MM-GBSA ΔG = −23.03 kcal/mol）一致。'
        )
        conclusion_cn.runs[0].bold = True

        docx_cn_path = REPORT_DIR / "Vancomycin_first_docking_report_CN.docx"
        doc_cn.save(docx_cn_path)
        print(f"Saved Chinese Word report: {docx_cn_path}")
    except Exception as e:
        print(f"Word report generation failed: {e}")


if __name__ == "__main__":
    main()
